import io
import os
import time
import sqlite3
import json
import matplotlib.pyplot as plt
import pandas as pd
import seaborn as sns
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import base64
from langchain.agents import AgentExecutor, create_tool_calling_agent
from langchain_core.output_parsers import PydanticOutputParser
from typing import List, Optional, Dict, Any, Union
from pydantic import BaseModel, Field
from langchain_core.exceptions import OutputParserException
from langchain_community.utilities import SQLDatabase
from sqlalchemy import create_engine, text
from langchain_community.agent_toolkits.sql.toolkit import SQLDatabaseToolkit
from langchain_community.agent_toolkits import create_sql_agent
from tenacity import retry, wait_exponential, stop_after_attempt, retry_if_exception_type
from langchain_openai import AzureChatOpenAI
from langchain_core.tools import tool
from langchain_core.prompts import ChatPromptTemplate, MessagesPlaceholder
import logging
import warnings
import numpy as np
from datetime import datetime
import uuid
import re

# Document processing imports with fallbacks
try:
    import PyPDF2
    import docx
    from docx import Document
    import openpyxl
    from openpyxl import load_workbook
    import xlrd
    import textract
    import fitz  # PyMuPDF for better PDF handling
    PDF_SUPPORT = True
except ImportError as e:
    print(f"⚠️ Some document processing libraries are missing: {e}")
    print("📦 Install with: pip install PyPDF2 python-docx openpyxl xlrd textract PyMuPDF")
    PDF_SUPPORT = False

# Suppress warnings
warnings.filterwarnings("ignore", category=UserWarning, module="langsmith")
warnings.filterwarnings("ignore", category=UserWarning, module="matplotlib")
os.environ["LANGCHAIN_TRACING_V2"] = "false"  # Disable LangSmith tracing

# Simple import solution - adjust path based on your file structure
import sys
import os

# Get the directory containing this file
current_dir = os.path.dirname(os.path.abspath(__file__))

# Add paths to find settings.py and services
sys.path.append(current_dir)  # Current directory
sys.path.append(os.path.join(current_dir, 'config'))  # config subdirectory
sys.path.append(os.path.dirname(current_dir))  # Parent directory
sys.path.append(os.path.join(os.path.dirname(current_dir), 'services'))  # services directory

try:
    from settings import get_settings
except ImportError:
    try:
        from config.settings import get_settings
    except ImportError:
        try:
            from backend.config.settings import get_settings
        except ImportError:
            print("❌ Could not import settings. Please check your file structure.")
            print(f"Current directory: {current_dir}")
            print("Please ensure settings.py is accessible.")
            sys.exit(1)

# Import your services
try:
    from backend.services.document_classifier import DocumentClassifier
    from backend.services.schema_manager import SchemaManager
    from backend.models.document import DocumentType, BusinessDocumentType
    SERVICES_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ Services not available: {e}")
    print("🔄 Running in standalone mode without classification services")
    SERVICES_AVAILABLE = False
    
    # Create fallback enums
    from enum import Enum
    
    class DocumentType(Enum):
        PDF = "pdf"
        DOCX = "docx"
        DOC = "doc"
        TXT = "txt"
        CSV = "csv"
        JSON = "json"
        XLSX = "xlsx"
        XLS = "xls"
    
    class BusinessDocumentType(Enum):
        INVOICE = "invoice"
        PAYROLL = "payroll"
        EMPLOYEE_RECORD = "employee_record"
        CUSTOMER_DATA = "customer_data"
        INVENTORY = "inventory"
        CONTRACT = "contract"
        PURCHASE_ORDER = "purchase_order"
        SALES_ORDER = "sales_order"
        FINANCIAL_STATEMENT = "financial_statement"
        REPORT = "report"
        GENERAL = "general"

# Configure logging
logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(name)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

class IntegratedDocumentProcessor:
    """Enhanced document processor with classification, schema management, and visualization"""
    
    def __init__(self, llm=None, db_path="integrated_database.db", user_id=None):
        """Initialize the integrated processor"""
        self.settings = get_settings()
        self.user_id = user_id or "default_user"  # Use provided user_id or default
        
        # Initialize Azure OpenAI LLM if not provided
        if llm is None:
            self.llm = self._create_azure_openai_llm()
        else:
            self.llm = llm
            
        self.db_path = db_path
        self.engine = create_engine(f"sqlite:///{self.db_path}")
        self.db = SQLDatabase(engine=self.engine)
        
        # Initialize services if available
        if SERVICES_AVAILABLE:
            try:
                self.document_classifier = DocumentClassifier()
                self.schema_manager = SchemaManager()
                logger.info("✅ Document classification and schema management services loaded")
            except Exception as e:
                logger.warning(f"⚠️ Failed to initialize services: {str(e)}")
                self.document_classifier = None
                self.schema_manager = None
        else:
            self.document_classifier = None
            self.schema_manager = None
            logger.warning("⚠️ Running without classification services")
        
        # Create tools and agents
        self.tools = self._create_tools()
        self.sql_agent = self._create_enhanced_sql_agent()
        self.current_table_name = None
        self.current_document_info = {}
        self.supported_formats = self.settings.ALLOWED_FILE_TYPES
        
        # Set up visualization defaults
        plt.style.use('default')
        try:
            sns.set_palette("husl")
        except:
            pass  # If seaborn not available
        
        logger.info(f"✅ IntegratedDocumentProcessor initialized for user: {self.user_id}")
        logger.info(f"📋 Supported file types: {', '.join(self.supported_formats)}")
        logger.info(f"🔧 AI Classification: {'Enabled' if SERVICES_AVAILABLE else 'Disabled'}")
    
    def _create_azure_openai_llm(self):
        """Create Azure OpenAI LLM instance"""
        try:
            azure_openai_config = self.settings.get_azure_openai_config()
            
            llm = AzureChatOpenAI(
                azure_endpoint=azure_openai_config["azure_endpoint"],
                api_key=azure_openai_config["api_key"],
                api_version=azure_openai_config["api_version"],
                azure_deployment=azure_openai_config["deployment_name"],
                model=azure_openai_config["model_name"],
                temperature=0.1,  # Lower temperature for consistent analysis
                max_tokens=azure_openai_config["max_tokens"]
            )
            
            logger.info("✅ Azure OpenAI LLM created successfully")
            return llm
            
        except Exception as e:
            logger.error(f"❌ Failed to create Azure OpenAI LLM: {str(e)}")
            raise
    
    def _sanitize_table_name(self, name: str) -> str:
        """Sanitize table name to be SQL-safe"""
        # Convert to lowercase
        name = name.lower()
        
        # Replace spaces and special characters with underscores
        name = re.sub(r'[^\w]', '_', name)
        
        # Remove multiple consecutive underscores
        name = re.sub(r'_+', '_', name)
        
        # Remove leading/trailing underscores
        name = name.strip('_')
        
        # Ensure it doesn't start with a number
        if name and name[0].isdigit():
            name = f"table_{name}"
        
        # Limit length
        if len(name) > 50:
            name = name[:50]
        
        # Ensure it's not empty
        if not name:
            name = "processed_table"
        
        return name
    
    def _extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file"""
        if not PDF_SUPPORT:
            return "PDF processing not available - install required libraries"
            
        try:
            # Try PyMuPDF first (better OCR and formatting)
            try:
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                return text
            except:
                # Fallback to PyPDF2
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    text = ""
                    for page in pdf_reader.pages:
                        text += page.extract_text()
                return text
        except Exception as e:
            logger.warning(f"⚠️ PDF extraction failed: {str(e)}")
            return f"Error extracting PDF content: {str(e)}"
    
    def _extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        if not PDF_SUPPORT:
            return "DOCX processing not available - install required libraries"
            
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text
        except Exception as e:
            logger.warning(f"⚠️ DOCX extraction failed: {str(e)}")
            return f"Error extracting DOCX content: {str(e)}"
    
    def _extract_text_from_doc(self, file_path: str) -> str:
        """Extract text from DOC file using textract"""
        if not PDF_SUPPORT:
            return "DOC processing not available - install required libraries"
            
        try:
            text = textract.process(file_path).decode('utf-8')
            return text
        except Exception as e:
            logger.warning(f"⚠️ DOC extraction failed: {str(e)}")
            return f"Error extracting DOC content: {str(e)}"
    
    def _extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        return file.read()
                except UnicodeDecodeError:
                    continue
            return "Error: Could not decode text file with supported encodings"
        except Exception as e:
            logger.warning(f"⚠️ TXT extraction failed: {str(e)}")
            return f"Error extracting TXT content: {str(e)}"
    
    def _load_json_file(self, file_path: str) -> Union[pd.DataFrame, str]:
        """Load JSON file as DataFrame with robust error handling"""
        try:
            # Strategy 1: Try standard JSON loading
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    data = json.load(file)
                return self._json_to_dataframe(data)
            except json.JSONDecodeError as e:
                logger.warning(f"Standard JSON parsing failed: {str(e)}")
                
            # Strategy 2: Try loading as JSONL (JSON Lines) - multiple JSON objects
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read().strip()
                    
                # Check if it's multiple JSON objects (JSONL format)
                lines = content.split('\n')
                json_objects = []
                
                for line in lines:
                    line = line.strip()
                    if line:  # Skip empty lines
                        try:
                            obj = json.loads(line)
                            json_objects.append(obj)
                        except json.JSONDecodeError:
                            # Try to fix common JSON issues
                            fixed_line = self._fix_json_line(line)
                            if fixed_line:
                                try:
                                    obj = json.loads(fixed_line)
                                    json_objects.append(obj)
                                except:
                                    continue
                
                if json_objects:
                    logger.info(f"Successfully parsed {len(json_objects)} JSON objects from JSONL format")
                    return self._json_to_dataframe(json_objects)
                    
            except Exception as e:
                logger.warning(f"JSONL parsing failed: {str(e)}")
            
            # Strategy 3: Try to fix malformed JSON
            try:
                with open(file_path, 'r', encoding='utf-8') as file:
                    content = file.read().strip()
                
                # Try to fix common JSON issues
                fixed_content = self._fix_malformed_json(content)
                if fixed_content:
                    data = json.loads(fixed_content)
                    logger.info("Successfully parsed fixed JSON")
                    return self._json_to_dataframe(data)
                    
            except Exception as e:
                logger.warning(f"JSON fixing failed: {str(e)}")
            
            # Strategy 4: Try different encodings
            encodings = ['utf-8-sig', 'latin-1', 'cp1252', 'utf-16']
            for encoding in encodings:
                try:
                    with open(file_path, 'r', encoding=encoding) as file:
                        data = json.load(file)
                    logger.info(f"Successfully parsed JSON with {encoding} encoding")
                    return self._json_to_dataframe(data)
                except:
                    continue
            
            # Strategy 5: Last resort - treat as text and extract what we can
            logger.warning("All JSON parsing strategies failed, extracting as text content")
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as file:
                content = file.read()
            
            return pd.DataFrame([{
                "content": content[:1000],  # First 1000 chars
                "file_type": "json_error",
                "filename": os.path.basename(file_path),
                "parsing_error": "Could not parse as valid JSON",
                "content_length": len(content)
            }])
                
        except Exception as e:
            logger.error(f"⚠️ JSON loading completely failed: {str(e)}")
            return f"Error loading JSON: {str(e)}"
    
    def _fix_json_line(self, line: str) -> str:
        """Try to fix common JSON line issues"""
        try:
            # Remove trailing commas
            line = line.rstrip(',').strip()
            
            # Try to add missing brackets
            if not line.startswith('{') and not line.startswith('['):
                if '=' in line and ':' not in line:
                    # Convert key=value to {"key": "value"}
                    parts = line.split('=')
                    if len(parts) == 2:
                        key, value = parts
                        return '{"' + key.strip() + '": "' + value.strip() + '"}'
            
            return line
            
        except:
            return None
    
    def _fix_malformed_json(self, content: str) -> str:
        """Try to fix malformed JSON content"""
        try:
            # Remove BOM if present
            if content.startswith('\ufeff'):
                content = content[1:]
            
            # Fix trailing commas in objects and arrays
            content = re.sub(r',(\s*[}\]])', r'\1', content)
            
            # Try to wrap multiple objects in an array
            lines = content.strip().split('\n')
            if len(lines) > 1:
                # Check if each line might be a separate JSON object
                json_objects = []
                for line in lines:
                    line = line.strip().rstrip(',')
                    if line and (line.startswith('{') or line.startswith('[')):
                        json_objects.append(line)
                
                if len(json_objects) > 1:
                    # Wrap in array
                    return '[' + ','.join(json_objects) + ']'
            
            return content
            
        except:
            return None
    
    def _flatten_dict(self, nested_dict: dict, separator: str = '_', max_depth: int = 3) -> dict:
        """Flatten a nested dictionary"""
        def _flatten_recursive(obj, parent_key='', depth=0):
            if depth >= max_depth:
                return {parent_key[:-1]: str(obj)}
            
            items = []
            if isinstance(obj, dict):
                for k, v in obj.items():
                    new_key = f"{parent_key}{k}{separator}"
                    if isinstance(v, (dict, list)) and depth < max_depth - 1:
                        items.extend(_flatten_recursive(v, new_key, depth + 1).items())
                    else:
                        items.append((new_key[:-1], v))
            elif isinstance(obj, list):
                for i, v in enumerate(obj[:10]):  # Limit to first 10 items
                    new_key = f"{parent_key}{i}{separator}"
                    if isinstance(v, (dict, list)) and depth < max_depth - 1:
                        items.extend(_flatten_recursive(v, new_key, depth + 1).items())
                    else:
                        items.append((new_key[:-1], v))
            else:
                items.append((parent_key[:-1], obj))
            
            return dict(items)
        
        return _flatten_recursive(nested_dict)
    
    def _json_to_dataframe(self, data: Any) -> pd.DataFrame:
        """Convert JSON data to DataFrame with better handling"""
        try:
            if isinstance(data, list):
                if not data:  # Empty list
                    return pd.DataFrame([{"message": "Empty JSON array"}])
                
                # Check if it's a list of dictionaries
                if all(isinstance(item, dict) for item in data):
                    return pd.DataFrame(data)
                else:
                    # Mixed types or non-dict items
                    processed_data = []
                    for i, item in enumerate(data):
                        if isinstance(item, dict):
                            processed_data.append(item)
                        else:
                            processed_data.append({"item_" + str(i): str(item)})
                    return pd.DataFrame(processed_data)
                    
            elif isinstance(data, dict):
                # Single object
                if not data:  # Empty dict
                    return pd.DataFrame([{"message": "Empty JSON object"}])
                
                # Check if it's a simple flat dictionary
                if all(isinstance(v, (str, int, float, bool, type(None))) for v in data.values()):
                    return pd.DataFrame([data])
                else:
                    # Complex nested structure - try to normalize
                    try:
                        return pd.json_normalize(data)
                    except:
                        # If normalization fails, flatten manually
                        flattened = self._flatten_dict(data)
                        return pd.DataFrame([flattened])
            else:
                # Single value or other type
                return pd.DataFrame([{"content": str(data), "type": type(data).__name__}])
                
        except Exception as e:
            logger.warning(f"JSON to DataFrame conversion failed: {str(e)}")
            return pd.DataFrame([{"content": str(data)[:500], "conversion_error": str(e)}])
    
    def _detect_json_format(self, file_path: str) -> str:
        """Detect the format of JSON file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                first_line = file.readline().strip()
                file.seek(0)
                second_line = file.readlines()[1].strip() if len(file.readlines()) > 1 else ""
            
            # Reset file pointer and read content
            with open(file_path, 'r', encoding='utf-8') as file:
                content = file.read().strip()
            
            # Detect format patterns
            if content.startswith('[') and content.endswith(']'):
                return "json_array"
            elif content.startswith('{') and content.endswith('}'):
                if '\n{' in content:
                    return "jsonl"  # Multiple JSON objects
                else:
                    return "json_object"
            elif first_line.startswith('{') and second_line.startswith('{'):
                return "jsonl"  # JSON Lines format
            else:
                return "unknown"
                
        except Exception as e:
            logger.warning(f"Could not detect JSON format: {str(e)}")
            return "unknown"
    
    def _parse_complex_json_structure(self, data: Any, max_depth: int = 3) -> pd.DataFrame:
        """Parse complex nested JSON structures"""
        try:
            def flatten_json(nested_json: dict, separator: str = '_', max_depth: int = 3, current_depth: int = 0) -> dict:
                """Flatten nested JSON with depth limit"""
                out = {}
                
                if current_depth >= max_depth:
                    return {str(current_depth): str(nested_json)}
                
                def flatten(x, name='', depth=0):
                    if depth >= max_depth:
                        out[name[:-1]] = str(x)
                        return
                    
                    if type(x) is dict:
                        for a in x:
                            flatten(x[a], name + a + separator, depth + 1)
                    elif type(x) is list:
                        i = 0
                        for a in x[:10]:  # Limit to first 10 items in arrays
                            flatten(a, name + str(i) + separator, depth + 1)
                            i += 1
                    else:
                        out[name[:-1]] = x
                
                flatten(nested_json)
                return out
            
            if isinstance(data, dict):
                flattened = flatten_json(data, max_depth=max_depth)
                return pd.DataFrame([flattened])
            elif isinstance(data, list):
                if not data:
                    return pd.DataFrame([{"message": "Empty array"}])
                
                flattened_list = []
                for item in data[:1000]:  # Limit to first 1000 items
                    if isinstance(item, dict):
                        flattened_list.append(flatten_json(item, max_depth=max_depth))
                    else:
                        flattened_list.append({"value": item})
                
                return pd.DataFrame(flattened_list)
            else:
                return pd.DataFrame([{"content": str(data)}])
                
        except Exception as e:
            logger.error(f"Complex JSON parsing failed: {str(e)}")
            return pd.DataFrame([{"parsing_error": str(e), "raw_content": str(data)[:500]}])
    
    def _load_excel_file(self, file_path: str) -> pd.DataFrame:
        """Load Excel file (XLS/XLSX)"""
        try:
            # Try different engines
            if file_path.endswith('.xlsx'):
                return pd.read_excel(file_path, engine='openpyxl')
            else:
                return pd.read_excel(file_path, engine='xlrd')
        except Exception as e:
            logger.warning(f"⚠️ Excel loading failed: {str(e)}")
            # Create a DataFrame with error info
            return pd.DataFrame([{"error": f"Failed to load Excel file: {str(e)}"}])
    
    def _load_csv_file(self, file_path: str) -> pd.DataFrame:
        """Load CSV file with multiple encoding attempts"""
        encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
        
        for encoding in encodings:
            try:
                return pd.read_csv(file_path, encoding=encoding, low_memory=False)
            except UnicodeDecodeError:
                continue
            except Exception as e:
                logger.warning(f"⚠️ CSV loading failed with {encoding}: {str(e)}")
                continue
        
        # If all encodings fail
        return pd.DataFrame([{"error": "Failed to load CSV file with supported encodings"}])
    
    def _convert_document_to_dataframe(self, file_path: str) -> pd.DataFrame:
        """Convert any document type to a structured DataFrame with enhanced error handling"""
        file_ext = os.path.splitext(file_path)[1].lower()
        
        try:
            logger.info(f"🔄 Processing {file_ext} file: {os.path.basename(file_path)}")
            
            if file_ext == '.csv':
                df = self._load_csv_file(file_path)
            
            elif file_ext in ['.xlsx', '.xls']:
                df = self._load_excel_file(file_path)
            
            elif file_ext == '.json':
                result = self._load_json_file(file_path)
                if isinstance(result, str):
                    # JSON loading failed, create DataFrame with error info
                    logger.error(f"JSON parsing failed: {result}")
                    return pd.DataFrame([{
                        "content": result, 
                        "file_type": "json", 
                        "filename": os.path.basename(file_path),
                        "status": "error",
                        "error_details": result
                    }])
                df = result
            
            elif file_ext == '.pdf':
                text = self._extract_text_from_pdf(file_path)
                df = self._create_text_dataframe(text, file_path, "pdf")
            
            elif file_ext == '.docx':
                text = self._extract_text_from_docx(file_path)
                df = self._create_text_dataframe(text, file_path, "docx")
            
            elif file_ext == '.doc':
                text = self._extract_text_from_doc(file_path)
                df = self._create_text_dataframe(text, file_path, "doc")
            
            elif file_ext == '.txt':
                text = self._extract_text_from_txt(file_path)
                df = self._create_text_dataframe(text, file_path, "txt")
            
            else:
                logger.warning(f"⚠️ Unsupported file type: {file_ext}")
                return pd.DataFrame([{
                    "error": f"Unsupported file type: {file_ext}",
                    "filename": os.path.basename(file_path),
                    "supported_types": ", ".join(self.supported_formats)
                }])
            
            # Validate and clean the resulting DataFrame
            df = self._validate_and_clean_dataframe(df, file_path, file_ext)
            
            logger.info(f"✅ Successfully converted {file_ext} to DataFrame: {df.shape}")
            return df
                
        except Exception as e:
            logger.error(f"❌ Error converting document: {str(e)}")
            return pd.DataFrame([{
                "error": f"Error processing file: {str(e)}",
                "filename": os.path.basename(file_path),
                "file_type": file_ext,
                "processing_timestamp": datetime.now().isoformat()
            }])
    
    def _validate_and_clean_dataframe(self, df: pd.DataFrame, file_path: str, file_ext: str) -> pd.DataFrame:
        """Validate and clean the DataFrame after processing"""
        try:
            if df.empty:
                logger.warning("DataFrame is empty, creating placeholder")
                return pd.DataFrame([{
                    "message": "No data found in file",
                    "filename": os.path.basename(file_path),
                    "file_type": file_ext
                }])
            
            # Check for error columns and handle them
            if 'error' in df.columns and len(df) == 1:
                logger.warning(f"Error found in processed data: {df.iloc[0]['error']}")
                return df  # Return as-is to show the error
            
            # Clean column names
            original_columns = df.columns.tolist()
            df.columns = df.columns.astype(str)  # Ensure all column names are strings
            df.columns = df.columns.str.strip()  # Remove whitespace
            
            # Replace problematic characters in column names
            df.columns = df.columns.str.replace(r'[^\w\s]', '_', regex=True)
            df.columns = df.columns.str.replace(r'\s+', '_', regex=True)
            
            # Handle duplicate column names
            if df.columns.duplicated().any():
                df.columns = [f"{col}_{i}" if df.columns.tolist()[:i].count(col) > 0 
                             else col for i, col in enumerate(df.columns)]
            
            # Log column name changes if any
            if original_columns != df.columns.tolist():
                logger.info(f"Column names cleaned: {len(original_columns)} columns")
            
            # Handle data types and clean content
            for col in df.columns:
                if df[col].dtype == 'object':
                    # Convert to string and handle NaN values
                    df[col] = df[col].astype(str)
                    df[col] = df[col].replace(['nan', 'None', 'null'], '')
                    df[col] = df[col].str.strip()
            
            # Remove completely empty rows
            df = df.dropna(how='all')
            
            # If still empty after cleaning
            if df.empty:
                logger.warning("DataFrame became empty after cleaning")
                return pd.DataFrame([{
                    "message": "No valid data remained after cleaning",
                    "filename": os.path.basename(file_path),
                    "file_type": file_ext
                }])
            
            # Add metadata columns
            df['source_filename'] = os.path.basename(file_path)
            df['file_extension'] = file_ext
            df['processing_timestamp'] = datetime.now().isoformat()
            
            logger.info(f"✅ DataFrame validated and cleaned: {df.shape}")
            return df
            
        except Exception as e:
            logger.error(f"❌ Error validating DataFrame: {str(e)}")
            # Return original DataFrame with error info
            error_df = pd.DataFrame([{
                "validation_error": str(e),
                "filename": os.path.basename(file_path),
                "file_type": file_ext,
                "original_shape": str(df.shape) if hasattr(df, 'shape') else 'unknown'
            }])
            return error_df
    
    def _create_text_dataframe(self, text: str, file_path: str, file_type: str) -> pd.DataFrame:
        """Create a structured DataFrame from extracted text"""
        try:
            # Split text into paragraphs/sections
            paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
            
            if not paragraphs:
                paragraphs = [p.strip() for p in text.split('\n') if p.strip()]
            
            # Create DataFrame with text analysis
            data = []
            for i, paragraph in enumerate(paragraphs[:1000]):  # Limit to first 1000 paragraphs
                data.append({
                    "document_id": i + 1,
                    "content": paragraph,
                    "content_length": len(paragraph),
                    "word_count": len(paragraph.split()),
                    "filename": os.path.basename(file_path),
                    "file_type": file_type,
                    "user_id": self.user_id,  # Add user context
                    "extraction_timestamp": pd.Timestamp.now()
                })
            
            if not data:
                # If no paragraphs found, create single row with full text
                data = [{
                    "document_id": 1,
                    "content": text[:10000],  # Limit content length
                    "content_length": len(text),
                    "word_count": len(text.split()),
                    "filename": os.path.basename(file_path),
                    "file_type": file_type,
                    "user_id": self.user_id,
                    "extraction_timestamp": pd.Timestamp.now()
                }]
            
            return pd.DataFrame(data)
            
        except Exception as e:
            logger.error(f"❌ Error creating text DataFrame: {str(e)}")
            return pd.DataFrame([{
                "content": text[:1000] if len(text) > 1000 else text,
                "filename": os.path.basename(file_path),
                "file_type": file_type,
                "user_id": self.user_id,
                "error": f"Error structuring text: {str(e)}"
            }])
    
    def _create_visualization_tools(self):
        """Create comprehensive visualization tools"""
        
        @tool
        def create_automatic_visualization(query_result: str, query_type: str = "auto") -> str:
            """Automatically create the most appropriate visualization based on data and query"""
            try:
                # Parse query result to get data
                data = self._parse_query_result_for_viz(query_result)
                if not data:
                    return "No data available for visualization"
                
                df = pd.DataFrame(data)
                if df.empty:
                    return "Empty dataset for visualization"
                
                # Determine best visualization type based on data structure and query
                viz_type = self._determine_best_visualization(df, query_result, query_type)
                chart_title = self._generate_chart_title(query_result, query_type)
                
                # Create the appropriate visualization
                if viz_type == "bar":
                    return self._create_bar_chart_plotly(df, chart_title)
                elif viz_type == "line":
                    return self._create_line_chart_plotly(df, chart_title)
                elif viz_type == "pie":
                    return self._create_pie_chart_plotly(df, chart_title)
                elif viz_type == "histogram":
                    return self._create_histogram_plotly(df, chart_title)
                elif viz_type == "scatter":
                    return self._create_scatter_plot_plotly(df, chart_title)
                elif viz_type == "heatmap":
                    return self._create_heatmap_plotly(df, chart_title)
                else:
                    return self._create_bar_chart_plotly(df, chart_title)  # Default fallback
                    
            except Exception as e:
                return f"Error creating visualization: {str(e)}"
        
        @tool
        def create_dashboard(query_result: str, dashboard_type: str = "summary") -> str:
            """Create a comprehensive dashboard with multiple visualizations"""
            try:
                data = self._parse_query_result_for_viz(query_result)
                if not data:
                    return "No data available for dashboard"
                
                df = pd.DataFrame(data)
                if df.empty:
                    return "Empty dataset for dashboard"
                
                # Create multi-panel dashboard
                dashboard_path = self._create_multi_panel_dashboard(df, dashboard_type)
                return f"Dashboard created and saved as {dashboard_path}"
                
            except Exception as e:
                return f"Error creating dashboard: {str(e)}"
        
        @tool
        def export_visualization_data(query_result: str, format_type: str = "csv") -> str:
            """Export visualization data in various formats"""
            try:
                data = self._parse_query_result_for_viz(query_result)
                if not data:
                    return "No data available for export"
                
                df = pd.DataFrame(data)
                if df.empty:
                    return "Empty dataset for export"
                
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                
                if format_type.lower() == "csv":
                    export_path = f"export_data_{timestamp}.csv"
                    df.to_csv(export_path, index=False)
                elif format_type.lower() == "json":
                    export_path = f"export_data_{timestamp}.json"
                    df.to_json(export_path, orient='records', indent=2)
                elif format_type.lower() == "excel":
                    export_path = f"export_data_{timestamp}.xlsx"
                    df.to_excel(export_path, index=False)
                else:
                    export_path = f"export_data_{timestamp}.csv"
                    df.to_csv(export_path, index=False)
                
                return f"Data exported to {export_path} with {len(df)} records"
                
            except Exception as e:
                return f"Error exporting data: {str(e)}"
        
        return [create_automatic_visualization, create_dashboard, export_visualization_data]
    
    def _determine_best_visualization(self, df: pd.DataFrame, query_result: str, query_type: str) -> str:
        """Determine the best visualization type based on data characteristics"""
        try:
            # Analyze data structure
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            categorical_cols = df.select_dtypes(include=['object']).columns
            date_cols = [col for col in df.columns if 'date' in col.lower() or 'time' in col.lower()]
            
            num_rows = len(df)
            num_numeric = len(numeric_cols)
            num_categorical = len(categorical_cols)
            
            # Analyze query context
            query_lower = query_result.lower()
            
            # Decision logic for visualization type
            if "over time" in query_lower or "trend" in query_lower or date_cols:
                return "line"
            elif "distribution" in query_lower and num_numeric > 0:
                return "histogram"
            elif "correlation" in query_lower and num_numeric >= 2:
                return "scatter"
            elif "percentage" in query_lower or "share" in query_lower or "proportion" in query_lower:
                return "pie"
            elif num_categorical > 0 and num_numeric > 0:
                if num_rows > 50:
                    return "bar"
                else:
                    return "bar"
            elif num_numeric >= 2 and num_categorical >= 1:
                return "heatmap"
            else:
                return "bar"  # Default
                
        except Exception as e:
            logger.warning(f"Could not determine visualization type: {e}")
            return "bar"
    
    def _generate_chart_title(self, query_result: str, query_type: str) -> str:
        """Generate an appropriate chart title based on query context"""
        try:
            # Extract key information from query result
            if self.current_table_name:
                base_title = f"Analysis of {self.current_table_name.replace('_', ' ').title()}"
            else:
                base_title = "Data Analysis"
            
            # Add query context
            if "count" in query_result.lower():
                base_title += " - Count Analysis"
            elif "average" in query_result.lower() or "avg" in query_result.lower():
                base_title += " - Average Values"
            elif "sum" in query_result.lower() or "total" in query_result.lower():
                base_title += " - Total Analysis"
            elif "distribution" in query_result.lower():
                base_title += " - Distribution"
            elif "comparison" in query_result.lower():
                base_title += " - Comparison"
            
            return base_title
            
        except Exception as e:
            return "Data Visualization"
    
    def _create_bar_chart_plotly(self, df: pd.DataFrame, title: str) -> str:
        """Create a bar chart using Plotly"""
        try:
            # Auto-detect columns
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            categorical_cols = df.select_dtypes(include=['object']).columns
            
            if len(categorical_cols) > 0 and len(numeric_cols) > 0:
                x_col = categorical_cols[0]
                y_col = numeric_cols[0]
            else:
                x_col = df.columns[0]
                y_col = df.columns[1] if len(df.columns) > 1 else df.columns[0]
            
            fig = px.bar(df, x=x_col, y=y_col, title=title,
                        color=y_col, color_continuous_scale='viridis',
                        labels={x_col: x_col.replace('_', ' ').title(),
                               y_col: y_col.replace('_', ' ').title()})
            
            fig.update_layout(showlegend=False, height=500, 
                            xaxis_title=x_col.replace('_', ' ').title(),
                            yaxis_title=y_col.replace('_', ' ').title())
            
            chart_path = f"chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(chart_path)
            
            return f"Bar chart created and saved as {chart_path}"
            
        except Exception as e:
            return f"Error creating bar chart: {str(e)}"
    
    def _create_line_chart_plotly(self, df: pd.DataFrame, title: str) -> str:
        """Create a line chart using Plotly"""
        try:
            x_col = df.columns[0]
            y_col = df.columns[1] if len(df.columns) > 1 else df.columns[0]
            
            fig = px.line(df, x=x_col, y=y_col, title=title, markers=True,
                         labels={x_col: x_col.replace('_', ' ').title(),
                                y_col: y_col.replace('_', ' ').title()})
            fig.update_layout(height=500)
            
            chart_path = f"line_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(chart_path)
            
            return f"Line chart created and saved as {chart_path}"
            
        except Exception as e:
            return f"Error creating line chart: {str(e)}"
    
    def _create_pie_chart_plotly(self, df: pd.DataFrame, title: str) -> str:
        """Create a pie chart using Plotly"""
        try:
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            categorical_cols = df.select_dtypes(include=['object']).columns
            
            if len(categorical_cols) > 0 and len(numeric_cols) > 0:
                names_col = categorical_cols[0]
                values_col = numeric_cols[0]
            else:
                names_col = df.columns[0]
                values_col = df.columns[1] if len(df.columns) > 1 else df.columns[0]
            
            fig = px.pie(df, values=values_col, names=names_col, title=title)
            fig.update_layout(height=500)
            
            chart_path = f"pie_chart_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(chart_path)
            
            return f"Pie chart created and saved as {chart_path}"
            
        except Exception as e:
            return f"Error creating pie chart: {str(e)}"
    
    def _create_histogram_plotly(self, df: pd.DataFrame, title: str) -> str:
        """Create a histogram using Plotly"""
        try:
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            column = numeric_cols[0] if len(numeric_cols) > 0 else df.columns[0]
            
            fig = px.histogram(df, x=column, title=title, nbins=20,
                              labels={column: column.replace('_', ' ').title()})
            fig.update_layout(height=500)
            
            chart_path = f"histogram_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(chart_path)
            
            return f"Histogram created and saved as {chart_path}"
            
        except Exception as e:
            return f"Error creating histogram: {str(e)}"
    
    def _create_scatter_plot_plotly(self, df: pd.DataFrame, title: str) -> str:
        """Create a scatter plot using Plotly"""
        try:
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            if len(numeric_cols) >= 2:
                x_col = numeric_cols[0]
                y_col = numeric_cols[1]
            else:
                x_col = df.columns[0]
                y_col = df.columns[1] if len(df.columns) > 1 else df.columns[0]
            
            fig = px.scatter(df, x=x_col, y=y_col, title=title, trendline="ols",
                            labels={x_col: x_col.replace('_', ' ').title(),
                                   y_col: y_col.replace('_', ' ').title()})
            fig.update_layout(height=500)
            
            chart_path = f"scatter_plot_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(chart_path)
            
            return f"Scatter plot created and saved as {chart_path}"
            
        except Exception as e:
            return f"Error creating scatter plot: {str(e)}"
    
    def _create_heatmap_plotly(self, df: pd.DataFrame, title: str) -> str:
        """Create a heatmap using Plotly"""
        try:
            # Select only numeric columns for correlation heatmap
            numeric_df = df.select_dtypes(include=[np.number])
            
            if len(numeric_df.columns) < 2:
                return "Not enough numeric columns for heatmap"
            
            corr_matrix = numeric_df.corr()
            
            fig = px.imshow(corr_matrix, title=f"{title} - Correlation Heatmap",
                            color_continuous_scale='RdBu_r', aspect="auto")
            fig.update_layout(height=500)
            
            chart_path = f"heatmap_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(chart_path)
            
            return f"Heatmap created and saved as {chart_path}"
            
        except Exception as e:
            return f"Error creating heatmap: {str(e)}"
    
    def _create_multi_panel_dashboard(self, df: pd.DataFrame, dashboard_type: str) -> str:
        """Create a comprehensive multi-panel dashboard"""
        try:
            from plotly.subplots import make_subplots
            
            # Determine the number of subplots based on data
            numeric_cols = df.select_dtypes(include=[np.number]).columns
            categorical_cols = df.select_dtypes(include=['object']).columns
            
            # Create subplots
            if len(numeric_cols) >= 2 and len(categorical_cols) >= 1:
                fig = make_subplots(
                    rows=2, cols=2,
                    subplot_titles=('Count Distribution', 'Trend Analysis', 'Category Breakdown', 'Summary Statistics'),
                    specs=[[{"secondary_y": False}, {"secondary_y": False}],
                           [{"type": "domain"}, {"secondary_y": False}]]
                )
            else:
                fig = make_subplots(
                    rows=1, cols=2,
                    subplot_titles=('Data Overview', 'Statistical Summary')
                )
            
            # Add different chart types to subplots
            if len(categorical_cols) > 0 and len(numeric_cols) > 0:
                cat_col = categorical_cols[0]
                num_col = numeric_cols[0]
                
                # Count by category
                count_data = df[cat_col].value_counts()
                fig.add_trace(
                    go.Bar(x=count_data.index, y=count_data.values, name="Count"),
                    row=1, col=1
                )
                
                # Trend or line chart
                if len(df) > 1:
                    fig.add_trace(
                        go.Scatter(x=df.index, y=df[num_col], mode='lines+markers', name="Trend"),
                        row=1, col=2
                    )
                
                # Pie chart for categories
                if len(numeric_cols) >= 2:
                    fig.add_trace(
                        go.Pie(labels=count_data.index, values=count_data.values, name="Distribution"),
                        row=2, col=1
                    )
                
                # Summary statistics
                stats_df = df[numeric_cols].describe()
                if len(stats_df) > 0:
                    fig.add_trace(
                        go.Bar(x=stats_df.columns, y=stats_df.loc['mean'], name="Mean Values"),
                        row=2, col=2
                    )
            
            # Update layout
            fig.update_layout(
                height=800,
                title_text=f"Dashboard: {self.current_table_name or 'Data Analysis'}",
                showlegend=True
            )
            
            dashboard_path = f"dashboard_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
            fig.write_html(dashboard_path)
            
            return dashboard_path
            
        except Exception as e:
            logger.error(f"Error creating dashboard: {str(e)}")
            return f"dashboard_error_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
    
    def _parse_query_result_for_viz(self, query_result: str) -> List[Dict]:
        """Parse SQL query result text to extract structured data for visualization"""
        try:
            # Try to extract data from the query result text
            lines = query_result.strip().split('\n')
            data = []
            
            # Look for patterns like: Category, Value or (Category, Value)
            for line in lines:
                line = line.strip()
                if not line or line.startswith('SELECT') or line.startswith('FROM'):
                    continue
                    
                # Handle different formats
                if line.startswith('(') and line.endswith(')'):
                    # Format: ('Category', 123)
                    line = line.strip('()')
                
                if ',' in line:
                    parts = [part.strip().strip("'\"()") for part in line.split(',')]
                    if len(parts) >= 2:
                        try:
                            # Try to convert second part to number
                            value = float(parts[1])
                            data.append({'category': parts[0], 'value': value})
                        except ValueError:
                            # If not numeric, treat as string
                            data.append({'category': parts[0], 'value': parts[1]})
            
            # If no data found from text parsing, try to get from current table directly
            if not data and self.current_table_name:
                try:
                    with self.engine.connect() as connection:
                        # Use quoted table name for safety
                        safe_table_name = self.current_table_name.replace('"', '""')
                        query = text(f'SELECT * FROM "{safe_table_name}" LIMIT 100')
                        result = connection.execute(query).fetchall()
                        if result:
                            columns = result[0]._mapping.keys()
                            data = [dict(row._mapping) for row in result]
                except Exception as e:
                    logger.warning(f"Could not fetch data from table: {str(e)}")
            
            return data
            
        except Exception as e:
            logger.warning(f"⚠️ Error parsing query result: {str(e)}")
            return []
    
    def _create_tools(self):
        """Create all tools including document processing and visualization"""
        doc_tools = self._create_document_tools()
        viz_tools = self._create_visualization_tools()
        return doc_tools + viz_tools
    
    def _create_document_tools(self):
        """Create document processing tools"""
        
        @tool
        def analyze_document(file_path: str) -> str:
            """Analyze any document and return comprehensive information"""
            try:
                file_ext = os.path.splitext(file_path)[1].lower()
                file_size = os.path.getsize(file_path)
                
                # Convert document to DataFrame
                df = self._convert_document_to_dataframe(file_path)
                
                # Classify document if services available
                classification_info = {}
                if SERVICES_AVAILABLE and self.document_classifier:
                    try:
                        # Prepare content for classification
                        content = {
                            'structured_data': df.to_dict('records') if not df.empty else [],
                            'metadata': {
                                'filename': os.path.basename(file_path),
                                'size_bytes': file_size,
                                'record_count': len(df)
                            }
                        }
                        
                        # Get AI classification
                        import asyncio
                        classification = asyncio.run(
                            self.document_classifier.classify_by_content(content, file_path)
                        )
                        
                        classification_info = {
                            'business_type': classification.business_type.value,
                            'confidence': classification.confidence,
                            'storage_container': classification.storage_container,
                            'requires_database': classification.requires_database
                        }
                        
                    except Exception as e:
                        logger.warning(f"Classification failed: {str(e)}")
                        classification_info = {'error': f"Classification failed: {str(e)}"}
                
                info = {
                    "filename": os.path.basename(file_path),
                    "file_type": file_ext,
                    "file_size_mb": round(file_size / (1024*1024), 2),
                    "dataframe_shape": df.shape,
                    "columns": list(df.columns),
                    "sample_data": df.head(3).to_dict(),
                    "classification": classification_info,
                    "user_id": self.user_id
                }
                return f"Document Analysis: {json.dumps(info, indent=2, default=str)}"
            except Exception as e:
                return f"Error analyzing document: {str(e)}"
        
        @tool
        def process_and_save_document(file_path: str, output_path: str = "processed_document.csv") -> str:
            """Process any document type and save as structured CSV with classification"""
            try:
                df = self._convert_document_to_dataframe(file_path)
                
                # Basic cleaning for all document types
                if 'content' in df.columns:
                    # Clean text content
                    df['content'] = df['content'].astype(str).str.strip()
                    df = df[df['content'] != '']  # Remove empty content
                
                # Add user context to all records
                df['user_id'] = self.user_id
                df['processing_timestamp'] = datetime.now().isoformat()
                df['source_file'] = os.path.basename(file_path)
                
                # Clean column names
                df.columns = df.columns.str.strip()
                df.columns = df.columns.str.replace(' ', '_', regex=False)
                df.columns = df.columns.str.replace('[^a-zA-Z0-9_]', '', regex=True)
                
                # Save processed data
                df.to_csv(output_path, index=False)
                
                # Store in schema manager if available
                schema_result = {}
                if SERVICES_AVAILABLE and self.schema_manager:
                    try:
                        # Get business type from classification
                        business_type = "general"
                        if hasattr(self, 'current_document_info') and 'business_type' in self.current_document_info:
                            business_type = self.current_document_info['business_type']
                        
                        # Create user table
                        import asyncio
                        schema_result = asyncio.run(
                            self.schema_manager.create_user_business_table(
                                user_id=self.user_id,
                                business_type=business_type,
                                data=df.to_dict('records'),
                                filename=os.path.basename(file_path)
                            )
                        )
                        
                    except Exception as e:
                        logger.warning(f"Schema management failed: {str(e)}")
                        schema_result = {'error': f"Schema management failed: {str(e)}"}
                
                return f"Processed document saved to {output_path}. Shape: {df.shape}, Columns: {list(df.columns)}. Schema result: {schema_result}"
                
            except Exception as e:
                return f"Error processing document: {str(e)}"
        
        return [analyze_document, process_and_save_document]
    
    def _create_enhanced_sql_agent(self):
        """Create enhanced SQL agent with visualization capabilities"""
        try:
            # Create custom prompt that includes visualization instructions
            enhanced_prompt = ChatPromptTemplate.from_messages([
                ("system", """You are an expert SQL analyst with advanced visualization capabilities. 
                
                When answering queries:
                1. First execute the SQL query to get the data
                2. Always provide the results in a clear, structured format
                3. AUTOMATICALLY create appropriate visualizations for the data
                4. Choose the best visualization type based on the data and query context
                5. Provide insights and analysis of the results
                
                Available visualization tools:
                - create_automatic_visualization: Automatically chooses the best chart type
                - create_dashboard: Creates comprehensive multi-panel dashboards
                - export_visualization_data: Exports data in various formats
                
                For ANY query that returns data, you MUST:
                1. Show the data results
                2. Create a visualization using create_automatic_visualization
                3. Provide analysis and insights
                
                Query types and their preferred visualizations:
                - Count/frequency queries → Bar charts
                - Trend/time series → Line charts  
                - Proportions/percentages → Pie charts
                - Distributions → Histograms
                - Correlations → Scatter plots
                - Multiple metrics → Dashboard
                
                Always end your response with a visualization."""),
                
                ("human", "{input}"),
                MessagesPlaceholder(variable_name="agent_scratchpad"),
            ])
            
            # Create SQL agent with enhanced capabilities
            agent = create_sql_agent(
                llm=self.llm,
                db=self.db,
                agent_type="tool-calling",
                verbose=True,
                extra_tools=self.tools,  # Include visualization tools
                prompt=enhanced_prompt
            )
            
            logger.info("✅ Enhanced SQL agent with visualization created successfully")
            return agent
            
        except Exception as e:
            logger.error(f"❌ Failed to create enhanced SQL agent: {str(e)}")
            # Fallback to basic SQL agent
            return create_sql_agent(
                llm=self.llm,
                db=self.db,
                agent_type="tool-calling",
                verbose=True
            )
    
    @retry(
        wait=wait_exponential(multiplier=1, min=4, max=60),
        stop=stop_after_attempt(3),
        retry=retry_if_exception_type((Exception,))
    )
    def _invoke_agent_with_retry(self, agent_executor, input_text):
        """Invoke an agent with retry mechanism"""
        try:
            return agent_executor.invoke({"input": input_text})
        except Exception as e:
            logger.warning(f"⚠️ Agent invocation failed: {str(e)}")
            raise
    
    async def process_document(self, file_path: str, table_name: Optional[str] = None) -> pd.DataFrame:
        """
        Process any document type with AI classification and schema management
        
        Args:
            file_path (str): Path to any supported document type
            table_name (str, optional): Name for the SQL table. If None, uses AI classification
            
        Returns:
            pd.DataFrame: The processed DataFrame
        """
        try:
            logger.info(f"📂 Starting integrated document processing: {file_path}")
            logger.info(f"👤 User ID: {self.user_id}")
            
            # Check file exists and size
            if not os.path.exists(file_path):
                raise FileNotFoundError(f"File not found: {file_path}")
            
            file_size = os.path.getsize(file_path)
            if file_size > self.settings.MAX_FILE_SIZE:
                raise ValueError(f"File too large: {file_size/1024/1024:.1f}MB > {self.settings.MAX_FILE_SIZE/1024/1024:.1f}MB")
            
            # Check file type
            file_ext = os.path.splitext(file_path)[1].lower().replace('.', '')
            if file_ext not in self.settings.ALLOWED_FILE_TYPES:
                raise ValueError(f"Unsupported file type: {file_ext}. Supported: {', '.join(self.settings.ALLOWED_FILE_TYPES)}")
            
            # Step 1: Convert document to DataFrame
            logger.info(f"📄 Converting {file_ext.upper()} document to structured data...")
            original_df = self._convert_document_to_dataframe(file_path)
            
            logger.info(f"📊 Original DataFrame shape: {original_df.shape}")
            logger.info(f"📋 Original DataFrame columns: {list(original_df.columns)}")
            
            # Step 2: AI Classification (if available)
            classification_result = None
            business_type = "general"
            
            if SERVICES_AVAILABLE and self.document_classifier:
                logger.info("🧠 Performing AI-powered document classification...")
                try:
                    # Prepare content for classification
                    content = {
                        'structured_data': original_df.to_dict('records') if not original_df.empty else [],
                        'raw_content': original_df.to_string() if len(original_df) < 100 else original_df.head(10).to_string(),
                        'metadata': {
                            'filename': os.path.basename(file_path),
                            'size_bytes': file_size,
                            'record_count': len(original_df),
                            'columns': list(original_df.columns)
                        }
                    }
                    
                    classification_result = await self.document_classifier.classify_by_content(content, file_path)
                    business_type = classification_result.business_type.value
                    
                    logger.info(f"🎯 Classification: {business_type} (confidence: {classification_result.confidence:.2f})")
                    
                    # Store classification info for later use
                    self.current_document_info = {
                        'business_type': business_type,
                        'confidence': classification_result.confidence,
                        'storage_container': classification_result.storage_container
                    }
                    
                except Exception as e:
                    logger.warning(f"⚠️ AI classification failed: {str(e)}")
                    classification_result = None
            
            # Step 3: Add user context and metadata
            processed_df = original_df.copy()
            processed_df['user_id'] = self.user_id
            processed_df['business_type'] = business_type
            processed_df['source_file'] = os.path.basename(file_path)
            processed_df['processing_timestamp'] = datetime.now().isoformat()
            processed_df['document_id'] = str(uuid.uuid4())
            
            # Basic data cleaning
            for col in processed_df.select_dtypes(include=['object']).columns:
                if processed_df[col].dtype == 'object':
                    processed_df[col] = processed_df[col].astype(str).replace('nan', '').str.strip()
            
            logger.info(f"✅ Enhanced DataFrame shape: {processed_df.shape}")
            logger.info(f"📋 Enhanced DataFrame columns: {list(processed_df.columns)}")
            
            # Step 4: Schema Management (if available)
            schema_result = None
            if SERVICES_AVAILABLE and self.schema_manager:
                logger.info("🗄️ Creating user-specific database schema...")
                try:
                    schema_result = await self.schema_manager.create_user_business_table(
                        user_id=self.user_id,
                        business_type=business_type,
                        data=processed_df.to_dict('records'),
                        filename=os.path.basename(file_path)
                    )
                    
                    if schema_result.get('success'):
                        logger.info(f"✅ Schema created: {schema_result.get('collection_name')}")
                    else:
                        logger.warning(f"⚠️ Schema creation issue: {schema_result.get('error')}")
                        
                except Exception as e:
                    logger.warning(f"⚠️ Schema management failed: {str(e)}")
            
            # Step 5: Load into local SQL database for querying
            if table_name is None:
                base_name = os.path.basename(file_path)
                # Create a cleaner table name
                clean_filename = os.path.splitext(base_name)[0]
                table_name_parts = [
                    self.user_id,
                    business_type,
                    clean_filename
                ]
                
                # Join and sanitize
                raw_table_name = "_".join(table_name_parts)
                table_name = self._sanitize_table_name(raw_table_name)
            
            # Store current table name
            self.current_table_name = table_name
            
            # Load into SQL database
            processed_df.to_sql(table_name, self.engine, if_exists='replace', index=False)
            logger.info(f"📊 Loaded processed data into SQL table: {table_name}")
            
            # Refresh the database connection
            self.db = SQLDatabase(engine=self.engine)
            self.sql_agent = self._create_enhanced_sql_agent()
            
            # Step 6: Generate processing summary
            summary = {
                "file_processed": os.path.basename(file_path),
                "user_id": self.user_id,
                "file_type": file_ext,
                "business_type": business_type,
                "records_processed": len(processed_df),
                "table_name": table_name,
                "classification": classification_result.__dict__ if classification_result else None,
                "schema_result": schema_result,
                "processing_timestamp": datetime.now().isoformat()
            }
            
            logger.info("=" * 70)
            logger.info("🎉 DOCUMENT PROCESSING COMPLETED SUCCESSFULLY")
            logger.info(f"📄 File: {os.path.basename(file_path)}")
            logger.info(f"👤 User: {self.user_id}")
            logger.info(f"🏷️ Type: {business_type}")
            logger.info(f"📊 Records: {len(processed_df)}")
            logger.info(f"🗄️ Table: {table_name}")
            logger.info("=" * 70)
            
            return processed_df
        
        except Exception as e:
            error_msg = f"❌ Error processing document: {str(e)}"
            logger.error(error_msg)
            raise
    
    def execute_query_with_visualization(self, query: str) -> Dict[str, Any]:
        """
        Execute a natural language query with automatic visualization
        
        Args:
            query (str): Natural language query
            
        Returns:
            Dict: Results including data, analysis, and visualization paths
        """
        try:
            if self.current_table_name is None:
                return {"error": "No document has been loaded. Please load a document first."}
            
            # Check if current table has data issues and try to fix them
            self._check_and_fix_table_data()
            
            # Enhance query with table context and visualization instructions
            enhanced_query = f"""
            Using the {self.current_table_name} table for user {self.user_id}, {query}
            
            After providing the data results, you MUST create a visualization using the create_automatic_visualization tool.
            Choose the most appropriate visualization type based on the data and query context.
            """
            
            logger.info(f"🔍 Executing enhanced query with auto-visualization...")
            logger.info(f"📊 Table: {self.current_table_name}")
            logger.info(f"👤 User: {self.user_id}")
            
            # Execute query through enhanced SQL agent
            result = self.sql_agent.invoke({"input": enhanced_query})
            
            logger.info("✅ Query with visualization executed successfully")
            
            # Extract visualization files created
            viz_files = []
            try:
                # Look for generated visualization files
                for file in os.listdir('.'):
                    if file.endswith('.html') and any(chart_type in file for chart_type in 
                                                     ['chart_', 'line_chart_', 'pie_chart_', 
                                                      'histogram_', 'scatter_plot_', 'heatmap_', 
                                                      'dashboard_']):
                        viz_files.append(file)
                        
                # Sort by creation time (most recent first)
                viz_files.sort(key=lambda x: os.path.getctime(x), reverse=True)
                
            except Exception as e:
                logger.warning(f"Could not detect visualization files: {str(e)}")
            
            # Return comprehensive result
            return {
                "query": query,
                "table_name": self.current_table_name,
                "user_id": self.user_id,
                "result": result,
                "visualizations_created": viz_files[:5],  # Last 5 visualizations
                "timestamp": datetime.now().isoformat(),
                "success": True
            }
        
        except Exception as e:
            error_msg = f"❌ Error executing query with visualization: {str(e)}"
            logger.error(error_msg)
            return {
                "error": error_msg,
                "query": query,
                "user_id": self.user_id,
                "timestamp": datetime.now().isoformat(),
                "success": False
            }
    
    def _check_and_fix_table_data(self):
        """Check if the current table has data issues and attempt to fix them"""
        try:
            if not self.current_table_name:
                return
            
            with self.engine.connect() as connection:
                # Use quoted table name to handle special characters
                safe_table_name = self.current_table_name.replace('"', '""')
                
                # Check if table exists first
                table_exists_query = text("""
                    SELECT name FROM sqlite_master 
                    WHERE type='table' AND name = :table_name
                """)
                
                table_exists = connection.execute(
                    table_exists_query, 
                    {"table_name": self.current_table_name}
                ).fetchone()
                
                if not table_exists:
                    logger.warning(f"Table {self.current_table_name} does not exist")
                    return
                
                # Check if table has content column before querying it
                schema_query = text(f'PRAGMA table_info("{safe_table_name}")')
                schema_result = connection.execute(schema_query).fetchall()
                
                columns = [row[1] for row in schema_result]  # Column names are in index 1
                
                if 'content' in columns:
                    # Check if table has error content
                    error_check_query = text(f"""
                        SELECT content, COUNT(*) as count 
                        FROM "{safe_table_name}"
                        WHERE content LIKE '%Error loading JSON%' OR content LIKE '%Extra data%'
                        GROUP BY content
                        LIMIT 5
                    """)
                    
                    try:
                        error_results = connection.execute(error_check_query).fetchall()
                        
                        if error_results:
                            logger.warning(f"🔧 Found data issues in table {self.current_table_name}")
                            self._attempt_data_recovery()
                            
                    except Exception as e:
                        logger.debug(f"Error check query failed: {str(e)}")
                else:
                    logger.debug(f"Table {self.current_table_name} does not have 'content' column")
                    
        except Exception as e:
            logger.warning(f"Table data check failed: {str(e)}")
    
    def _attempt_data_recovery(self):
        """Attempt to recover data from a corrupted table"""
        try:
            logger.info("🔄 Attempting data recovery...")
            
            with self.engine.connect() as connection:
                # Get a sample of the data to understand the structure
                safe_table_name = self.current_table_name.replace('"', '""')
                sample_query = text(f'SELECT * FROM "{safe_table_name}" LIMIT 10')
                sample_data = connection.execute(sample_query).fetchall()
                
                if sample_data:
                    columns = sample_data[0]._mapping.keys()
                    logger.info(f"📋 Table structure: {list(columns)}")
                    
                    # Check if we can clean the data
                    if 'content' in columns:
                        # Try to extract useful information from error messages
                        logger.info("🔧 Attempting to clean error content...")
                
        except Exception as e:
            logger.warning(f"Data recovery attempt failed: {str(e)}")
    
    def diagnose_json_file(self, file_path: str) -> Dict[str, Any]:
        """Diagnose JSON file issues and provide recommendations"""
        try:
            diagnosis = {
                "file_path": file_path,
                "file_exists": os.path.exists(file_path),
                "file_size": 0,
                "encoding_detected": None,
                "json_format": None,
                "parse_issues": [],
                "recommendations": [],
                "sample_content": "",
                "can_be_fixed": False
            }
            
            if not os.path.exists(file_path):
                diagnosis["recommendations"].append("File does not exist")
                return diagnosis
            
            # Get file size
            diagnosis["file_size"] = os.path.getsize(file_path)
            
            # Try to detect encoding
            try:
                import chardet
                with open(file_path, 'rb') as file:
                    raw_data = file.read(1024)  # Read first 1KB
                    detected = chardet.detect(raw_data)
                    diagnosis["encoding_detected"] = detected.get('encoding', 'utf-8')
            except ImportError:
                logger.warning("chardet not available, using utf-8 as default encoding")
                diagnosis["encoding_detected"] = "utf-8"
            except Exception:
                diagnosis["encoding_detected"] = "utf-8"
            
            # Read sample content
            try:
                with open(file_path, 'r', encoding=diagnosis["encoding_detected"], errors='ignore') as file:
                    sample = file.read(500)
                    diagnosis["sample_content"] = sample
            except:
                pass
            
            # Detect JSON format
            diagnosis["json_format"] = self._detect_json_format(file_path)
            
            # Try parsing and identify issues
            try:
                with open(file_path, 'r', encoding=diagnosis["encoding_detected"]) as file:
                    json.load(file)
                diagnosis["parse_issues"].append("No issues - JSON parses correctly")
                diagnosis["can_be_fixed"] = True
            except json.JSONDecodeError as e:
                diagnosis["parse_issues"].append(f"JSON Decode Error: {str(e)}")
                
                # Analyze the specific error
                if "Extra data" in str(e):
                    diagnosis["parse_issues"].append("Multiple JSON objects detected (not in array format)")
                    diagnosis["recommendations"].append("Convert to JSON array format or parse as JSON Lines")
                    diagnosis["can_be_fixed"] = True
                elif "Expecting" in str(e):
                    diagnosis["parse_issues"].append("Malformed JSON syntax")
                    diagnosis["recommendations"].append("Check for missing quotes, brackets, or commas")
                    diagnosis["can_be_fixed"] = True
                elif "Unterminated" in str(e):
                    diagnosis["parse_issues"].append("Unterminated JSON structure")
                    diagnosis["recommendations"].append("File may be truncated or corrupted")
                    diagnosis["can_be_fixed"] = False
            
            # Additional checks
            if diagnosis["file_size"] == 0:
                diagnosis["parse_issues"].append("File is empty")
                diagnosis["recommendations"].append("File contains no data")
                diagnosis["can_be_fixed"] = False
            elif diagnosis["file_size"] > 100 * 1024 * 1024:  # 100MB
                diagnosis["parse_issues"].append("Very large file")
                diagnosis["recommendations"].append("Consider splitting into smaller files")
            
            return diagnosis
            
        except Exception as e:
            return {
                "file_path": file_path,
                "diagnosis_error": str(e),
                "recommendations": ["Could not diagnose file - check if file is accessible"]
            }
    
    def get_processing_recommendations(self, file_path: str) -> List[str]:
        """Get recommendations for processing a specific file"""
        try:
            diagnosis = self.diagnose_json_file(file_path)
            recommendations = []
            
            if not diagnosis.get("can_be_fixed", False):
                recommendations.append("❌ File cannot be automatically fixed")
                recommendations.extend(diagnosis.get("recommendations", []))
                return recommendations
            
            # File-specific recommendations
            if "mtcars" in file_path.lower():
                recommendations.append("🎯 Detected mtcars dataset - using specialized parser")
            
            if "parquet" in file_path.lower():
                recommendations.append("📊 Detected Parquet JSON format - using Arrow parser")
            
            json_format = diagnosis.get("json_format", "unknown")
            if json_format == "jsonl":
                recommendations.append("📝 JSON Lines format detected - parsing line by line")
            elif json_format == "json_array":
                recommendations.append("📋 JSON array format detected - standard parsing")
            elif json_format == "json_object":
                recommendations.append("📄 Single JSON object detected")
            
            if "Extra data" in str(diagnosis.get("parse_issues", [])):
                recommendations.append("🔧 Multiple JSON objects detected - will attempt to parse as JSONL or fix structure")
            
            return recommendations
            
        except Exception as e:
            return [f"❌ Could not generate recommendations: {str(e)}"]
    
    def get_user_document_summary(self) -> Dict[str, Any]:
        """Get comprehensive summary of user's processed documents"""
        try:
            summary = {
                "user_id": self.user_id,
                "current_table": self.current_table_name,
                "current_document_info": self.current_document_info,
                "available_tables": self.list_tables(),
                "services_status": {
                    "classification_service": SERVICES_AVAILABLE and self.document_classifier is not None,
                    "schema_management": SERVICES_AVAILABLE and self.schema_manager is not None,
                    "pdf_support": PDF_SUPPORT
                },
                "supported_formats": self.supported_formats,
                "timestamp": datetime.now().isoformat()
            }
            
            # Add table info if current table exists
            if self.current_table_name:
                table_info = self.get_table_info()
                summary["current_table_info"] = table_info
            
            # Add user tables from schema manager if available
            if SERVICES_AVAILABLE and self.schema_manager:
                try:
                    import asyncio
                    user_tables = asyncio.run(self.schema_manager.get_user_tables(self.user_id))
                    summary["user_tables_in_cosmos"] = user_tables
                except Exception as e:
                    summary["user_tables_error"] = str(e)
            
            return summary
            
        except Exception as e:
            logger.error(f"❌ Error getting user summary: {str(e)}")
            return {"error": str(e), "user_id": self.user_id}
    
    def get_table_info(self) -> Dict[str, Any]:
        """Get information about the current table"""
        if self.current_table_name is None:
            return {"error": "No table loaded"}
        
        try:
            with self.engine.connect() as connection:
                # Use parameterized query or properly escape table name
                safe_table_name = self.current_table_name.replace('"', '""')
                schema_query = text(f'PRAGMA table_info("{safe_table_name}")')
                schema_result = connection.execute(schema_query).fetchall()
                
                count_query = text(f'SELECT COUNT(*) FROM "{safe_table_name}"')
                count_result = connection.execute(count_query).fetchone()
            
            return {
                "table_name": self.current_table_name,
                "user_id": self.user_id,
                "schema": [dict(row._mapping) for row in schema_result],
                "row_count": count_result[0] if count_result else 0,
                "document_info": self.current_document_info
            }
            
        except Exception as e:
            logger.error(f"❌ Error getting table info: {str(e)}")
            return {"error": f"Error getting table info: {str(e)}"}
    
    def list_tables(self) -> List[str]:
        """List all tables in the local database"""
        try:
            with self.engine.connect() as connection:
                tables_query = text("SELECT name FROM sqlite_master WHERE type='table'")
                tables_result = connection.execute(tables_query).fetchall()
            
            return [row[0] for row in tables_result]
            
        except Exception as e:
            logger.error(f"❌ Error listing tables: {str(e)}")
            return []

