# Suppress protobuf warnings at the very beginning
import warnings
import os
import sys

# Suppress all common warnings first
warnings.filterwarnings('ignore', category=UserWarning, module='google.protobuf.runtime_version')
warnings.filterwarnings('ignore', category=DeprecationWarning, module='websockets')
warnings.filterwarnings('ignore', category=DeprecationWarning, module='uvicorn')
warnings.filterwarnings('ignore', message='.*swigvarlink.*')
warnings.filterwarnings('ignore', message='.*destructor.*')
warnings.filterwarnings('ignore', category=ResourceWarning, module='neo4j')

# Disable HuggingFace Hub symlinks warning
os.environ['HF_HUB_DISABLE_SYMLINKS_WARNING'] = '1'
os.environ['PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION'] = 'python'

# Check for critical dependencies first
try:
    import fastapi
    import uvicorn
    import pydantic
except ImportError as e:
    print(f"❌ Critical dependency missing: {e}")
    print("Please install required packages:")
    print("pip install fastapi uvicorn pydantic")
    sys.exit(1)

import asyncio
import logging
from typing import List, Dict, Any, Optional, Tuple, Union, Set
from pathlib import Path
import json
import uuid
from datetime import datetime
import time
from contextlib import asynccontextmanager
import re
from collections import defaultdict, Counter
import numpy as np
import concurrent.futures
from functools import partial
import itertools

# Try to import optional dependencies with fallbacks
try:
    import networkx as nx
    from networkx.algorithms import community
    NETWORKX_AVAILABLE = True
except ImportError:
    print("⚠️ NetworkX not available - multi-hop reasoning will be limited")
    NETWORKX_AVAILABLE = False

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    print("⚠️ Spacy or en_core_web_sm model not available")
    print("Install with: pip install spacy && python -m spacy download en_core_web_sm")
    SPACY_AVAILABLE = False
    nlp = None

# Core imports with error handling
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    print("⚠️ PyMuPDF not available. Install with: pip install PyMuPDF")
    PYMUPDF_AVAILABLE = False

try:
    import pymupdf
except ImportError:
    print("⚠️ pymupdf not available")

try:
    from unstructured.partition.auto import partition
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    from unstructured.partition.text import partition_text
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    print("⚠️ Unstructured not available. Install with: pip install unstructured[pdf]")
    UNSTRUCTURED_AVAILABLE = False

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

# Docling imports for advanced document processing
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
    DOCLING_AVAILABLE = True
except ImportError:
    print("⚠️ Docling not available. Install with: pip install docling")
    DOCLING_AVAILABLE = False

# FastAPI imports
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Security
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

# CLIP and Vision imports
try:
    from transformers import CLIPProcessor, CLIPModel
    from PIL import Image
    import torch
    CLIP_AVAILABLE = True
except ImportError:
    print("⚠️ CLIP dependencies not available. Install with: pip install transformers torch pillow")
    CLIP_AVAILABLE = False

try:
    import base64
    import io
    BASIC_AVAILABLE = True
except ImportError:
    BASIC_AVAILABLE = False

# LlamaIndex imports
try:
    from llama_index.core import (
        VectorStoreIndex, 
        Document as LlamaDocument,
        Settings,
        StorageContext,
        SimpleDirectoryReader,
        load_index_from_storage
    )
    from llama_index.core.node_parser import SimpleNodeParser, SentenceSplitter
    from llama_index.core.extractors import TitleExtractor, QuestionsAnsweredExtractor
    from llama_index.core.ingestion import IngestionPipeline
    from llama_index.core.schema import BaseNode, TextNode
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.llms.azure_openai import AzureOpenAI
    from llama_index.vector_stores.weaviate import WeaviateVectorStore as LlamaWeaviateVectorStore
    LLAMAINDEX_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ LlamaIndex not available: {e}")
    LLAMAINDEX_AVAILABLE = False
    # Define dummy classes to prevent NameError
    class BaseNode:
        pass
    class TextNode:
        pass
    class LlamaDocument:
        pass
    class HuggingFaceEmbedding:
        pass
    class AzureOpenAI:
        pass

# HuggingFace imports
try:
    from sentence_transformers import SentenceTransformer
    from transformers import AutoTokenizer, AutoModel
    HUGGINGFACE_AVAILABLE = True
except ImportError:
    print("⚠️ HuggingFace transformers not available. Please install: pip install transformers sentence-transformers")
    HUGGINGFACE_AVAILABLE = False

# Weaviate imports
try:
    import weaviate
    from langchain_weaviate.vectorstores import WeaviateVectorStore
    WEAVIATE_AVAILABLE = True
except ImportError:
    print("⚠️ Weaviate not available. Using file-based storage.")
    WEAVIATE_AVAILABLE = False

# Vector database and ML imports
import numpy as np
try:
    from sentence_transformers import SentenceTransformer, CrossEncoder
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    print("⚠️ SentenceTransformers not available. Reranking disabled.")
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except ImportError:
    print("⚠️ BM25 not available. Using simple retrieval.")
    BM25_AVAILABLE = False

# Neo4j imports
try:
    from neo4j import GraphDatabase
    NEO4J_AVAILABLE = True
except ImportError:
    print("⚠️ Neo4j not available. Using fallback knowledge tracking.")
    NEO4J_AVAILABLE = False

# Environment and utilities
try:
    from dotenv import load_dotenv
    DOTENV_AVAILABLE = True
except ImportError:
    print("⚠️ python-dotenv not available. Install with: pip install python-dotenv")
    DOTENV_AVAILABLE = False

try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    print("⚠️ aiofiles not available. Install with: pip install aiofiles")
    AIOFILES_AVAILABLE = False

try:
    from sklearn.metrics.pairwise import cosine_similarity
    from sklearn.feature_extraction.text import TfidfVectorizer
    from sklearn.cluster import KMeans
    SKLEARN_AVAILABLE = True
except ImportError:
    print("⚠️ scikit-learn not available. Install with: pip install scikit-learn")
    SKLEARN_AVAILABLE = False

# Load environment variables if available
if DOTENV_AVAILABLE:
    load_dotenv()
else:
    print("⚠️ Environment variables will be loaded from system only")

# Configure logging
logging.basicConfig(
    level=getattr(logging, os.getenv('LOG_LEVEL', 'INFO')),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Pydantic models
class QueryRequest(BaseModel):
    query: str
    top_k: int = 10
    use_graph: bool = True
    use_reranking: bool = True
    use_multimodal: bool = True
    use_rag_fusion: bool = True
    use_hybrid_search: bool = True
    summary_length: str = "medium"
    enable_multi_hop: bool = True
    include_recommendations: bool = True
    fusion_queries: int = 3
    dense_weight: float = 0.7
    sparse_weight: float = 0.3

class HybridScores(BaseModel):
    query_type: Optional[str] = None
    dense_weight: Optional[float] = None
    sparse_weight: Optional[float] = None
    dense_count: Optional[int] = None
    sparse_count: Optional[int] = None
    fused_count: Optional[int] = None
    final_count: Optional[int] = None
    fusion_used: Optional[bool] = None
    fallback_dense: Optional[bool] = None

class GraphCommunity(BaseModel):
    id: str
    algorithm: str
    nodes: List[str]
    size: int
    entities: List[Dict[str, Any]]

class QueryResponse(BaseModel):
    answer: str
    summary: str
    conclusion: str
    sources: List[Dict[str, Any]]
    entities: List[Dict[str, Any]]
    recommendations: List[str]
    confidence_score: float
    processing_time: float
    multi_hop_reasoning: List[str]
    image_sources: List[Dict[str, Any]]
    fusion_queries: List[str]
    hybrid_scores: HybridScores
    graph_communities: List[GraphCommunity]

class DocumentUploadResponse(BaseModel):
    message: str
    document_id: str
    chunks_created: int
    images_processed: int
    entities_extracted: int
    processing_time: float
    graph_nodes_created: int
    graph_relationships_created: int

# Configuration class
class Config:
    def __init__(self):
        # Azure OpenAI Configuration
        self.azure_openai_endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')
        self.azure_openai_api_key = os.getenv('AZURE_OPENAI_API_KEY')
        self.azure_openai_api_version = os.getenv('AZURE_OPENAI_API_VERSION', '2025-01-01-preview')
        self.azure_openai_deployment_name = os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME')
        self.azure_openai_model_name = os.getenv('AZURE_OPENAI_MODEL_NAME')
        self.azure_openai_temperature = float(os.getenv('AZURE_OPENAI_TEMPERATURE', '0.1'))
        self.azure_openai_max_tokens = int(os.getenv('AZURE_OPENAI_MAX_TOKENS', '4000'))
        
        # HuggingFace Configuration
        self.hf_embedding_model = os.getenv('HF_EMBEDDING_MODEL', 'BAAI/bge-large-en-v1.5')
        self.hf_reranker_model = os.getenv('HF_RERANKER_MODEL', 'BAAI/bge-reranker-large')
        # Use a simple local cache directory that's easier to create
        self.hf_cache_dir = os.getenv('HF_CACHE_DIR', './hf_cache')
        self.hf_device = os.getenv('HF_DEVICE', 'cpu')
        
        # CLIP Configuration
        self.clip_model_name = os.getenv('CLIP_MODEL_NAME', 'openai/clip-vit-base-patch32')
        
        # Neo4j Configuration
        self.neo4j_uri = os.getenv('NEO4J_URI')
        self.neo4j_username = os.getenv('NEO4J_USERNAME')
        self.neo4j_password = os.getenv('NEO4J_PASSWORD')
        self.neo4j_database = os.getenv('NEO4J_DATABASE')
        
        # Weaviate Configuration
        self.weaviate_url = os.getenv('WEAVIATE_URL', 'http://localhost:8080')
        self.weaviate_api_key = os.getenv('WEAVIATE_API_KEY')
        
        # Storage Configuration
        self.persist_dir = os.getenv('PERSIST_DIR', './storage')
        self.index_name = os.getenv('INDEX_NAME', 'rag_documents')
        
        # Security Configuration
        self.api_key_header = os.getenv('API_KEY_HEADER', 'X-API-Key')
        self.cors_origins = json.loads(os.getenv('CORS_ORIGINS', '["*"]'))
        self.debug = os.getenv('DEBUG', 'true').lower() == 'true'
        self.max_file_size = int(os.getenv('MAX_FILE_SIZE', '104857600'))
        self.allowed_file_types = json.loads(os.getenv('ALLOWED_FILE_TYPES', '["csv", "json", "xlsx", "xls", "pdf", "txt", "docx", "doc"]'))
        self.batch_size = int(os.getenv('BATCH_SIZE', '100'))
        self.max_retry_attempts = int(os.getenv('MAX_RETRY_ATTEMPTS', '3'))
        self.processing_timeout = int(os.getenv('PROCESSING_TIMEOUT', '300'))

config = Config()

# RAG Fusion Query Expander
class RAGFusionQueryExpander:
    """Generate multiple query variations for RAG Fusion"""
    
    def __init__(self, llm):
        self.llm = llm
        self.query_expansion_prompt = PromptTemplate(
            template="""You are an expert at creating search query variations. Given an original query, generate {num_queries} different but related search queries that would help find comprehensive information about the topic.

Original Query: {original_query}

Requirements:
1. Each query should approach the topic from a different angle
2. Use synonyms and related terms
3. Include more specific and more general versions
4. Consider different contexts and perspectives
5. Keep queries focused and searchable

Generate {num_queries} alternative queries (one per line):""",
            input_variables=["original_query", "num_queries"]
        )
    
    def expand_query(self, query: str, num_queries: int = 3) -> List[str]:
        """Generate multiple query variations"""
        if not self.llm:
            # Fallback to simple variations
            return self._generate_simple_variations(query, num_queries)
        
        try:
            messages = [
                SystemMessage(content="You are an expert at query expansion for information retrieval."),
                HumanMessage(content=self.query_expansion_prompt.format(
                    original_query=query,
                    num_queries=num_queries
                ))
            ]
            
            response = self.llm.invoke(messages)
            queries = [q.strip() for q in response.content.split('\n') if q.strip()]
            
            # Always include original query
            all_queries = [query] + queries
            return all_queries[:num_queries + 1]
            
        except Exception as e:
            logger.warning(f"Query expansion failed: {e}")
            return self._generate_simple_variations(query, num_queries)
    
    def _generate_simple_variations(self, query: str, num_queries: int) -> List[str]:
        """Fallback method for generating query variations"""
        variations = [query]
        
        # Add question variations
        if not query.endswith('?'):
            variations.append(f"What is {query}?")
            variations.append(f"How does {query} work?")
        
        # Add specific terms
        words = query.split()
        if len(words) > 1:
            variations.append(' '.join(words[:len(words)//2]))  # First half
            variations.append(' '.join(words[len(words)//2:]))  # Second half
        
        return variations[:num_queries + 1]

# Reciprocal Rank Fusion
class ReciprocalRankFusion:
    """Implements Reciprocal Rank Fusion for combining multiple ranked lists"""
    
    def __init__(self, k: int = 60):
        self.k = k  # RRF constant
    
    def fuse_results(self, ranked_lists: List[List[Tuple[Any, float]]], weights: List[float] = None) -> List[Tuple[Any, float]]:
        """
        Combine multiple ranked lists using RRF
        
        Args:
            ranked_lists: List of ranked results [(item, score), ...]
            weights: Optional weights for each ranked list
        
        Returns:
            Fused ranked list [(item, fused_score), ...]
        """
        if not ranked_lists:
            return []
        
        if weights is None:
            weights = [1.0] * len(ranked_lists)
        
        # Collect all unique items
        all_items = set()
        for ranked_list in ranked_lists:
            for item, _ in ranked_list:
                all_items.add(self._get_item_key(item))
        
        # Calculate RRF scores
        rrf_scores = {}
        
        for item_key in all_items:
            rrf_score = 0.0
            
            for i, ranked_list in enumerate(ranked_lists):
                # Find rank of item in this list (1-indexed)
                rank = None
                for j, (item, _) in enumerate(ranked_list):
                    if self._get_item_key(item) == item_key:
                        rank = j + 1
                        break
                
                if rank is not None:
                    rrf_score += weights[i] * (1.0 / (self.k + rank))
            
            rrf_scores[item_key] = rrf_score
        
        # Create final ranked list
        fused_results = []
        item_map = {}
        
        # Create item mapping
        for ranked_list in ranked_lists:
            for item, _ in ranked_list:
                item_map[self._get_item_key(item)] = item
        
        # Sort by RRF score
        for item_key, score in sorted(rrf_scores.items(), key=lambda x: x[1], reverse=True):
            fused_results.append((item_map[item_key], score))
        
        return fused_results
    
    def _get_item_key(self, item) -> str:
        """Generate a unique key for an item"""
        if hasattr(item, 'page_content'):
            return hash(item.page_content[:100])
        elif isinstance(item, dict):
            return hash(str(item.get('content', item)))
        else:
            return hash(str(item))

# BM25 Sparse Retriever
class BM25SparseRetriever:
    """BM25 sparse retrieval implementation"""
    
    def __init__(self):
        self.bm25 = None
        self.documents = []
        self.tokenized_docs = []
    
    def fit(self, documents: List[LangChainDocument]):
        """Fit BM25 on document corpus"""
        if not BM25_AVAILABLE:
            logger.warning("BM25 not available")
            return False
        
        try:
            self.documents = documents
            
            # Tokenize documents
            self.tokenized_docs = []
            for doc in documents:
                tokens = self._tokenize(doc.page_content)
                self.tokenized_docs.append(tokens)
            
            # Fit BM25
            self.bm25 = BM25Okapi(self.tokenized_docs)
            logger.info(f"BM25 fitted on {len(documents)} documents")
            return True
            
        except Exception as e:
            logger.error(f"BM25 fitting failed: {e}")
            return False
    
    def search(self, query: str, top_k: int = 10) -> List[Tuple[LangChainDocument, float]]:
        """Search using BM25"""
        if not self.bm25 or not self.documents:
            return []
        
        try:
            query_tokens = self._tokenize(query)
            scores = self.bm25.get_scores(query_tokens)
            
            # Get top k results
            top_indices = np.argsort(scores)[::-1][:top_k]
            
            results = []
            for idx in top_indices:
                if scores[idx] > 0:  # Only include non-zero scores
                    results.append((self.documents[idx], float(scores[idx])))
            
            return results
            
        except Exception as e:
            logger.error(f"BM25 search failed: {e}")
            return []
    
    def _tokenize(self, text: str) -> List[str]:
        """Simple tokenization"""
        # Basic tokenization - can be enhanced with better tokenizers
        text = text.lower()
        tokens = re.findall(r'\b\w+\b', text)
        return tokens

# Cross-Encoder Reranker
class CrossEncoderReranker:
    """Cross-encoder for reranking retrieved documents"""
    
    def __init__(self, model_name: str = None):
        self.model_name = model_name or config.hf_reranker_model
        self.reranker = None
        self.setup_reranker()
    
    def setup_reranker(self):
        """Initialize cross-encoder model"""
        if not SENTENCE_TRANSFORMERS_AVAILABLE:
            logger.warning("SentenceTransformers not available for reranking")
            return False
        
        try:
            from sentence_transformers import CrossEncoder
            self.reranker = CrossEncoder(self.model_name)
            logger.info(f"✅ Cross-encoder reranker initialized: {self.model_name}")
            return True
        except Exception as e:
            logger.error(f"Cross-encoder initialization failed: {e}")
            return False
    
    def rerank(self, query: str, documents: List[Tuple[LangChainDocument, float]], top_k: int = None) -> List[Tuple[LangChainDocument, float]]:
        """Rerank documents using cross-encoder"""
        if not self.reranker or not documents:
            return documents
        
        try:
            # Prepare query-document pairs
            pairs = []
            for doc, _ in documents:
                pairs.append([query, doc.page_content[:512]])  # Limit content length
            
            # Get relevance scores
            scores = self.reranker.predict(pairs)
            
            # Combine with documents and sort
            reranked = []
            for i, (doc, _) in enumerate(documents):
                reranked.append((doc, float(scores[i])))
            
            # Sort by relevance score
            reranked.sort(key=lambda x: x[1], reverse=True)
            
            if top_k:
                reranked = reranked[:top_k]
            
            logger.info(f"Reranked {len(documents)} documents")
            return reranked
            
        except Exception as e:
            logger.error(f"Reranking failed: {e}")
            return documents

# Knowledge Graph Manager with Community Detection
class GraphRAGManager:
    """Enhanced knowledge graph management with community detection"""
    
    def __init__(self, neo4j_uri=None, neo4j_username=None, neo4j_password=None):
        self.driver = None
        self.graph = nx.Graph()  # NetworkX graph for analysis
        self.communities = []
        self.entity_embeddings = {}
        
        if NEO4J_AVAILABLE and neo4j_uri:
            self.setup_neo4j(neo4j_uri, neo4j_username, neo4j_password)
        else:
            logger.info("Using NetworkX for graph operations")
    
    def setup_neo4j(self, uri: str, username: str, password: str):
        """Setup Neo4j connection"""
        try:
            self.driver = GraphDatabase.driver(uri, auth=(username, password))
            # Test connection
            with self.driver.session() as session:
                session.run("RETURN 1")
            logger.info("✅ Neo4j connected successfully")
        except Exception as e:
            logger.error(f"Neo4j connection failed: {e}")
            self.driver = None
    
    def add_entities_and_relationships(self, entities: List[Dict], document_id: str) -> Tuple[int, int]:
        """Add entities and relationships to knowledge graph"""
        nodes_created = 0
        relationships_created = 0
        
        try:
            # Add to NetworkX graph
            for entity in entities:
                entity_id = f"{entity['text']}_{entity['label']}"
                
                # Add node to NetworkX graph
                self.graph.add_node(
                    entity_id,
                    text=entity['text'],
                    label=entity['label'],
                    confidence=entity.get('confidence', 0.0),
                    documents=[document_id]
                )
                nodes_created += 1
                
                # Store entity embedding if available
                if 'embedding' in entity:
                    self.entity_embeddings[entity_id] = entity['embedding']
            
            # Create relationships between entities in the same document
            entity_ids = [f"{e['text']}_{e['label']}" for e in entities]
            for i, entity1 in enumerate(entity_ids):
                for entity2 in entity_ids[i+1:]:
                    self.graph.add_edge(entity1, entity2, document=document_id, weight=1.0)
                    relationships_created += 1
            
            # Add to Neo4j if available
            if self.driver:
                nodes_neo4j, rels_neo4j = self._add_to_neo4j(entities, document_id)
                nodes_created = max(nodes_created, nodes_neo4j)
                relationships_created = max(relationships_created, rels_neo4j)
            
            logger.info(f"Added {nodes_created} nodes and {relationships_created} relationships")
            return nodes_created, relationships_created
            
        except Exception as e:
            logger.error(f"Error adding entities to graph: {e}")
            return 0, 0
    
    def _add_to_neo4j(self, entities: List[Dict], document_id: str) -> Tuple[int, int]:
        """Add entities to Neo4j"""
        if not self.driver:
            return 0, 0
        
        nodes_created = 0
        relationships_created = 0
        
        try:
            with self.driver.session() as session:
                # Create entities
                for entity in entities:
                    result = session.run("""
                        MERGE (e:Entity {text: $text, label: $label})
                        ON CREATE SET e.confidence = $confidence, e.documents = [$document_id]
                        ON MATCH SET e.documents = e.documents + $document_id
                        RETURN e
                    """, 
                    text=entity['text'], 
                    label=entity['label'],
                    confidence=entity.get('confidence', 0.0),
                    document_id=document_id
                    )
                    nodes_created += len(list(result))
                
                # Create relationships
                for i, entity1 in enumerate(entities):
                    for entity2 in entities[i+1:]:
                        result = session.run("""
                            MATCH (e1:Entity {text: $text1, label: $label1})
                            MATCH (e2:Entity {text: $text2, label: $label2})
                            MERGE (e1)-[r:CO_OCCURS]->(e2)
                            ON CREATE SET r.weight = 1, r.documents = [$document_id]
                            ON MATCH SET r.weight = r.weight + 1, r.documents = r.documents + $document_id
                            RETURN r
                        """,
                        text1=entity1['text'], label1=entity1['label'],
                        text2=entity2['text'], label2=entity2['label'],
                        document_id=document_id
                        )
                        relationships_created += len(list(result))
            
            return nodes_created, relationships_created
            
        except Exception as e:
            logger.error(f"Neo4j operations failed: {e}")
            return 0, 0
    
    def detect_communities(self) -> List[Dict[str, Any]]:
        """Detect communities in the knowledge graph"""
        if not NETWORKX_AVAILABLE or len(self.graph.nodes) < 3:
            return []
        
        try:
            # Use different community detection algorithms
            communities_detected = []
            
            # Louvain community detection
            if len(self.graph.nodes) > 0:
                communities_louvain = community.greedy_modularity_communities(self.graph)
                
                for i, comm in enumerate(communities_louvain):
                    if len(comm) >= 2:  # Only include communities with 2+ nodes
                        community_info = {
                            'id': f'community_{i}',
                            'algorithm': 'louvain',
                            'nodes': list(comm),
                            'size': len(comm),
                            'entities': []
                        }
                        
                        # Get entity information
                        for node in comm:
                            if node in self.graph.nodes:
                                node_data = self.graph.nodes[node]
                                community_info['entities'].append({
                                    'text': node_data.get('text', ''),
                                    'label': node_data.get('label', ''),
                                    'confidence': node_data.get('confidence', 0.0)
                                })
                        
                        communities_detected.append(community_info)
            
            self.communities = communities_detected
            logger.info(f"Detected {len(communities_detected)} communities")
            return communities_detected
            
        except Exception as e:
            logger.error(f"Community detection failed: {e}")
            return []
    
    def get_subgraph(self, entities: List[str], max_hops: int = 2) -> nx.Graph:
        """Extract subgraph around given entities"""
        if not NETWORKX_AVAILABLE:
            return nx.Graph()
        
        try:
            # Find all entity nodes matching the given entities
            entity_nodes = []
            for node in self.graph.nodes():
                node_data = self.graph.nodes[node]
                if node_data.get('text', '').lower() in [e.lower() for e in entities]:
                    entity_nodes.append(node)
            
            if not entity_nodes:
                return nx.Graph()
            
            # Get subgraph with multi-hop neighbors
            subgraph_nodes = set(entity_nodes)
            
            for hop in range(max_hops):
                new_nodes = set()
                for node in subgraph_nodes:
                    neighbors = list(self.graph.neighbors(node))
                    new_nodes.update(neighbors)
                subgraph_nodes.update(new_nodes)
            
            subgraph = self.graph.subgraph(subgraph_nodes).copy()
            logger.info(f"Extracted subgraph with {len(subgraph.nodes)} nodes")
            return subgraph
            
        except Exception as e:
            logger.error(f"Subgraph extraction failed: {e}")
            return nx.Graph()
    
    def get_related_entities(self, entity: str, max_results: int = 10) -> List[Dict[str, Any]]:
        """Get entities related to a given entity"""
        related = []
        
        try:
            # Find the entity node
            target_node = None
            for node in self.graph.nodes():
                node_data = self.graph.nodes[node]
                if node_data.get('text', '').lower() == entity.lower():
                    target_node = node
                    break
            
            if not target_node:
                return []
            
            # Get neighbors and their weights
            neighbors = list(self.graph.neighbors(target_node))
            
            for neighbor in neighbors[:max_results]:
                neighbor_data = self.graph.nodes[neighbor]
                edge_data = self.graph.edges[target_node, neighbor]
                
                related.append({
                    'text': neighbor_data.get('text', ''),
                    'label': neighbor_data.get('label', ''),
                    'confidence': neighbor_data.get('confidence', 0.0),
                    'relationship_weight': edge_data.get('weight', 1.0)
                })
            
            return related
            
        except Exception as e:
            logger.error(f"Error getting related entities: {e}")
            return []
    
    def close(self):
        """Close Neo4j connection"""
        if self.driver:
            self.driver.close()
            logger.info("Neo4j connection closed")

# Hybrid Retrieval Engine
class HybridRetrievalEngine:
    """Combines dense and sparse retrieval with adaptive weighting"""
    
    def __init__(self, dense_retriever, sparse_retriever=None, reranker=None):
        self.dense_retriever = dense_retriever
        self.sparse_retriever = sparse_retriever
        self.reranker = reranker
        self.rrf = ReciprocalRankFusion()
        
        # Query type classifier for adaptive weighting
        self.query_patterns = {
            'factual': re.compile(r'\b(what|who|when|where|which)\b', re.IGNORECASE),
            'procedural': re.compile(r'\b(how|step|process|method)\b', re.IGNORECASE),
            'analytical': re.compile(r'\b(why|analyze|compare|evaluate)\b', re.IGNORECASE),
            'conceptual': re.compile(r'\b(explain|define|describe|concept)\b', re.IGNORECASE)
        }
    
    def classify_query_type(self, query: str) -> str:
        """Classify query type for adaptive weighting"""
        query_lower = query.lower()
        
        for query_type, pattern in self.query_patterns.items():
            if pattern.search(query_lower):
                return query_type
        
        return 'general'
    
    def get_adaptive_weights(self, query: str) -> Tuple[float, float]:
        """Get adaptive weights based on query type"""
        query_type = self.classify_query_type(query)
        
        weight_mapping = {
            'factual': (0.3, 0.7),      # Favor sparse for facts
            'procedural': (0.6, 0.4),   # Favor dense for procedures
            'analytical': (0.8, 0.2),   # Strong favor for dense
            'conceptual': (0.7, 0.3),   # Favor dense for concepts
            'general': (0.6, 0.4)       # Default
        }
        
        return weight_mapping.get(query_type, (0.6, 0.4))
    
    def retrieve_hybrid(self, query: str, top_k: int = 10, dense_weight: float = None, sparse_weight: float = None) -> Dict[str, Any]:
        """Perform hybrid retrieval combining dense and sparse methods"""
        results = {
            'documents': [],
            'scores': {},
            'query_type': self.classify_query_type(query),
            'weights_used': {}
        }
        
        try:
            # Get adaptive weights if not provided
            if dense_weight is None or sparse_weight is None:
                dense_weight, sparse_weight = self.get_adaptive_weights(query)
            
            results['weights_used'] = {
                'dense': dense_weight,
                'sparse': sparse_weight
            }
            
            # Dense retrieval
            dense_results = []
            if self.dense_retriever:
                try:
                    dense_docs = self.dense_retriever.search_similar(query, k=top_k * 2)
                    dense_results = [(doc, 1.0) for doc in dense_docs]  # Normalize scores
                except Exception as e:
                    logger.warning(f"Dense retrieval failed: {e}")
            
            # Sparse retrieval
            sparse_results = []
            if self.sparse_retriever:
                try:
                    sparse_results = self.sparse_retriever.search(query, top_k=top_k * 2)
                except Exception as e:
                    logger.warning(f"Sparse retrieval failed: {e}")
            
            # Combine results using RRF
            if dense_results and sparse_results:
                fused_results = self.rrf.fuse_results(
                    [dense_results, sparse_results],
                    [dense_weight, sparse_weight]
                )
            elif dense_results:
                fused_results = dense_results
            elif sparse_results:
                fused_results = sparse_results
            else:
                fused_results = []
            
            # Rerank if reranker available
            if self.reranker and fused_results:
                try:
                    reranked_results = self.reranker.rerank(query, fused_results, top_k)
                    fused_results = reranked_results
                except Exception as e:
                    logger.warning(f"Reranking failed: {e}")
            
            # Prepare final results
            results['documents'] = [doc for doc, score in fused_results[:top_k]]
            results['scores'] = {
                'dense_count': len(dense_results),
                'sparse_count': len(sparse_results),
                'fused_count': len(fused_results),
                'final_count': len(results['documents'])
            }
            
            logger.info(f"Hybrid retrieval: {results['scores']}")
            return results
            
        except Exception as e:
            logger.error(f"Hybrid retrieval failed: {e}")
            return results

# RAG Fusion Pipeline
class RAGFusionPipeline:
    """Complete RAG Fusion implementation with parallel processing"""
    
    def __init__(self, query_expander, hybrid_retriever, llm):
        self.query_expander = query_expander
        self.hybrid_retriever = hybrid_retriever
        self.llm = llm
        self.rrf = ReciprocalRankFusion()
    
    async def process_fusion_query(self, original_query: str, num_queries: int = 3, top_k: int = 10) -> Dict[str, Any]:
        """Process query using RAG Fusion approach"""
        start_time = datetime.now()
        
        try:
            # Step 1: Query expansion
            expanded_queries = self.query_expander.expand_query(original_query, num_queries)
            logger.info(f"Generated {len(expanded_queries)} fusion queries")
            
            # Step 2: Parallel retrieval
            all_results = []
            
            # Use ThreadPoolExecutor for parallel retrieval
            with concurrent.futures.ThreadPoolExecutor(max_workers=min(len(expanded_queries), 5)) as executor:
                future_to_query = {}
                
                for query in expanded_queries:
                    future = executor.submit(
                        self.hybrid_retriever.retrieve_hybrid,
                        query,
                        top_k * 2  # Get more results for fusion
                    )
                    future_to_query[future] = query
                
                # Collect results
                for future in concurrent.futures.as_completed(future_to_query):
                    query = future_to_query[future]
                    try:
                        result = future.result()
                        if result['documents']:
                            # Convert to (document, score) tuples
                            doc_scores = [(doc, 1.0) for doc in result['documents']]
                            all_results.append(doc_scores)
                    except Exception as e:
                        logger.warning(f"Retrieval failed for query '{query}': {e}")
            
            # Step 3: Fusion using RRF
            if all_results:
                fused_results = self.rrf.fuse_results(all_results)
                final_documents = [doc for doc, score in fused_results[:top_k]]
            else:
                final_documents = []
            
            # Step 4: Deduplicate based on content similarity
            deduplicated_docs = self._deduplicate_documents(final_documents)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                'documents': deduplicated_docs,
                'fusion_queries': expanded_queries,
                'retrieval_results_count': len(all_results),
                'fused_results_count': len(fused_results) if all_results else 0,
                'final_count': len(deduplicated_docs),
                'processing_time': processing_time
            }
            
        except Exception as e:
            logger.error(f"RAG Fusion processing failed: {e}")
            return {
                'documents': [],
                'fusion_queries': [original_query],
                'retrieval_results_count': 0,
                'fused_results_count': 0,
                'final_count': 0,
                'processing_time': 0
            }
    
    def _deduplicate_documents(self, documents: List[LangChainDocument], similarity_threshold: float = 0.8) -> List[LangChainDocument]:
        """Remove near-duplicate documents based on content similarity"""
        if not documents or not SKLEARN_AVAILABLE:
            return documents
        
        try:
            # Create TF-IDF vectors for documents
            texts = [doc.page_content for doc in documents]
            vectorizer = TfidfVectorizer(max_features=1000, stop_words='english')
            tfidf_matrix = vectorizer.fit_transform(texts)
            
            # Calculate similarity matrix
            similarity_matrix = cosine_similarity(tfidf_matrix)
            
            # Find duplicates
            to_remove = set()
            for i in range(len(documents)):
                if i in to_remove:
                    continue
                for j in range(i + 1, len(documents)):
                    if j in to_remove:
                        continue
                    if similarity_matrix[i][j] > similarity_threshold:
                        to_remove.add(j)  # Remove the later document
            
            # Return deduplicated documents
            deduplicated = [doc for i, doc in enumerate(documents) if i not in to_remove]
            
            if len(to_remove) > 0:
                logger.info(f"Removed {len(to_remove)} duplicate documents")
            
            return deduplicated
            
        except Exception as e:
            logger.warning(f"Document deduplication failed: {e}")
            return documents

# Enhanced Entity Extractor with Graph Integration
class EnhancedEntityExtractor:
    """Enhanced entity extraction with graph integration"""
    
    def __init__(self, embedding_manager=None):
        self.nlp = nlp if SPACY_AVAILABLE else None
        self.embedding_manager = embedding_manager
        self.custom_patterns = [
            # Financial patterns
            (r'\$[\d,]+(?:\.\d{2})?', 'MONEY'),
            (r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', 'DATE'),
            (r'\b\d+%\b', 'PERCENTAGE'),
            # Technical patterns
            (r'\b[A-Z]{2,10}\b', 'ACRONYM'),
            (r'\b\w+@\w+\.\w+\b', 'EMAIL'),
            (r'\bhttps?://[^\s]+\b', 'URL'),
            # Business patterns
            (r'\b[A-Z][a-z]+ (?:Inc|Corp|LLC|Ltd|Co)\b', 'COMPANY'),
            (r'\b(?:CEO|CTO|CFO|VP|President|Director)\b', 'TITLE'),
        ]
    
    def extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract entities from text with embeddings"""
        entities = []
        
        # Spacy NER
        if self.nlp:
            try:
                doc = self.nlp(text)
                for ent in doc.ents:
                    entity_data = {
                        'text': ent.text,
                        'label': ent.label_,
                        'start': ent.start_char,
                        'end': ent.end_char,
                        'confidence': 1.0,
                        'method': 'spacy'
                    }
                    
                    # Add embedding if available
                    if self.embedding_manager:
                        try:
                            embedding = self.embedding_manager.embed_query(ent.text)
                            entity_data['embedding'] = embedding
                        except Exception as e:
                            logger.warning(f"Entity embedding failed for '{ent.text}': {e}")
                    
                    entities.append(entity_data)
            except Exception as e:
                logger.warning(f"Spacy NER failed: {e}")
        
        # Custom regex patterns
        for pattern, label in self.custom_patterns:
            try:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    entity_data = {
                        'text': match.group(),
                        'label': label,
                        'start': match.start(),
                        'end': match.end(),
                        'confidence': 0.8,
                        'method': 'regex'
                    }
                    
                    # Add embedding if available
                    if self.embedding_manager:
                        try:
                            embedding = self.embedding_manager.embed_query(match.group())
                            entity_data['embedding'] = embedding
                        except Exception as e:
                            logger.warning(f"Entity embedding failed for '{match.group()}': {e}")
                    
                    entities.append(entity_data)
            except Exception as e:
                logger.warning(f"Regex pattern {pattern} failed: {e}")
        
        # Remove duplicates and sort by confidence
        unique_entities = []
        seen = set()
        for entity in sorted(entities, key=lambda x: x['confidence'], reverse=True):
            key = (entity['text'].lower(), entity['label'])
            if key not in seen:
                seen.add(key)
                unique_entities.append(entity)
        
        return unique_entities[:30]  # Increased limit for better graph construction

# Semantic Chunker using CLIP embeddings
class SemanticChunker:
    """Advanced semantic chunking based on content similarity"""
    
    def __init__(self, clip_model, clip_processor, max_chunk_size=1000, similarity_threshold=0.7):
        self.clip_model = clip_model
        self.clip_processor = clip_processor
        self.max_chunk_size = max_chunk_size
        self.similarity_threshold = similarity_threshold
        self.basic_splitter = RecursiveCharacterTextSplitter(
            chunk_size=200,
            chunk_overlap=50,
            separators=["\n\n", "\n", ". ", " "]
        )
    
    def embed_text(self, text: str) -> np.ndarray:
        """Create CLIP embedding for text"""
        try:
            inputs = self.clip_processor(
                text=text,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=77
            )
            
            with torch.no_grad():
                features = self.clip_model.get_text_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
            
            return features.squeeze().numpy()
        except Exception as e:
            logger.warning(f"CLIP text embedding failed: {e}")
            return np.zeros(512)  # Default embedding size
    
    def semantic_chunk(self, text: str) -> List[str]:
        """Create semantically coherent chunks"""
        # First, create basic chunks
        basic_chunks = self.basic_splitter.split_text(text)
        
        if len(basic_chunks) <= 1:
            return basic_chunks
        
        # Create embeddings for each chunk
        chunk_embeddings = []
        for chunk in basic_chunks:
            embedding = self.embed_text(chunk)
            chunk_embeddings.append(embedding)
        
        # Group similar chunks together
        semantic_chunks = []
        current_chunk = [basic_chunks[0]]
        current_size = len(basic_chunks[0])
        
        for i in range(1, len(basic_chunks)):
            # Calculate similarity with current chunk
            if len(chunk_embeddings[i-1]) > 0 and len(chunk_embeddings[i]) > 0:
                if SKLEARN_AVAILABLE:
                    similarity = cosine_similarity(
                        chunk_embeddings[i-1].reshape(1, -1),
                        chunk_embeddings[i].reshape(1, -1)
                    )[0][0]
                else:
                    # Fallback similarity calculation
                    dot_product = np.dot(chunk_embeddings[i-1], chunk_embeddings[i])
                    norm_a = np.linalg.norm(chunk_embeddings[i-1])
                    norm_b = np.linalg.norm(chunk_embeddings[i])
                    similarity = dot_product / (norm_a * norm_b) if norm_a > 0 and norm_b > 0 else 0.5
            else:
                similarity = 0.5
            
            # Check if we should merge chunks
            new_size = current_size + len(basic_chunks[i])
            
            if (similarity > self.similarity_threshold and 
                new_size <= self.max_chunk_size):
                current_chunk.append(basic_chunks[i])
                current_size = new_size
            else:
                # Finalize current chunk and start new one
                semantic_chunks.append(" ".join(current_chunk))
                current_chunk = [basic_chunks[i]]
                current_size = len(basic_chunks[i])
        
        # Add the last chunk
        if current_chunk:
            semantic_chunks.append(" ".join(current_chunk))
        
        return semantic_chunks

# Multi-hop Reasoning Engine with Graph Enhancement
class GraphEnhancedMultiHopReasoner:
    """Enhanced multi-hop reasoning using knowledge graph"""
    
    def __init__(self, retriever, llm, graph_manager):
        self.retriever = retriever
        self.llm = llm
        self.graph_manager = graph_manager
        self.reasoning_graph = nx.DiGraph()
    
    def decompose_query(self, query: str) -> List[str]:
        """Decompose complex query into sub-questions"""
        if not self.llm:
            return [query]
        
        try:
            decomposition_prompt = f"""
            Break down this complex question into 2-4 simpler sub-questions that need to be answered step by step:
            
            Original question: {query}
            
            Sub-questions (one per line):
            """
            
            messages = [
                SystemMessage(content="You are an expert at breaking down complex questions into simpler sub-questions."),
                HumanMessage(content=decomposition_prompt)
            ]
            
            response = self.llm.invoke(messages)
            sub_questions = [q.strip() for q in response.content.split('\n') if q.strip() and not q.strip().startswith('Sub-questions')]
            
            return sub_questions[:4] if sub_questions else [query]
            
        except Exception as e:
            logger.warning(f"Query decomposition failed: {e}")
            return [query]
    
    def reason_step_with_graph(self, question: str, context: str = "", entities: List[str] = None) -> Dict[str, Any]:
        """Perform one reasoning step enhanced with graph information"""
        try:
            # Get graph-enhanced context
            graph_context = ""
            if entities and self.graph_manager:
                subgraph = self.graph_manager.get_subgraph(entities, max_hops=2)
                if subgraph.nodes:
                    graph_entities = []
                    for node in subgraph.nodes:
                        node_data = subgraph.nodes[node]
                        graph_entities.append(f"{node_data.get('text', '')} ({node_data.get('label', '')})")
                    graph_context = f"Related entities from knowledge graph: {', '.join(graph_entities[:10])}"
            
            # Retrieve relevant documents
            if hasattr(self.retriever, 'retrieve_hybrid'):
                results = self.retriever.retrieve_hybrid(question, top_k=5)
                retrieved_docs = results.get('documents', [])
            else:
                retrieved_docs = []
            
            # Prepare context
            doc_context = "\n".join([doc.page_content for doc in retrieved_docs[:3]])
            full_context = f"{context}\n{doc_context}\n{graph_context}".strip()
            
            # Generate answer
            if self.llm and full_context:
                reasoning_prompt = f"""
                Based on the following context, answer this question step by step:
                
                Context: {full_context}
                Question: {question}
                
                Provide a clear, factual answer with reasoning:
                """
                
                messages = [
                    SystemMessage(content="You are an expert analyst who provides step-by-step reasoning using both document context and knowledge graph information."),
                    HumanMessage(content=reasoning_prompt)
                ]
                
                response = self.llm.invoke(messages)
                answer = response.content
            else:
                answer = f"Insufficient context to answer: {question}"
            
            return {
                'question': question,
                'answer': answer,
                'context': full_context,
                'sources': retrieved_docs,
                'graph_entities': entities or []
            }
            
        except Exception as e:
            logger.error(f"Graph-enhanced reasoning step failed: {e}")
            return {
                'question': question,
                'answer': f"Error in reasoning: {str(e)}",
                'context': context,
                'sources': [],
                'graph_entities': []
            }
    
    def multi_hop_reason(self, query: str) -> List[Dict[str, Any]]:
        """Perform multi-hop reasoning with graph enhancement"""
        # Decompose query
        sub_questions = self.decompose_query(query)
        reasoning_steps = []
        accumulated_context = ""
        
        # Extract entities from original query for graph traversal
        if hasattr(self, 'entity_extractor'):
            query_entities = [ent['text'] for ent in self.entity_extractor.extract_entities(query)]
        else:
            query_entities = []
        
        for i, question in enumerate(sub_questions):
            step_result = self.reason_step_with_graph(
                question, 
                accumulated_context, 
                query_entities
            )
            reasoning_steps.append(step_result)
            
            # Accumulate context for next step
            accumulated_context += f"\n\nStep {i+1}: {step_result['answer']}"
        
        return reasoning_steps

# Recommendation Engine
class RecommendationEngine:
    """Generate recommendations based on query and context"""
    
    def __init__(self, llm, entity_extractor):
        self.llm = llm
        self.entity_extractor = entity_extractor
    
    def generate_recommendations(self, query: str, answer: str, entities: List[Dict], sources: List) -> List[str]:
        """Generate contextual recommendations"""
        if not self.llm:
            return ["Enable LLM to get personalized recommendations"]
        
        try:
            # Extract key topics from entities
            key_topics = [ent['text'] for ent in entities if ent['label'] in ['ORG', 'PRODUCT', 'TECH']][:5]
            
            # Create recommendation prompt
            recommendation_prompt = f"""
            Based on the user's query and the information provided, suggest 3-5 practical next steps or related topics they might want to explore:
            
            User Query: {query}
            Answer Context: {answer[:500]}...
            Key Topics: {', '.join(key_topics)}
            
            Provide actionable recommendations (one per line):
            """
            
            messages = [
                SystemMessage(content="You are an expert advisor who provides practical, actionable recommendations."),
                HumanMessage(content=recommendation_prompt)
            ]
            
            response = self.llm.invoke(messages)
            recommendations = [rec.strip() for rec in response.content.split('\n') if rec.strip() and not rec.strip().startswith('Provide')]
            
            return recommendations[:5]
            
        except Exception as e:
            logger.warning(f"Recommendation generation failed: {e}")
            return ["Explore related documents", "Search for more specific information", "Consider different perspectives on this topic"]

# CLIP Multimodal Manager
class CLIPMultimodalManager:
    """Manages CLIP model for unified text and image embeddings"""
    
    def __init__(self, model_name: str = None):
        self.model_name = model_name or config.clip_model_name
        self.clip_model = None
        self.clip_processor = None
        self.image_data_store = {}
        
        self.setup_clip()
    
    def setup_clip(self):
        """Initialize CLIP model and processor"""
        try:
            self.clip_model = CLIPModel.from_pretrained(self.model_name)
            self.clip_processor = CLIPProcessor.from_pretrained(self.model_name)
            self.clip_model.eval()
            logger.info(f"✅ CLIP model initialized: {self.model_name}")
        except Exception as e:
            logger.error(f"Failed to initialize CLIP: {e}")
            self.clip_model = None
            self.clip_processor = None
    
    def embed_image(self, image_data) -> np.ndarray:
        """Embed image using CLIP"""
        if not self.clip_model:
            return np.zeros(512)
        
        try:
            if isinstance(image_data, str):  # If path
                image = Image.open(image_data).convert("RGB")
            else:  # If PIL Image
                image = image_data
            
            inputs = self.clip_processor(images=image, return_tensors="pt")
            
            with torch.no_grad():
                features = self.clip_model.get_image_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
            
            return features.squeeze().numpy()
        except Exception as e:
            logger.warning(f"Image embedding failed: {e}")
            return np.zeros(512)
    
    def embed_text(self, text: str) -> np.ndarray:
        """Embed text using CLIP"""
        if not self.clip_model:
            return np.zeros(512)
        
        try:
            inputs = self.clip_processor(
                text=text,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=77
            )
            
            with torch.no_grad():
                features = self.clip_model.get_text_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
            
            return features.squeeze().numpy()
        except Exception as e:
            logger.warning(f"Text embedding failed: {e}")
            return np.zeros(512)
    
    def store_image(self, image_id: str, pil_image: Image.Image) -> str:
        """Store image as base64 for GPT-4V"""
        try:
            buffered = io.BytesIO()
            pil_image.save(buffered, format="PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode()
            self.image_data_store[image_id] = img_base64
            return img_base64
        except Exception as e:
            logger.error(f"Failed to store image {image_id}: {e}")
            return ""
    
    def get_stored_image(self, image_id: str) -> str:
        """Retrieve stored image base64"""
        return self.image_data_store.get(image_id, "")

# Advanced Docling Document Processor
class DoclingProcessor:
    """Advanced document processing using Docling"""
    
    def __init__(self):
        self.converter = None
        if DOCLING_AVAILABLE:
            try:
                # Configure Docling with advanced options
                pipeline_options = PdfPipelineOptions()
                pipeline_options.do_ocr = True
                pipeline_options.do_table_structure = True
                pipeline_options.table_structure_options.do_cell_matching = True
                
                self.converter = DocumentConverter(
                    format_options={
                        InputFormat.PDF: pipeline_options,
                    }
                )
                logger.info("✅ Docling processor initialized")
            except Exception as e:
                logger.warning(f"Docling initialization failed: {e}")
                self.converter = None
        else:
            logger.info("Docling not available")
    
    def process_document_advanced(self, file_path: str) -> Dict[str, Any]:
        """Process document using Docling for advanced structure extraction"""
        if not self.converter:
            return self.fallback_processing(file_path)
        
        try:
            # Convert document
            result = self.converter.convert(file_path)
            
            # Extract structured content
            doc_content = {
                'text': result.document.export_to_markdown(),
                'tables': [],
                'figures': [],
                'metadata': {
                    'title': getattr(result.document, 'title', ''),
                    'page_count': len(result.document.pages) if hasattr(result.document, 'pages') else 0,
                    'processing_method': 'docling'
                }
            }
            
            # Extract tables and figures if available
            if hasattr(result.document, 'tables'):
                for table in result.document.tables:
                    doc_content['tables'].append({
                        'content': str(table),
                        'page': getattr(table, 'page', 0)
                    })
            
            if hasattr(result.document, 'figures'):
                for figure in result.document.figures:
                    doc_content['figures'].append({
                        'content': str(figure),
                        'page': getattr(figure, 'page', 0)
                    })
            
            return doc_content
            
        except Exception as e:
            logger.warning(f"Docling processing failed for {file_path}: {e}")
            return self.fallback_processing(file_path)
    
    def fallback_processing(self, file_path: str) -> Dict[str, Any]:
        """Fallback processing when Docling is not available"""
        try:
            if file_path.lower().endswith('.pdf'):
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                return {
                    'text': text,
                    'tables': [],
                    'figures': [],
                    'metadata': {
                        'title': Path(file_path).stem,
                        'page_count': len(doc) if 'doc' in locals() else 0,
                        'processing_method': 'fallback_pdf'
                    }
                }
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                
                return {
                    'text': text,
                    'tables': [],
                    'figures': [],
                    'metadata': {
                        'title': Path(file_path).stem,
                        'page_count': 1,
                        'processing_method': 'fallback_text'
                    }
                }
        except Exception as e:
            logger.error(f"Fallback processing failed: {e}")
            return {
                'text': "",
                'tables': [],
                'figures': [],
                'metadata': {'processing_method': 'failed'}
            }

# Enhanced Document Processor with Multimodal Support
class EnhancedDocumentProcessor:
    """Enhanced document processor with multimodal and advanced processing capabilities"""
    
    def __init__(self, clip_manager: CLIPMultimodalManager, entity_extractor: EnhancedEntityExtractor):
        self.clip_manager = clip_manager
        self.entity_extractor = entity_extractor
        self.docling_processor = DoclingProcessor()
        
        # Initialize semantic chunker
        if self.clip_manager.clip_model:
            self.semantic_chunker = SemanticChunker(
                self.clip_manager.clip_model,
                self.clip_manager.clip_processor
            )
        else:
            self.semantic_chunker = None
        
        # Fallback text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    async def process_pdf_multimodal(self, file_path: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Process PDF with text, images, and advanced structure extraction"""
        documents = []
        image_docs = []
        entities_list = []
        
        try:
            # Use Docling for advanced processing
            docling_result = self.docling_processor.process_document_advanced(file_path)
            
            # Process text content
            if docling_result['text']:
                # Extract entities from full text
                entities = self.entity_extractor.extract_entities(docling_result['text'])
                entities_list.extend(entities)
                
                # Create semantic chunks if available
                if self.semantic_chunker:
                    chunks = self.semantic_chunker.semantic_chunk(docling_result['text'])
                else:
                    chunks = self.text_splitter.split_text(docling_result['text'])
                
                # Create document objects for each chunk
                for i, chunk in enumerate(chunks):
                    chunk_entities = self.entity_extractor.extract_entities(chunk)
                    
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "type": "text",
                            "chunk_id": i,
                            "chunk_size": len(chunk),
                            "entities": [ent['text'] for ent in chunk_entities],
                            "processing_method": docling_result['metadata']['processing_method']
                        }
                    )
                    documents.append(doc)
            
            # Process images from PDF using PyMuPDF
            try:
                pdf_doc = fitz.open(file_path)
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc.load_page(page_num)
                    
                    for img_index, img in enumerate(page.get_images(full=True)):
                        try:
                            xref = img[0]
                            base_image = pdf_doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # Convert to PIL Image
                            pil_image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                            
                            # Create unique identifier
                            image_id = f"page_{page_num}_img_{img_index}"
                            
                            # Store image
                            img_base64 = self.clip_manager.store_image(image_id, pil_image)
                            
                            # Create image document
                            image_doc = {
                                "image_id": image_id,
                                "page": page_num,
                                "base64": img_base64,
                                "embedding": self.clip_manager.embed_image(pil_image),
                                "metadata": {
                                    "source": file_path,
                                    "type": "image",
                                    "page": page_num,
                                    "image_index": img_index
                                }
                            }
                            image_docs.append(image_doc)
                            
                        except Exception as e:
                            logger.warning(f"Error processing image {img_index} on page {page_num}: {e}")
                            continue
                
                pdf_doc.close()
                
            except Exception as e:
                logger.warning(f"Image extraction failed: {e}")
            
            return documents, image_docs, entities_list
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    async def process_document(self, file_path: str, file_type: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Main document processing method"""
        if file_type.lower() == 'pdf':
            return await self.process_pdf_multimodal(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return await self.process_docx_enhanced(file_path)
        elif file_type.lower() == 'txt':
            return await self.process_txt_enhanced(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def process_docx_enhanced(self, file_path: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Enhanced DOCX processing"""
        documents = []
        entities_list = []
        
        try:
            # Use Docling if available
            docling_result = self.docling_processor.process_document_advanced(file_path)
            
            if docling_result['text']:
                entities = self.entity_extractor.extract_entities(docling_result['text'])
                entities_list.extend(entities)
                
                if self.semantic_chunker:
                    chunks = self.semantic_chunker.semantic_chunk(docling_result['text'])
                else:
                    chunks = self.text_splitter.split_text(docling_result['text'])
                
                for i, chunk in enumerate(chunks):
                    chunk_entities = self.entity_extractor.extract_entities(chunk)
                    
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "type": "docx",
                            "chunk_id": i,
                            "chunk_size": len(chunk),
                            "entities": [ent['text'] for ent in chunk_entities],
                            "processing_method": docling_result['metadata']['processing_method']
                        }
                    )
                    documents.append(doc)
            
            return documents, [], entities_list  # No images for DOCX yet
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def process_txt_enhanced(self, file_path: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Enhanced text processing"""
        documents = []
        entities_list = []
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            entities = self.entity_extractor.extract_entities(content)
            entities_list.extend(entities)
            
            if self.semantic_chunker:
                chunks = self.semantic_chunker.semantic_chunk(content)
            else:
                chunks = self.text_splitter.split_text(content)
            
            for i, chunk in enumerate(chunks):
                chunk_entities = self.entity_extractor.extract_entities(chunk)
                
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "type": "txt",
                        "chunk_id": i,
                        "chunk_size": len(chunk),
                        "entities": [ent['text'] for ent in chunk_entities],
                        "processing_method": "enhanced_text"
                    }
                )
                documents.append(doc)
            
            return documents, [], entities_list
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise

# HuggingFace Embedding Manager (Enhanced)
class HuggingFaceEmbeddingManager:
    """Enhanced HuggingFace embeddings with fallback options"""
    
    def __init__(self, model_name: str = None, cache_dir: str = None, device: str = 'cpu'):
        self.model_name = model_name or config.hf_embedding_model
        self.cache_dir = cache_dir or config.hf_cache_dir
        self.device = device or config.hf_device
        self.embedding_model = None
        self.llamaindex_embedding = None
        
        Path(self.cache_dir).mkdir(parents=True, exist_ok=True)
        self.setup_embeddings()
    
    def setup_embeddings(self):
        """Setup HuggingFace embeddings"""
        if not HUGGINGFACE_AVAILABLE:
            logger.error("HuggingFace transformers not available")
            return False
        
        try:
            self.embedding_model = SentenceTransformer(
                self.model_name,
                cache_folder=self.cache_dir,
                device=self.device
            )
            
            if LLAMAINDEX_AVAILABLE:
                try:
                    self.llamaindex_embedding = HuggingFaceEmbedding(
                        model_name=self.model_name,
                        cache_folder=self.cache_dir,
                        device=self.device
                    )
                    
                    if hasattr(Settings, 'embed_model'):
                        Settings.embed_model = self.llamaindex_embedding
                    
                    logger.info(f"✅ LlamaIndex HuggingFace embeddings initialized")
                except Exception as e:
                    logger.warning(f"LlamaIndex embedding setup failed: {e}")
                    self.llamaindex_embedding = None
            
            logger.info(f"✅ HuggingFace embeddings initialized: {self.model_name}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to initialize HuggingFace embeddings: {e}")
            return False
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a single query text"""
        if not self.embedding_model:
            raise ValueError("Embedding model not initialized")
        
        try:
            embedding = self.embedding_model.encode(text, normalize_embeddings=True)
            return embedding.tolist()
        except Exception as e:
            logger.error(f"Query embedding failed: {e}")
            raise
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed multiple documents"""
        if not self.embedding_model:
            raise ValueError("Embedding model not initialized")
        
        try:
            embeddings = self.embedding_model.encode(texts, normalize_embeddings=True)
            return embeddings.tolist()
        except Exception as e:
            logger.error(f"Document embedding failed: {e}")
            raise

# Enhanced Weaviate Manager with Hybrid Support
class EnhancedWeaviateManager:
    """Enhanced Weaviate manager with multimodal and hybrid search support"""
    
    def __init__(self):
        self.client = None
        self.vector_store = None
        self.connection_type = None
        self.image_documents = []
        self.bm25_retriever = BM25SparseRetriever()
        
        if WEAVIATE_AVAILABLE:
            self.setup_weaviate()
        else:
            logger.info("Weaviate not available - will use fallback storage")
    
    def setup_weaviate(self):
        """Setup Weaviate with proper connection"""
        if not WEAVIATE_AVAILABLE:
            return False
        
        try:
            self.client = weaviate.connect_to_local()
            
            if self.client.is_ready():
                logger.info("✅ Weaviate connected successfully (local)")
                self.connection_type = "local"
                return True
            else:
                self.client = None
                
        except Exception as e:
            logger.warning(f"Local Weaviate connection failed: {e}")
            self.client = None
        
        if config.weaviate_api_key and not self.client:
            try:
                self.client = weaviate.connect_to_weaviate_cloud(
                    cluster_url=config.weaviate_url,
                    auth_credentials=weaviate.auth.AuthApiKey(config.weaviate_api_key)
                )
                
                if self.client.is_ready():
                    logger.info("✅ Weaviate connected successfully (cloud)")
                    self.connection_type = "cloud"
                    return True
                else:
                    self.client = None
                    
            except Exception as e:
                logger.warning(f"Weaviate Cloud connection failed: {e}")
                self.client = None
        
        return False
    
    def create_vector_store(self, documents: List[LangChainDocument], embeddings):
        """Create vector store from documents and setup BM25"""
        if not self.client or not documents:
            return None
        
        try:
            self.vector_store = WeaviateVectorStore.from_documents(
                documents,
                embeddings,
                client=self.client,
                index_name="RagDocuments"
            )
            
            # Setup BM25 for sparse retrieval
            self.bm25_retriever.fit(documents)
            
            logger.info(f"Vector store created with {len(documents)} documents")
            logger.info("BM25 sparse retriever initialized")
            return self.vector_store
            
        except Exception as e:
            logger.error(f"Failed to create vector store: {e}")
            return None
    
    def add_documents_to_store(self, documents: List[LangChainDocument]):
        """Add documents to existing vector store and update BM25"""
        if not self.vector_store or not documents:
            return False
        
        try:
            self.vector_store.add_documents(documents)
            
            # Update BM25 with all documents
            all_docs = self.get_all_documents() + documents
            self.bm25_retriever.fit(all_docs)
            
            logger.info(f"Added {len(documents)} documents to vector store")
            return True
        except Exception as e:
            logger.error(f"Failed to add documents to vector store: {e}")
            return False
    
    def get_all_documents(self) -> List[LangChainDocument]:
        """Get all documents from vector store (for BM25 updates)"""
        # This is a simplified implementation
        # In practice, you might need to query Weaviate to get all documents
        return []
    
    def add_image_documents(self, image_docs: List[Dict]):
        """Store image documents separately"""
        self.image_documents.extend(image_docs)
        logger.info(f"Added {len(image_docs)} image documents")
    
    def search_similar(self, query: str, k: int = 5):
        """Dense vector search for similar documents"""
        if not self.vector_store:
            return []
        
        try:
            results = self.vector_store.similarity_search(query, k=k)
            return results
        except Exception as e:
            logger.error(f"Vector search failed: {e}")
            return []
    
    def search_sparse(self, query: str, k: int = 5):
        """Sparse BM25 search"""
        try:
            results = self.bm25_retriever.search(query, top_k=k)
            return [doc for doc, score in results]
        except Exception as e:
            logger.error(f"Sparse search failed: {e}")
            return []
    
    def search_images(self, query: str, clip_manager: CLIPMultimodalManager, k: int = 3):
        """Search for similar images using CLIP"""
        if not self.image_documents or not clip_manager.clip_model:
            return []
        
        try:
            query_embedding = clip_manager.embed_text(query)
            
            # Calculate similarities
            similarities = []
            for img_doc in self.image_documents:
                img_embedding = img_doc.get('embedding', np.zeros(512))
                if len(img_embedding) > 0 and len(query_embedding) > 0:
                    similarity = cosine_similarity(
                        query_embedding.reshape(1, -1),
                        img_embedding.reshape(1, -1)
                    )[0][0]
                    similarities.append((img_doc, similarity))
            
            # Sort by similarity and return top k
            similarities.sort(key=lambda x: x[1], reverse=True)
            return [img_doc for img_doc, sim in similarities[:k]]
            
        except Exception as e:
            logger.error(f"Image search failed: {e}")
            return []
    
    def close(self):
        """Close Weaviate connection"""
        if self.client:
            try:
                self.client.close()
                logger.info("Weaviate connection closed")
            except Exception as e:
                logger.warning(f"Error closing Weaviate connection: {e}")

# Enhanced RAG Bot with All Advanced Features
class EnhancedRAGBot:
    """Enhanced RAG Bot with RAG Fusion, Hybrid Search, GraphRAG, and all advanced features"""
    
    def __init__(self):
        # Initialize core components
        self.embedding_manager = HuggingFaceEmbeddingManager()
        self.clip_manager = CLIPMultimodalManager()
        self.entity_extractor = EnhancedEntityExtractor(self.embedding_manager)
        self.weaviate_manager = EnhancedWeaviateManager()
        
        # Initialize graph manager
        self.graph_manager = GraphRAGManager(
            config.neo4j_uri,
            config.neo4j_username,
            config.neo4j_password
        )
        
        # Initialize enhanced document processor
        self.document_processor = EnhancedDocumentProcessor(
            self.clip_manager, 
            self.entity_extractor
        )
        
        # Initialize Azure OpenAI LLM
        try:
            from langchain_openai import AzureChatOpenAI
            self.llm = AzureChatOpenAI(
                azure_endpoint=config.azure_openai_endpoint,
                api_key=config.azure_openai_api_key,
                api_version=config.azure_openai_api_version,
                deployment_name=config.azure_openai_deployment_name,
                model_name=config.azure_openai_model_name,
                temperature=config.azure_openai_temperature,
                max_tokens=config.azure_openai_max_tokens
            )
            logger.info("✅ Azure OpenAI LLM initialized")
        except Exception as e:
            logger.error(f"Failed to initialize LLM: {e}")
            self.llm = None
        
        # Initialize advanced components
        self.query_expander = RAGFusionQueryExpander(self.llm)
        self.reranker = CrossEncoderReranker()
        self.hybrid_retriever = HybridRetrievalEngine(
            self.weaviate_manager,
            self.weaviate_manager.bm25_retriever,
            self.reranker
        )
        self.rag_fusion_pipeline = RAGFusionPipeline(
            self.query_expander,
            self.hybrid_retriever,
            self.llm
        )
        
        # Initialize reasoning and recommendation engines
        self.multi_hop_reasoner = GraphEnhancedMultiHopReasoner(
            self.hybrid_retriever,
            self.llm,
            self.graph_manager
        )
        self.recommendation_engine = RecommendationEngine(self.llm, self.entity_extractor)
        
        # Initialize document storage
        self.documents = []
        self.all_entities = []
        
        # Setup enhanced prompts
        self.setup_enhanced_prompts()
    
    def setup_enhanced_prompts(self):
        """Setup enhanced prompt templates"""
        self.multimodal_qa_prompt = PromptTemplate(
            template="""You are an advanced AI assistant with comprehensive knowledge analysis capabilities. Analyze text, visual, graph, and reasoning information to provide authoritative answers.

Text Context from Multiple Sources:
{text_context}

Image Context:
{image_context}

Knowledge Graph Entities:
{entities}

Multi-hop Reasoning Chain:
{reasoning_steps}

RAG Fusion Queries Used:
{fusion_queries}

Hybrid Search Information:
- Query Type: {query_type}
- Dense/Sparse Weights: {search_weights}

Question: {question}

Instructions:
1. Synthesize information from all sources (text, images, entities, reasoning, graph)
2. Provide a comprehensive, well-structured answer
3. Reference specific sources and reasoning steps
4. If images are relevant, describe their relationship to the question
5. Use the knowledge graph connections to provide deeper insights
6. Build upon the multi-hop reasoning chain
7. Be precise, factual, and avoid speculation
8. Highlight any conflicting information from different sources

Comprehensive Answer:""",
            input_variables=["text_context", "image_context", "entities", "reasoning_steps", 
                           "fusion_queries", "query_type", "search_weights", "question"]
        )
    
    async def process_and_store_document(self, file_path: str, file_type: str) -> Tuple[str, int, int, int, float, int, int]:
        """Enhanced document processing with graph construction"""
        start_time = datetime.now()
        
        try:
            doc_id = str(uuid.uuid4())
            
            # Process document with enhanced processor
            text_docs, image_docs, entities = await self.document_processor.process_document(file_path, file_type)
            
            if not text_docs and not image_docs:
                raise ValueError("No content extracted from document")
            
            # Store text documents in Weaviate
            if text_docs:
                if self.weaviate_manager.vector_store is None:
                    embedding_wrapper = HuggingFaceEmbeddingsWrapper(self.embedding_manager)
                    self.weaviate_manager.create_vector_store(text_docs, embedding_wrapper)
                else:
                    self.weaviate_manager.add_documents_to_store(text_docs)
                
                self.documents.extend(text_docs)
            
            # Store image documents
            if image_docs:
                self.weaviate_manager.add_image_documents(image_docs)
            
            # Store entities and build knowledge graph
            nodes_created = 0
            relationships_created = 0
            if entities:
                self.all_entities.extend(entities)
                nodes_created, relationships_created = self.graph_manager.add_entities_and_relationships(
                    entities, doc_id
                )
                
            # Update community detection
            communities = self.graph_manager.detect_communities()
            
            processing_time = (datetime.now() - start_time).total_seconds()
            logger.info(f"Document {doc_id} processed successfully in {processing_time:.2f}s")
            logger.info(f"Graph: {nodes_created} nodes, {relationships_created} relationships, {len(communities)} communities")
            
            return doc_id, len(text_docs), len(image_docs), len(entities), processing_time, nodes_created, relationships_created
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    async def enhanced_query_with_all_features(self, query: str, top_k: int = 10, 
                                             use_multimodal: bool = True, use_rag_fusion: bool = True,
                                             use_hybrid_search: bool = True, enable_multi_hop: bool = True, 
                                             include_recommendations: bool = True, fusion_queries: int = 3,
                                             dense_weight: float = 0.7, sparse_weight: float = 0.3) -> Dict[str, Any]:
        """Ultimate enhanced query processing with all advanced features"""
        start_time = datetime.now()
        
        try:
            if not self.documents and not self.weaviate_manager.image_documents:
                raise ValueError("No documents have been processed yet")
            
            # Initialize response structure
            response = {
                'fusion_queries': [query],
                'hybrid_scores': {},
                'graph_communities': []
            }
            
            # Step 1: RAG Fusion Query Processing
            fusion_results = {}
            if use_rag_fusion:
                fusion_results = await self.rag_fusion_pipeline.process_fusion_query(
                    query, fusion_queries, top_k
                )
                response['fusion_queries'] = fusion_results.get('fusion_queries', [query])
            
            # Step 2: Hybrid Search (if not using RAG Fusion or as fallback)
            text_docs = []
            hybrid_scores_data = HybridScores()
            
            if use_hybrid_search and fusion_results.get('documents'):
                text_docs = fusion_results['documents']
                hybrid_scores_data = HybridScores(
                    fusion_used=True,
                    final_count=len(text_docs)
                )
            elif use_hybrid_search:
                hybrid_results = self.hybrid_retriever.retrieve_hybrid(
                    query, top_k, dense_weight, sparse_weight
                )
                text_docs = hybrid_results['documents']
                weights_used = hybrid_results.get('weights_used', {})
                scores = hybrid_results.get('scores', {})
                
                hybrid_scores_data = HybridScores(
                    query_type=hybrid_results.get('query_type'),
                    dense_weight=weights_used.get('dense'),
                    sparse_weight=weights_used.get('sparse'),
                    dense_count=scores.get('dense_count'),
                    sparse_count=scores.get('sparse_count'),
                    fused_count=scores.get('fused_count'),
                    final_count=scores.get('final_count')
                )
            else:
                # Fallback to simple dense retrieval
                text_docs = self.weaviate_manager.search_similar(query, k=top_k)
                hybrid_scores_data = HybridScores(
                    fallback_dense=True,
                    final_count=len(text_docs)
                )
            
            # Step 3: Multimodal Image Search
            image_docs = []
            if use_multimodal and self.weaviate_manager.image_documents:
                image_docs = self.weaviate_manager.search_images(query, self.clip_manager, k=3)
            
            # Step 4: Graph-Enhanced Entity Processing
            query_entities = self.entity_extractor.extract_entities(query)
            
            # Get related entities from knowledge graph
            graph_entities = []
            for entity in query_entities:
                related = self.graph_manager.get_related_entities(entity['text'], max_results=5)
                graph_entities.extend(related)
            
            # Get graph communities
            communities = self.graph_manager.detect_communities()
            graph_communities_data = []
            for community in communities[:5]:  # Top 5 communities
                graph_communities_data.append(GraphCommunity(
                    id=community.get('id', ''),
                    algorithm=community.get('algorithm', ''),
                    nodes=community.get('nodes', []),
                    size=community.get('size', 0),
                    entities=community.get('entities', [])
                ))
            
            # Step 5: Multi-hop Reasoning with Graph Enhancement
            reasoning_steps = []
            if enable_multi_hop and self.llm:
                self.multi_hop_reasoner.entity_extractor = self.entity_extractor
                reasoning_results = self.multi_hop_reasoner.multi_hop_reason(query)
                reasoning_steps = [f"Step {i+1}: {step['answer']}" for i, step in enumerate(reasoning_results)]
            
            # Step 6: Prepare Contexts
            text_context = ""
            if text_docs:
                text_context = "\n\n".join([
                    f"Source {i+1}: {doc.page_content}"
                    for i, doc in enumerate(text_docs)
                ])
            
            image_context = ""
            if image_docs:
                image_descriptions = []
                for i, img_doc in enumerate(image_docs):
                    img_desc = f"Image {i+1} (Page {img_doc['page']}): Visual content from document"
                    image_descriptions.append(img_desc)
                image_context = "\n".join(image_descriptions)
            
            # Combine original and graph entities
            all_entities = query_entities + graph_entities
            entities_context = ", ".join([f"{ent['text']} ({ent['label']})" for ent in all_entities[:15]])
            reasoning_context = "\n".join(reasoning_steps)
            
            # Step 7: Generate Comprehensive Answer
            answer = ""
            if self.llm:
                try:
                    # Create multimodal message
                    content = []
                    
                    # Add text prompt
                    prompt_text = self.multimodal_qa_prompt.format(
                        text_context=text_context,
                        image_context=image_context,
                        entities=entities_context,
                        reasoning_steps=reasoning_context,
                        fusion_queries=", ".join(response['fusion_queries']),
                        query_type=response['hybrid_scores'].get('query_type', 'general'),
                        search_weights=str(response['hybrid_scores'].get('weights_used', {})),
                        question=query
                    )
                    
                    content.append({
                        "type": "text",
                        "text": prompt_text
                    })
                    
                    # Add relevant images
                    for img_doc in image_docs[:2]:  # Limit to 2 images
                        if img_doc.get('base64'):
                            content.append({
                                "type": "text",
                                "text": f"\n[Image from page {img_doc['page']}]:\n"
                            })
                            content.append({
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_doc['base64']}"
                                }
                            })
                    
                    # Generate response
                    message = HumanMessage(content=content)
                    llm_response = self.llm.invoke([message])
                    answer = llm_response.content
                    
                except Exception as e:
                    logger.error(f"LLM invocation failed: {e}")
                    answer = f"Found relevant information but failed to generate response: {str(e)}"
            else:
                answer = "LLM not available for answer generation"
            
            # Step 8: Generate Summary and Conclusion
            summary = answer[:200] + "..." if len(answer) > 200 else answer
            conclusion = "Comprehensive analysis completed using advanced RAG fusion, hybrid search, and knowledge graph integration."
            
            # Step 9: Generate Recommendations
            recommendations = []
            if include_recommendations:
                recommendations = self.recommendation_engine.generate_recommendations(
                    query, answer, all_entities, text_docs + image_docs
                )
            
            # Step 10: Prepare Sources
            sources = []
            for i, doc in enumerate(text_docs):
                sources.append({
                    "id": i,
                    "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    "metadata": doc.metadata,
                    "type": "text"
                })
            
            image_sources = []
            for i, img_doc in enumerate(image_docs):
                image_sources.append({
                    "id": i,
                    "image_id": img_doc['image_id'],
                    "page": img_doc['page'],
                    "metadata": img_doc['metadata'],
                    "type": "image"
                })
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            # Final response
            final_response = {
                "answer": answer,
                "summary": summary,
                "conclusion": conclusion,
                "sources": sources,
                "entities": [{"text": ent['text'], "label": ent['label'], "confidence": ent['confidence']} for ent in all_entities[:20]],
                "recommendations": recommendations,
                "confidence_score": 0.9,  # Higher confidence with all features
                "processing_time": processing_time,
                "multi_hop_reasoning": reasoning_steps,
                "image_sources": image_sources,
                "fusion_queries": response['fusion_queries'],
                "hybrid_scores": hybrid_scores_data,
                "graph_communities": graph_communities_data
            }
            
            return final_response
            
        except Exception as e:
            logger.error(f"Error processing enhanced query with all features: {e}")
            raise
    
    def close(self):
        """Clean up resources"""
        try:
            if self.weaviate_manager:
                self.weaviate_manager.close()
            if self.graph_manager:
                self.graph_manager.close()
        except Exception as e:
            logger.warning(f"Error closing RAG Bot resources: {e}")
        
        logger.info("Enhanced RAG Bot resources closed")

# LangChain compatible embeddings wrapper
class HuggingFaceEmbeddingsWrapper:
    """Wrapper to make HuggingFace embeddings compatible with LangChain"""
    
    def __init__(self, embedding_manager: HuggingFaceEmbeddingManager):
        self.embedding_manager = embedding_manager
    
    def embed_query(self, text: str) -> List[float]:
        """Embed query text"""
        return self.embedding_manager.embed_query(text)
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed multiple documents"""
        return self.embedding_manager.embed_documents(texts)

# Global variable
enhanced_rag_bot = None

# Lifespan event handler
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    global enhanced_rag_bot
    logger.info("🚀 Enhanced RAG Bot with RAG Fusion, Hybrid Search & GraphRAG starting up...")
    
    try:
        enhanced_rag_bot = EnhancedRAGBot()
        logger.info("✅ Enhanced RAG Bot with all advanced features initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize Enhanced RAG Bot: {e}")
        enhanced_rag_bot = None
    
    yield
    
    # Shutdown
    logger.info("🛑 Enhanced RAG Bot API shutting down...")
    if enhanced_rag_bot:
        try:
            enhanced_rag_bot.close()
            logger.info("✅ Enhanced RAG Bot closed successfully")
        except Exception as e:
            logger.warning(f"Error during shutdown: {e}")

# FastAPI application
app = FastAPI(
    title="Enhanced Multimodal RAG Bot with RAG Fusion, Hybrid Search & GraphRAG", 
    version="4.0.0",
    description="Advanced multimodal RAG system with RAG Fusion, Hybrid Search (Dense+Sparse+RRF), GraphRAG with community detection, CLIP, multi-hop reasoning, and comprehensive features",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=config.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer(auto_error=False)

def verify_token(credentials: HTTPAuthorizationCredentials = Security(security)):
    return credentials

# API Endpoints
@app.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Upload and process a document with all advanced features"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in config.allowed_file_types:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_extension} not supported. Allowed types: {config.allowed_file_types}"
            )
        
        # Save uploaded file
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        file_path = upload_dir / f"{uuid.uuid4()}_{file.filename}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            if len(content) > config.max_file_size:
                raise HTTPException(status_code=400, detail="File too large")
            await f.write(content)
        
        # Process document with graph construction
        doc_id, chunks_created, images_processed, entities_extracted, processing_time, nodes_created, relationships_created = await enhanced_rag_bot.process_and_store_document(
            str(file_path), file_extension
        )
        
        # Clean up uploaded file
        file_path.unlink()
        
        return DocumentUploadResponse(
            message="Document processed successfully with RAG Fusion, Hybrid Search & GraphRAG features",
            document_id=doc_id,
            chunks_created=chunks_created,
            images_processed=images_processed,
            entities_extracted=entities_extracted,
            processing_time=processing_time,
            graph_nodes_created=nodes_created,
            graph_relationships_created=relationships_created
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}")
        if 'file_path' in locals() and file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query", response_model=QueryResponse)
async def query_documents(
    request: QueryRequest,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Enhanced query with RAG Fusion, Hybrid Search, and GraphRAG"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        result = await enhanced_rag_bot.enhanced_query_with_all_features(
            query=request.query,
            top_k=request.top_k,
            use_multimodal=request.use_multimodal,
            use_rag_fusion=request.use_rag_fusion,
            use_hybrid_search=request.use_hybrid_search,
            enable_multi_hop=request.enable_multi_hop,
            include_recommendations=request.include_recommendations,
            fusion_queries=request.fusion_queries,
            dense_weight=request.dense_weight,
            sparse_weight=request.sparse_weight
        )
        
        return QueryResponse(**result)
        
    except Exception as e:
        logger.error(f"Query error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    """Comprehensive health check endpoint"""
    status_info = {
        "status": "healthy" if enhanced_rag_bot else "degraded", 
        "timestamp": datetime.now().isoformat(),
        "components": {}
    }
    
    if enhanced_rag_bot:
        status_info["components"] = {
            "weaviate_connected": enhanced_rag_bot.weaviate_manager.client is not None,
            "weaviate_connection_type": enhanced_rag_bot.weaviate_manager.connection_type,
            "clip_initialized": enhanced_rag_bot.clip_manager.clip_model is not None,
            "docling_available": enhanced_rag_bot.document_processor.docling_processor.converter is not None,
            "spacy_available": enhanced_rag_bot.entity_extractor.nlp is not None,
            "neo4j_connected": enhanced_rag_bot.graph_manager.driver is not None,
            "bm25_available": BM25_AVAILABLE,
            "reranker_available": enhanced_rag_bot.reranker.reranker is not None,
            "llm_available": enhanced_rag_bot.llm is not None,
            "rag_fusion_available": True,
            "hybrid_search_available": True,
            "graph_rag_available": NETWORKX_AVAILABLE
        }
    
    status_info.update({
        "enhanced_rag_bot_initialized": enhanced_rag_bot is not None,
        "libraries_available": {
            "huggingface_embeddings": HUGGINGFACE_AVAILABLE,
            "sentence_transformers": SENTENCE_TRANSFORMERS_AVAILABLE,
            "llamaindex": LLAMAINDEX_AVAILABLE,
            "docling": DOCLING_AVAILABLE,
            "networkx": NETWORKX_AVAILABLE,
            "neo4j": NEO4J_AVAILABLE,
            "bm25": BM25_AVAILABLE,
            "sklearn": SKLEARN_AVAILABLE
        }
    })
    
    return status_info

@app.get("/stats")
async def get_comprehensive_stats(credentials: HTTPAuthorizationCredentials = Depends(verify_token)):
    """Get comprehensive system statistics"""
    if enhanced_rag_bot is None:
        return {
            "error": "Enhanced RAG Bot not initialized",
            "total_documents": 0,
            "total_images": 0,
            "total_entities": 0,
            "graph_nodes": 0,
            "graph_relationships": 0
        }
    
    graph_stats = {}
    if enhanced_rag_bot.graph_manager.graph.nodes:
        graph_stats = {
            "nodes": len(enhanced_rag_bot.graph_manager.graph.nodes),
            "edges": len(enhanced_rag_bot.graph_manager.graph.edges),
            "communities": len(enhanced_rag_bot.graph_manager.communities),
            "density": nx.density(enhanced_rag_bot.graph_manager.graph) if NETWORKX_AVAILABLE else 0
        }
    
    return {
        "total_text_documents": len(enhanced_rag_bot.documents),
        "total_image_documents": len(enhanced_rag_bot.weaviate_manager.image_documents),
        "total_entities": len(enhanced_rag_bot.all_entities),
        "vector_store_initialized": enhanced_rag_bot.weaviate_manager.vector_store is not None,
        "weaviate_connection": enhanced_rag_bot.weaviate_manager.connection_type,
        "knowledge_graph": graph_stats,
        "models": {
            "clip_model": enhanced_rag_bot.clip_manager.model_name,
            "embedding_model": enhanced_rag_bot.embedding_manager.model_name,
            "reranker_model": enhanced_rag_bot.reranker.model_name if enhanced_rag_bot.reranker.reranker else None,
            "huggingface_device": enhanced_rag_bot.embedding_manager.device
        },
        "advanced_capabilities": {
            "rag_fusion": True,
            "hybrid_search": {
                "dense_retrieval": True,
                "sparse_retrieval": BM25_AVAILABLE,
                "reciprocal_rank_fusion": True,
                "cross_encoder_reranking": enhanced_rag_bot.reranker.reranker is not None
            },
            "graph_rag": {
                "knowledge_graph": True,
                "community_detection": NETWORKX_AVAILABLE,
                "multi_hop_reasoning": True,
                "subgraph_extraction": NETWORKX_AVAILABLE
            },
            "multimodal": {
                "text_processing": True,
                "image_processing": enhanced_rag_bot.clip_manager.clip_model is not None,
                "entity_extraction": enhanced_rag_bot.entity_extractor.nlp is not None,
                "semantic_chunking": enhanced_rag_bot.document_processor.semantic_chunker is not None
            }
        }
    }

@app.post("/test/rag-fusion")
async def test_rag_fusion(
    query: str,
    num_queries: int = 3,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Test RAG Fusion query expansion and parallel retrieval"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        start_time = datetime.now()
        
        # Test query expansion
        expanded_queries = enhanced_rag_bot.query_expander.expand_query(query, num_queries)
        
        # Test fusion pipeline
        fusion_results = await enhanced_rag_bot.rag_fusion_pipeline.process_fusion_query(
            query, num_queries, top_k=5
        )
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "original_query": query,
            "expanded_queries": expanded_queries,
            "fusion_results": {
                "documents_found": fusion_results.get('final_count', 0),
                "retrieval_results": fusion_results.get('retrieval_results_count', 0),
                "processing_time": fusion_results.get('processing_time', 0)
            },
            "total_processing_time": processing_time,
            "rag_fusion_available": True
        }
        
    except Exception as e:
        logger.error(f"RAG Fusion test error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/test/hybrid-search")
async def test_hybrid_search(
    query: str,
    top_k: int = 10,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Test hybrid search combining dense and sparse retrieval"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        start_time = datetime.now()
        
        # Test hybrid retrieval
        hybrid_results = enhanced_rag_bot.hybrid_retriever.retrieve_hybrid(query, top_k)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "query": query,
            "query_type": hybrid_results.get('query_type', 'unknown'),
            "weights_used": hybrid_results.get('weights_used', {}),
            "retrieval_scores": hybrid_results.get('scores', {}),
            "documents_found": len(hybrid_results.get('documents', [])),
            "processing_time": processing_time,
            "components_available": {
                "dense_retrieval": enhanced_rag_bot.weaviate_manager.vector_store is not None,
                "sparse_retrieval": BM25_AVAILABLE,
                "reranking": enhanced_rag_bot.reranker.reranker is not None,
                "reciprocal_rank_fusion": True
            }
        }
        
    except Exception as e:
        logger.error(f"Hybrid search test error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/test/graph-rag")
async def test_graph_rag(
    query: str,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Test GraphRAG capabilities"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        start_time = datetime.now()
        
        # Extract entities from query
        query_entities = enhanced_rag_bot.entity_extractor.extract_entities(query)
        entity_texts = [ent['text'] for ent in query_entities]
        
        # Test subgraph extraction
        subgraph = enhanced_rag_bot.graph_manager.get_subgraph(entity_texts, max_hops=2)
        
        # Test community detection
        communities = enhanced_rag_bot.graph_manager.detect_communities()
        
        # Test related entity finding
        related_entities = []
        for entity_text in entity_texts[:3]:  # Limit to first 3 entities
            related = enhanced_rag_bot.graph_manager.get_related_entities(entity_text, max_results=5)
            related_entities.extend(related)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "query": query,
            "extracted_entities": [{"text": ent['text'], "label": ent['label']} for ent in query_entities],
            "subgraph_stats": {
                "nodes": len(subgraph.nodes) if subgraph else 0,
                "edges": len(subgraph.edges) if subgraph else 0
            },
            "communities_detected": len(communities),
            "related_entities_found": len(related_entities),
            "total_graph_stats": {
                "total_nodes": len(enhanced_rag_bot.graph_manager.graph.nodes),
                "total_edges": len(enhanced_rag_bot.graph_manager.graph.edges),
                "total_communities": len(enhanced_rag_bot.graph_manager.communities)
            },
            "processing_time": processing_time,
            "graph_rag_features": {
                "knowledge_graph_construction": True,
                "community_detection": NETWORKX_AVAILABLE,
                "subgraph_extraction": NETWORKX_AVAILABLE,
                "entity_relationships": True,
                "multi_hop_graph_reasoning": True
            }
        }
        
    except Exception as e:
        logger.error(f"GraphRAG test error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/graph/communities")
async def get_graph_communities(
    limit: int = 10,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Get detected communities from knowledge graph"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        communities = enhanced_rag_bot.graph_manager.detect_communities()
        
        # Sort by size and limit
        communities_sorted = sorted(communities, key=lambda x: x['size'], reverse=True)
        
        return {
            "total_communities": len(communities),
            "communities": communities_sorted[:limit],
            "graph_stats": {
                "total_nodes": len(enhanced_rag_bot.graph_manager.graph.nodes),
                "total_edges": len(enhanced_rag_bot.graph_manager.graph.edges)
            }
        }
        
    except Exception as e:
        logger.error(f"Error getting graph communities: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/config/advanced")
async def get_advanced_config(credentials: HTTPAuthorizationCredentials = Depends(verify_token)):
    """Get comprehensive system configuration"""
    return {
        "rag_fusion": {
            "query_expansion": True,
            "parallel_retrieval": True,
            "result_aggregation": True,
            "deduplication": SKLEARN_AVAILABLE
        },
        "hybrid_search": {
            "dense_embeddings": {
                "model": config.hf_embedding_model,
                "available": HUGGINGFACE_AVAILABLE
            },
            "sparse_retrieval": {
                "bm25_available": BM25_AVAILABLE,
                "tfidf_available": SKLEARN_AVAILABLE
            },
            "fusion_methods": {
                "reciprocal_rank_fusion": True,
                "adaptive_weighting": True,
                "query_type_classification": True
            },
            "reranking": {
                "cross_encoder_model": config.hf_reranker_model,
                "available": SENTENCE_TRANSFORMERS_AVAILABLE
            }
        },
        "graph_rag": {
            "knowledge_graph": {
                "networkx_available": NETWORKX_AVAILABLE,
                "neo4j_available": NEO4J_AVAILABLE,
                "neo4j_configured": bool(config.neo4j_uri and config.neo4j_username)
            },
            "community_detection": {
                "louvain_algorithm": NETWORKX_AVAILABLE,
                "modularity_optimization": NETWORKX_AVAILABLE
            },
            "graph_traversal": {
                "multi_hop_reasoning": NETWORKX_AVAILABLE,
                "subgraph_extraction": NETWORKX_AVAILABLE,
                "entity_relationships": True
            }
        },
        "multimodal_capabilities": {
            "clip_model": config.clip_model_name,
            "text_embeddings": config.hf_embedding_model,
            "image_processing": CLIP_AVAILABLE,
            "docling_available": DOCLING_AVAILABLE,
            "spacy_available": SPACY_AVAILABLE
        },
        "vector_databases": {
            "weaviate_available": WEAVIATE_AVAILABLE,
            "weaviate_url": config.weaviate_url
        },
        "language_models": {
            "azure_openai_configured": bool(config.azure_openai_endpoint and config.azure_openai_api_key),
            "llamaindex_available": LLAMAINDEX_AVAILABLE
        },
        "system_settings": {
            "debug": config.debug,
            "hf_device": config.hf_device,
            "persist_dir": config.persist_dir,
            "max_file_size": config.max_file_size,
            "allowed_file_types": config.allowed_file_types
        }
    }

# Error handlers
@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    from fastapi.responses import JSONResponse
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "timestamp": datetime.now().isoformat()
        }
    )

@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    from fastapi.responses import JSONResponse
    logger.error(f"Unhandled exception: {exc}")
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "status_code": 500,
            "timestamp": datetime.now().isoformat(),
            "detail": str(exc) if config.debug else "Internal server error"
        }
    )

if __name__ == "__main__":
    import sys
    
    # Suppress specific warnings
    warnings.filterwarnings("ignore", category=DeprecationWarning, module="websockets")
    warnings.filterwarnings("ignore", category=DeprecationWarning, module="uvicorn")
    warnings.filterwarnings("ignore", message=".*swigvarlink.*")
    
    # Validate required environment variables
    required_vars = [
        'AZURE_OPENAI_ENDPOINT',
        'AZURE_OPENAI_API_KEY',
        'AZURE_OPENAI_DEPLOYMENT_NAME'
    ]
    
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    if missing_vars:
        logger.error(f"Missing required environment variables: {missing_vars}")
        logger.error("Please set these variables in your .env file or environment")
        sys.exit(1)
    
    # Create required directories
    for directory in ["uploads", config.persist_dir, config.hf_cache_dir]:
        Path(directory).mkdir(parents=True, exist_ok=True)  # Added parents=True
    
    logger.info("🚀 Starting Enhanced RAG Bot with Advanced Features...")
    logger.info("📋 Feature Availability:")
    logger.info(f"  ✅ RAG Fusion: Query Expansion + Parallel Retrieval + RRF")
    logger.info(f"  ✅ Hybrid Search: Dense ({HUGGINGFACE_AVAILABLE}) + Sparse ({BM25_AVAILABLE}) + Reranking ({SENTENCE_TRANSFORMERS_AVAILABLE})")
    logger.info(f"  ✅ GraphRAG: Knowledge Graph + Community Detection ({NETWORKX_AVAILABLE}) + Neo4j ({NEO4J_AVAILABLE})")
    logger.info(f"  ✅ Multimodal: CLIP ({CLIP_AVAILABLE}) + Docling ({DOCLING_AVAILABLE}) + Spacy ({SPACY_AVAILABLE})")
    logger.info(f"  ✅ Vector Store: Weaviate ({WEAVIATE_AVAILABLE}) + LlamaIndex ({LLAMAINDEX_AVAILABLE})")
    logger.info(f"  📊 Models: Embedding ({config.hf_embedding_model}) + Reranker ({config.hf_reranker_model})")
    logger.info(f"  💾 Storage: {config.persist_dir} | Device: {config.hf_device}")
    
    try:
        if config.debug:
            # Update this line to match your actual filename
            # If you saved as main.py, change to "main:app"
            # If you saved as enhanced_multimodal_rag.py, keep as is
            uvicorn.run(
                "RAG:app",  # Change this if your filename is different
                host="0.0.0.0",
                port=8000,
                reload=True,
                log_level="info"
            )
        else:
            uvicorn.run(
                app,
                host="0.0.0.0",
                port=8000,
                log_level="info"
            )
    except KeyboardInterrupt:
        logger.info("Server stopped by user")
    except Exception as e:
        logger.error(f"Server error: {e}")
        sys.exit(1)