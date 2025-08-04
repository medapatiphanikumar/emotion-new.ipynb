# Suppress protobuf warnings at the very beginning
import warnings
import os
import sys

# Suppress all common warnings first
warnings.filterwarnings('ignore', category=UserWarning, module='google.protobuf.runtime_version')
warnings.filterwarnings('ignore', category=DeprecationWarning, module='websockets')
warnings.filterwarnings('ignore', category=DeprecationWarning, module='uvicorn')
warnings.filterwarnings('ignore', message='.*swigvarlink.*')
warnings.filterwarnings('ignore', message='.*destructor.*')
warnings.filterwarnings('ignore', category=ResourceWarning, module='neo4j')

os.environ['PROTOCOL_BUFFERS_PYTHON_IMPLEMENTATION'] = 'python'

# Check for critical dependencies first
try:
    import fastapi
    import uvicorn
    import pydantic
except ImportError as e:
    print(f"❌ Critical dependency missing: {e}")
    print("Please install required packages:")
    print("pip install fastapi uvicorn pydantic")
    sys.exit(1)

import asyncio
import logging
from typing import List, Dict, Any, Optional, Tuple, Union
from pathlib import Path
import json
import uuid
from datetime import datetime
import time
from contextlib import asynccontextmanager
import re
from collections import defaultdict, Counter
import numpy as np

# Try to import optional dependencies with fallbacks
try:
    import networkx as nx
    NETWORKX_AVAILABLE = True
except ImportError:
    print("⚠️ NetworkX not available - multi-hop reasoning will be limited")
    NETWORKX_AVAILABLE = False

try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    print("⚠️ Spacy or en_core_web_sm model not available")
    print("Install with: pip install spacy && python -m spacy download en_core_web_sm")
    SPACY_AVAILABLE = False
    nlp = None

# Core imports with error handling
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    print("⚠️ PyMuPDF not available. Install with: pip install PyMuPDF")
    PYMUPDF_AVAILABLE = False

try:
    import pymupdf
except ImportError:
    print("⚠️ pymupdf not available")

try:
    from unstructured.partition.auto import partition
    from unstructured.partition.pdf import partition_pdf
    from unstructured.partition.docx import partition_docx
    from unstructured.partition.text import partition_text
    UNSTRUCTURED_AVAILABLE = True
except ImportError:
    print("⚠️ Unstructured not available. Install with: pip install unstructured[pdf]")
    UNSTRUCTURED_AVAILABLE = False

try:
    import docx
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("⚠️ python-docx not available. Install with: pip install python-docx")
    DOCX_AVAILABLE = False

# Docling imports for advanced document processing
try:
    from docling.document_converter import DocumentConverter
    from docling.datamodel.base_models import InputFormat
    from docling.datamodel.pipeline_options import PdfPipelineOptions
    from docling.backend.pypdfium2_backend import PyPdfiumDocumentBackend
    DOCLING_AVAILABLE = True
except ImportError:
    print("⚠️ Docling not available. Install with: pip install docling")
    DOCLING_AVAILABLE = False

# FastAPI imports
from fastapi import FastAPI, File, UploadFile, HTTPException, Depends, Security
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn

# LangChain imports
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document as LangChainDocument
from langchain_core.prompts import PromptTemplate
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage

# CLIP and Vision imports
try:
    from transformers import CLIPProcessor, CLIPModel
    from PIL import Image
    import torch
    CLIP_AVAILABLE = True
except ImportError:
    print("⚠️ CLIP dependencies not available. Install with: pip install transformers torch pillow")
    CLIP_AVAILABLE = False

try:
    import base64
    import io
    BASIC_AVAILABLE = True
except ImportError:
    BASIC_AVAILABLE = False

# LlamaIndex imports
try:
    from llama_index.core import (
        VectorStoreIndex, 
        Document as LlamaDocument,
        Settings,
        StorageContext,
        SimpleDirectoryReader,
        load_index_from_storage
    )
    from llama_index.core.node_parser import SimpleNodeParser, SentenceSplitter
    from llama_index.core.extractors import TitleExtractor, QuestionsAnsweredExtractor
    from llama_index.core.ingestion import IngestionPipeline
    from llama_index.core.schema import BaseNode, TextNode
    from llama_index.embeddings.huggingface import HuggingFaceEmbedding
    from llama_index.llms.azure_openai import AzureOpenAI
    from llama_index.vector_stores.weaviate import WeaviateVectorStore as LlamaWeaviateVectorStore
    LLAMAINDEX_AVAILABLE = True
except ImportError as e:
    print(f"⚠️ LlamaIndex not available: {e}")
    LLAMAINDEX_AVAILABLE = False
    # Define dummy classes to prevent NameError
    class BaseNode:
        pass
    class TextNode:
        pass
    class LlamaDocument:
        pass
    class HuggingFaceEmbedding:
        pass
    class AzureOpenAI:
        pass

# HuggingFace imports
try:
    from sentence_transformers import SentenceTransformer
    from transformers import AutoTokenizer, AutoModel
    HUGGINGFACE_AVAILABLE = True
except ImportError:
    print("⚠️ HuggingFace transformers not available. Please install: pip install transformers sentence-transformers")
    HUGGINGFACE_AVAILABLE = False

# Weaviate imports
try:
    import weaviate
    from langchain_weaviate.vectorstores import WeaviateVectorStore
    WEAVIATE_AVAILABLE = True
except ImportError:
    print("⚠️ Weaviate not available. Using file-based storage.")
    WEAVIATE_AVAILABLE = False

# Vector database and ML imports
import numpy as np
try:
    from sentence_transformers import SentenceTransformer, CrossEncoder
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    print("⚠️ SentenceTransformers not available. Reranking disabled.")
    SENTENCE_TRANSFORMERS_AVAILABLE = False

try:
    from rank_bm25 import BM25Okapi
    BM25_AVAILABLE = True
except ImportError:
    print("⚠️ BM25 not available. Using simple retrieval.")
    BM25_AVAILABLE = False

# Neo4j imports
try:
    from neo4j import GraphDatabase
    NEO4J_AVAILABLE = True
except ImportError:
    print("⚠️ Neo4j not available. Using fallback knowledge tracking.")
    NEO4J_AVAILABLE = False

# Spacy for NER
try:
    import spacy
    # Try to load the model
    nlp = spacy.load("en_core_web_sm")
    SPACY_AVAILABLE = True
except (ImportError, OSError):
    print("⚠️ Spacy or en_core_web_sm model not available. Please install: python -m spacy download en_core_web_sm")
    SPACY_AVAILABLE = False
    nlp = None

# Environment and utilities
try:
    from dotenv import load_dotenv
    DOTENV_AVAILABLE = True
except ImportError:
    print("⚠️ python-dotenv not available. Install with: pip install python-dotenv")
    DOTENV_AVAILABLE = False

try:
    import aiofiles
    AIOFILES_AVAILABLE = True
except ImportError:
    print("⚠️ aiofiles not available. Install with: pip install aiofiles")
    AIOFILES_AVAILABLE = False

try:
    from sklearn.metrics.pairwise import cosine_similarity
    SKLEARN_AVAILABLE = True
except ImportError:
    print("⚠️ scikit-learn not available. Install with: pip install scikit-learn")
    SKLEARN_AVAILABLE = False

# Load environment variables if available
if DOTENV_AVAILABLE:
    load_dotenv()
else:
    print("⚠️ Environment variables will be loaded from system only")

# Configure logging
logging.basicConfig(
    level=getattr(logging, os.getenv('LOG_LEVEL', 'INFO')),
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Pydantic models
class QueryRequest(BaseModel):
    query: str
    top_k: int = 5
    use_graph: bool = True
    use_reranking: bool = True
    use_multimodal: bool = True
    summary_length: str = "medium"
    enable_multi_hop: bool = True
    include_recommendations: bool = True

class QueryResponse(BaseModel):
    answer: str
    summary: str
    conclusion: str
    sources: List[Dict[str, Any]]
    entities: List[Dict[str, Any]]
    recommendations: List[str]
    confidence_score: float
    processing_time: float
    multi_hop_reasoning: List[str]
    image_sources: List[Dict[str, Any]]

class DocumentUploadResponse(BaseModel):
    message: str
    document_id: str
    chunks_created: int
    images_processed: int
    entities_extracted: int
    processing_time: float

# Configuration class
class Config:
    def __init__(self):
        # Azure OpenAI Configuration
        self.azure_openai_endpoint = os.getenv('AZURE_OPENAI_ENDPOINT')
        self.azure_openai_api_key = os.getenv('AZURE_OPENAI_API_KEY')
        self.azure_openai_api_version = os.getenv('AZURE_OPENAI_API_VERSION', '2025-01-01-preview')
        self.azure_openai_deployment_name = os.getenv('AZURE_OPENAI_DEPLOYMENT_NAME')
        self.azure_openai_model_name = os.getenv('AZURE_OPENAI_MODEL_NAME')
        self.azure_openai_temperature = float(os.getenv('AZURE_OPENAI_TEMPERATURE', '0.1'))
        self.azure_openai_max_tokens = int(os.getenv('AZURE_OPENAI_MAX_TOKENS', '4000'))
        
        # HuggingFace Configuration
        self.hf_embedding_model = os.getenv('HF_EMBEDDING_MODEL', 'BAAI/bge-small-en-v1.5')
        self.hf_cache_dir = os.getenv('HF_CACHE_DIR', './hf_cache')
        self.hf_device = os.getenv('HF_DEVICE', 'cpu')
        
        # CLIP Configuration
        self.clip_model_name = os.getenv('CLIP_MODEL_NAME', 'openai/clip-vit-base-patch32')
        
        # Neo4j Configuration
        self.neo4j_uri = os.getenv('NEO4J_URI')
        self.neo4j_username = os.getenv('NEO4J_USERNAME')
        self.neo4j_password = os.getenv('NEO4J_PASSWORD')
        self.neo4j_database = os.getenv('NEO4J_DATABASE')
        
        # Weaviate Configuration
        self.weaviate_url = os.getenv('WEAVIATE_URL', 'http://localhost:8080')
        self.weaviate_api_key = os.getenv('WEAVIATE_API_KEY')
        
        # Storage Configuration
        self.persist_dir = os.getenv('PERSIST_DIR', './storage')
        self.index_name = os.getenv('INDEX_NAME', 'rag_documents')
        
        # Security Configuration
        self.api_key_header = os.getenv('API_KEY_HEADER', 'X-API-Key')
        self.cors_origins = json.loads(os.getenv('CORS_ORIGINS', '["*"]'))
        self.debug = os.getenv('DEBUG', 'true').lower() == 'true'
        self.max_file_size = int(os.getenv('MAX_FILE_SIZE', '104857600'))
        self.allowed_file_types = json.loads(os.getenv('ALLOWED_FILE_TYPES', '["csv", "json", "xlsx", "xls", "pdf", "txt", "docx", "doc"]'))
        self.batch_size = int(os.getenv('BATCH_SIZE', '100'))
        self.max_retry_attempts = int(os.getenv('MAX_RETRY_ATTEMPTS', '3'))
        self.processing_timeout = int(os.getenv('PROCESSING_TIMEOUT', '300'))

config = Config()

# Enhanced Entity Extractor
class EntityExtractor:
    """Advanced entity extraction with custom rules and NER"""
    
    def __init__(self):
        self.nlp = nlp if SPACY_AVAILABLE else None
        self.custom_patterns = [
            # Financial patterns
            (r'\$[\d,]+(?:\.\d{2})?', 'MONEY'),
            (r'\b\d{1,2}/\d{1,2}/\d{2,4}\b', 'DATE'),
            (r'\b\d+%\b', 'PERCENTAGE'),
            # Technical patterns
            (r'\b[A-Z]{2,10}\b', 'ACRONYM'),
            (r'\b\w+@\w+\.\w+\b', 'EMAIL'),
            (r'\bhttps?://[^\s]+\b', 'URL'),
        ]
    
    def extract_entities(self, text: str) -> List[Dict[str, Any]]:
        """Extract entities from text using multiple approaches"""
        entities = []
        
        # Spacy NER
        if self.nlp:
            try:
                doc = self.nlp(text)
                for ent in doc.ents:
                    entities.append({
                        'text': ent.text,
                        'label': ent.label_,
                        'start': ent.start_char,
                        'end': ent.end_char,
                        'confidence': 1.0,
                        'method': 'spacy'
                    })
            except Exception as e:
                logger.warning(f"Spacy NER failed: {e}")
        
        # Custom regex patterns
        for pattern, label in self.custom_patterns:
            try:
                matches = re.finditer(pattern, text, re.IGNORECASE)
                for match in matches:
                    entities.append({
                        'text': match.group(),
                        'label': label,
                        'start': match.start(),
                        'end': match.end(),
                        'confidence': 0.8,
                        'method': 'regex'
                    })
            except Exception as e:
                logger.warning(f"Regex pattern {pattern} failed: {e}")
        
        # Remove duplicates and sort by confidence
        unique_entities = []
        seen = set()
        for entity in sorted(entities, key=lambda x: x['confidence'], reverse=True):
            key = (entity['text'].lower(), entity['label'])
            if key not in seen:
                seen.add(key)
                unique_entities.append(entity)
        
        return unique_entities[:20]  # Limit to top 20 entities

# Semantic Chunker using CLIP embeddings
class SemanticChunker:
    """Advanced semantic chunking based on content similarity"""
    
    def __init__(self, clip_model, clip_processor, max_chunk_size=1000, similarity_threshold=0.7):
        self.clip_model = clip_model
        self.clip_processor = clip_processor
        self.max_chunk_size = max_chunk_size
        self.similarity_threshold = similarity_threshold
        self.basic_splitter = RecursiveCharacterTextSplitter(
            chunk_size=200,
            chunk_overlap=50,
            separators=["\n\n", "\n", ". ", " "]
        )
    
    def embed_text(self, text: str) -> np.ndarray:
        """Create CLIP embedding for text"""
        try:
            inputs = self.clip_processor(
                text=text,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=77
            )
            
            with torch.no_grad():
                features = self.clip_model.get_text_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
            
            return features.squeeze().numpy()
        except Exception as e:
            logger.warning(f"CLIP text embedding failed: {e}")
            return np.zeros(512)  # Default embedding size
    
    def semantic_chunk(self, text: str) -> List[str]:
        """Create semantically coherent chunks"""
        # First, create basic chunks
        basic_chunks = self.basic_splitter.split_text(text)
        
        if len(basic_chunks) <= 1:
            return basic_chunks
        
        # Create embeddings for each chunk
        chunk_embeddings = []
        for chunk in basic_chunks:
            embedding = self.embed_text(chunk)
            chunk_embeddings.append(embedding)
        
        # Group similar chunks together
        semantic_chunks = []
        current_chunk = [basic_chunks[0]]
        current_size = len(basic_chunks[0])
        
        for i in range(1, len(basic_chunks)):
            # Calculate similarity with current chunk
            if len(chunk_embeddings[i-1]) > 0 and len(chunk_embeddings[i]) > 0:
                if SKLEARN_AVAILABLE:
                    similarity = cosine_similarity(
                        chunk_embeddings[i-1].reshape(1, -1),
                        chunk_embeddings[i].reshape(1, -1)
                    )[0][0]
                else:
                    # Fallback similarity calculation
                    dot_product = np.dot(chunk_embeddings[i-1], chunk_embeddings[i])
                    norm_a = np.linalg.norm(chunk_embeddings[i-1])
                    norm_b = np.linalg.norm(chunk_embeddings[i])
                    similarity = dot_product / (norm_a * norm_b) if norm_a > 0 and norm_b > 0 else 0.5
            else:
                similarity = 0.5
            
            # Check if we should merge chunks
            new_size = current_size + len(basic_chunks[i])
            
            if (similarity > self.similarity_threshold and 
                new_size <= self.max_chunk_size):
                current_chunk.append(basic_chunks[i])
                current_size = new_size
            else:
                # Finalize current chunk and start new one
                semantic_chunks.append(" ".join(current_chunk))
                current_chunk = [basic_chunks[i]]
                current_size = len(basic_chunks[i])
        
        # Add the last chunk
        if current_chunk:
            semantic_chunks.append(" ".join(current_chunk))
        
        return semantic_chunks

# Multi-hop Reasoning Engine
class MultiHopReasoner:
    """Implements multi-hop reasoning for complex queries"""
    
    def __init__(self, retriever, llm):
        self.retriever = retriever
        self.llm = llm
        self.reasoning_graph = nx.DiGraph()
    
    def decompose_query(self, query: str) -> List[str]:
        """Decompose complex query into sub-questions"""
        if not self.llm:
            return [query]
        
        try:
            decomposition_prompt = f"""
            Break down this complex question into 2-3 simpler sub-questions that need to be answered step by step:
            
            Original question: {query}
            
            Sub-questions (one per line):
            """
            
            messages = [
                SystemMessage(content="You are an expert at breaking down complex questions into simpler sub-questions."),
                HumanMessage(content=decomposition_prompt)
            ]
            
            response = self.llm.invoke(messages)
            sub_questions = [q.strip() for q in response.content.split('\n') if q.strip() and not q.strip().startswith('Sub-questions')]
            
            return sub_questions[:3] if sub_questions else [query]
            
        except Exception as e:
            logger.warning(f"Query decomposition failed: {e}")
            return [query]
    
    def reason_step(self, question: str, context: str = "") -> Dict[str, Any]:
        """Perform one reasoning step"""
        # Retrieve relevant documents
        try:
            if hasattr(self.retriever, 'retrieve_hybrid'):
                results = self.retriever.retrieve_hybrid(question, top_k=3)
                retrieved_docs = results.get('langchain_docs', [])
                additional_context = results.get('llamaindex_response', '')
            else:
                retrieved_docs = []
                additional_context = ""
            
            # Prepare context
            doc_context = "\n".join([doc.page_content for doc in retrieved_docs[:3]])
            full_context = f"{context}\n{doc_context}\n{additional_context}".strip()
            
            # Generate answer
            if self.llm and full_context:
                reasoning_prompt = f"""
                Based on the following context, answer this question step by step:
                
                Context: {full_context}
                Question: {question}
                
                Provide a clear, factual answer with reasoning:
                """
                
                messages = [
                    SystemMessage(content="You are an expert analyst who provides step-by-step reasoning."),
                    HumanMessage(content=reasoning_prompt)
                ]
                
                response = self.llm.invoke(messages)
                answer = response.content
            else:
                answer = f"Insufficient context to answer: {question}"
            
            return {
                'question': question,
                'answer': answer,
                'context': full_context,
                'sources': retrieved_docs
            }
            
        except Exception as e:
            logger.error(f"Reasoning step failed: {e}")
            return {
                'question': question,
                'answer': f"Error in reasoning: {str(e)}",
                'context': context,
                'sources': []
            }
    
    def multi_hop_reason(self, query: str) -> List[Dict[str, Any]]:
        """Perform multi-hop reasoning"""
        # Decompose query
        sub_questions = self.decompose_query(query)
        reasoning_steps = []
        accumulated_context = ""
        
        for i, question in enumerate(sub_questions):
            step_result = self.reason_step(question, accumulated_context)
            reasoning_steps.append(step_result)
            
            # Accumulate context for next step
            accumulated_context += f"\n\nStep {i+1}: {step_result['answer']}"
        
        return reasoning_steps

# Recommendation Engine
class RecommendationEngine:
    """Generate recommendations based on query and context"""
    
    def __init__(self, llm, entity_extractor):
        self.llm = llm
        self.entity_extractor = entity_extractor
    
    def generate_recommendations(self, query: str, answer: str, entities: List[Dict], sources: List) -> List[str]:
        """Generate contextual recommendations"""
        if not self.llm:
            return ["Enable LLM to get personalized recommendations"]
        
        try:
            # Extract key topics from entities
            key_topics = [ent['text'] for ent in entities if ent['label'] in ['ORG', 'PRODUCT', 'TECH']][:5]
            
            # Create recommendation prompt
            recommendation_prompt = f"""
            Based on the user's query and the information provided, suggest 3-5 practical next steps or related topics they might want to explore:
            
            User Query: {query}
            Answer Context: {answer[:500]}...
            Key Topics: {', '.join(key_topics)}
            
            Provide actionable recommendations (one per line):
            """
            
            messages = [
                SystemMessage(content="You are an expert advisor who provides practical, actionable recommendations."),
                HumanMessage(content=recommendation_prompt)
            ]
            
            response = self.llm.invoke(messages)
            recommendations = [rec.strip() for rec in response.content.split('\n') if rec.strip() and not rec.strip().startswith('Provide')]
            
            return recommendations[:5]
            
        except Exception as e:
            logger.warning(f"Recommendation generation failed: {e}")
            return ["Explore related documents", "Search for more specific information", "Consider different perspectives on this topic"]

# CLIP Multimodal Manager
class CLIPMultimodalManager:
    """Manages CLIP model for unified text and image embeddings"""
    
    def __init__(self, model_name: str = None):
        self.model_name = model_name or config.clip_model_name
        self.clip_model = None
        self.clip_processor = None
        self.image_data_store = {}
        
        self.setup_clip()
    
    def setup_clip(self):
        """Initialize CLIP model and processor"""
        try:
            self.clip_model = CLIPModel.from_pretrained(self.model_name)
            self.clip_processor = CLIPProcessor.from_pretrained(self.model_name)
            self.clip_model.eval()
            logger.info(f"✅ CLIP model initialized: {self.model_name}")
        except Exception as e:
            logger.error(f"Failed to initialize CLIP: {e}")
            self.clip_model = None
            self.clip_processor = None
    
    def embed_image(self, image_data) -> np.ndarray:
        """Embed image using CLIP"""
        if not self.clip_model:
            return np.zeros(512)
        
        try:
            if isinstance(image_data, str):  # If path
                image = Image.open(image_data).convert("RGB")
            else:  # If PIL Image
                image = image_data
            
            inputs = self.clip_processor(images=image, return_tensors="pt")
            
            with torch.no_grad():
                features = self.clip_model.get_image_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
            
            return features.squeeze().numpy()
        except Exception as e:
            logger.warning(f"Image embedding failed: {e}")
            return np.zeros(512)
    
    def embed_text(self, text: str) -> np.ndarray:
        """Embed text using CLIP"""
        if not self.clip_model:
            return np.zeros(512)
        
        try:
            inputs = self.clip_processor(
                text=text,
                return_tensors="pt",
                padding=True,
                truncation=True,
                max_length=77
            )
            
            with torch.no_grad():
                features = self.clip_model.get_text_features(**inputs)
                features = features / features.norm(dim=-1, keepdim=True)
            
            return features.squeeze().numpy()
        except Exception as e:
            logger.warning(f"Text embedding failed: {e}")
            return np.zeros(512)
    
    def store_image(self, image_id: str, pil_image: Image.Image) -> str:
        """Store image as base64 for GPT-4V"""
        try:
            buffered = io.BytesIO()
            pil_image.save(buffered, format="PNG")
            img_base64 = base64.b64encode(buffered.getvalue()).decode()
            self.image_data_store[image_id] = img_base64
            return img_base64
        except Exception as e:
            logger.error(f"Failed to store image {image_id}: {e}")
            return ""
    
    def get_stored_image(self, image_id: str) -> str:
        """Retrieve stored image base64"""
        return self.image_data_store.get(image_id, "")

# Advanced Docling Document Processor
class DoclingProcessor:
    """Advanced document processing using Docling"""
    
    def __init__(self):
        self.converter = None
        if DOCLING_AVAILABLE:
            try:
                # Configure Docling with advanced options
                pipeline_options = PdfPipelineOptions()
                pipeline_options.do_ocr = True
                pipeline_options.do_table_structure = True
                pipeline_options.table_structure_options.do_cell_matching = True
                
                self.converter = DocumentConverter(
                    format_options={
                        InputFormat.PDF: pipeline_options,
                    }
                )
                logger.info("✅ Docling processor initialized")
            except Exception as e:
                logger.warning(f"Docling initialization failed: {e}")
                self.converter = None
        else:
            logger.info("Docling not available")
    
    def process_document_advanced(self, file_path: str) -> Dict[str, Any]:
        """Process document using Docling for advanced structure extraction"""
        if not self.converter:
            return self.fallback_processing(file_path)
        
        try:
            # Convert document
            result = self.converter.convert(file_path)
            
            # Extract structured content
            doc_content = {
                'text': result.document.export_to_markdown(),
                'tables': [],
                'figures': [],
                'metadata': {
                    'title': getattr(result.document, 'title', ''),
                    'page_count': len(result.document.pages) if hasattr(result.document, 'pages') else 0,
                    'processing_method': 'docling'
                }
            }
            
            # Extract tables and figures if available
            if hasattr(result.document, 'tables'):
                for table in result.document.tables:
                    doc_content['tables'].append({
                        'content': str(table),
                        'page': getattr(table, 'page', 0)
                    })
            
            if hasattr(result.document, 'figures'):
                for figure in result.document.figures:
                    doc_content['figures'].append({
                        'content': str(figure),
                        'page': getattr(figure, 'page', 0)
                    })
            
            return doc_content
            
        except Exception as e:
            logger.warning(f"Docling processing failed for {file_path}: {e}")
            return self.fallback_processing(file_path)
    
    def fallback_processing(self, file_path: str) -> Dict[str, Any]:
        """Fallback processing when Docling is not available"""
        try:
            if file_path.lower().endswith('.pdf'):
                doc = fitz.open(file_path)
                text = ""
                for page in doc:
                    text += page.get_text()
                doc.close()
                
                return {
                    'text': text,
                    'tables': [],
                    'figures': [],
                    'metadata': {
                        'title': Path(file_path).stem,
                        'page_count': len(doc) if 'doc' in locals() else 0,
                        'processing_method': 'fallback_pdf'
                    }
                }
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
                
                return {
                    'text': text,
                    'tables': [],
                    'figures': [],
                    'metadata': {
                        'title': Path(file_path).stem,
                        'page_count': 1,
                        'processing_method': 'fallback_text'
                    }
                }
        except Exception as e:
            logger.error(f"Fallback processing failed: {e}")
            return {
                'text': "",
                'tables': [],
                'figures': [],
                'metadata': {'processing_method': 'failed'}
            }

# Enhanced Document Processor with Multimodal Support
class EnhancedDocumentProcessor:
    """Enhanced document processor with multimodal and advanced processing capabilities"""
    
    def __init__(self, clip_manager: CLIPMultimodalManager, entity_extractor: EntityExtractor):
        self.clip_manager = clip_manager
        self.entity_extractor = entity_extractor
        self.docling_processor = DoclingProcessor()
        
        # Initialize semantic chunker
        if self.clip_manager.clip_model:
            self.semantic_chunker = SemanticChunker(
                self.clip_manager.clip_model,
                self.clip_manager.clip_processor
            )
        else:
            self.semantic_chunker = None
        
        # Fallback text splitter
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    async def process_pdf_multimodal(self, file_path: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Process PDF with text, images, and advanced structure extraction"""
        documents = []
        image_docs = []
        entities_list = []
        
        try:
            # Use Docling for advanced processing
            docling_result = self.docling_processor.process_document_advanced(file_path)
            
            # Process text content
            if docling_result['text']:
                # Extract entities from full text
                entities = self.entity_extractor.extract_entities(docling_result['text'])
                entities_list.extend(entities)
                
                # Create semantic chunks if available
                if self.semantic_chunker:
                    chunks = self.semantic_chunker.semantic_chunk(docling_result['text'])
                else:
                    chunks = self.text_splitter.split_text(docling_result['text'])
                
                # Create document objects for each chunk
                for i, chunk in enumerate(chunks):
                    chunk_entities = self.entity_extractor.extract_entities(chunk)
                    
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "type": "text",
                            "chunk_id": i,
                            "chunk_size": len(chunk),
                            "entities": [ent['text'] for ent in chunk_entities],
                            "processing_method": docling_result['metadata']['processing_method']
                        }
                    )
                    documents.append(doc)
            
            # Process images from PDF using PyMuPDF
            try:
                pdf_doc = fitz.open(file_path)
                for page_num in range(pdf_doc.page_count):
                    page = pdf_doc.load_page(page_num)
                    
                    for img_index, img in enumerate(page.get_images(full=True)):
                        try:
                            xref = img[0]
                            base_image = pdf_doc.extract_image(xref)
                            image_bytes = base_image["image"]
                            
                            # Convert to PIL Image
                            pil_image = Image.open(io.BytesIO(image_bytes)).convert("RGB")
                            
                            # Create unique identifier
                            image_id = f"page_{page_num}_img_{img_index}"
                            
                            # Store image
                            img_base64 = self.clip_manager.store_image(image_id, pil_image)
                            
                            # Create image document
                            image_doc = {
                                "image_id": image_id,
                                "page": page_num,
                                "base64": img_base64,
                                "embedding": self.clip_manager.embed_image(pil_image),
                                "metadata": {
                                    "source": file_path,
                                    "type": "image",
                                    "page": page_num,
                                    "image_index": img_index
                                }
                            }
                            image_docs.append(image_doc)
                            
                        except Exception as e:
                            logger.warning(f"Error processing image {img_index} on page {page_num}: {e}")
                            continue
                
                pdf_doc.close()
                
            except Exception as e:
                logger.warning(f"Image extraction failed: {e}")
            
            return documents, image_docs, entities_list
            
        except Exception as e:
            logger.error(f"Error processing PDF {file_path}: {e}")
            raise
    
    async def process_document(self, file_path: str, file_type: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Main document processing method"""
        if file_type.lower() == 'pdf':
            return await self.process_pdf_multimodal(file_path)
        elif file_type.lower() in ['docx', 'doc']:
            return await self.process_docx_enhanced(file_path)
        elif file_type.lower() == 'txt':
            return await self.process_txt_enhanced(file_path)
        else:
            raise ValueError(f"Unsupported file type: {file_type}")
    
    async def process_docx_enhanced(self, file_path: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Enhanced DOCX processing"""
        documents = []
        entities_list = []
        
        try:
            # Use Docling if available
            docling_result = self.docling_processor.process_document_advanced(file_path)
            
            if docling_result['text']:
                entities = self.entity_extractor.extract_entities(docling_result['text'])
                entities_list.extend(entities)
                
                if self.semantic_chunker:
                    chunks = self.semantic_chunker.semantic_chunk(docling_result['text'])
                else:
                    chunks = self.text_splitter.split_text(docling_result['text'])
                
                for i, chunk in enumerate(chunks):
                    chunk_entities = self.entity_extractor.extract_entities(chunk)
                    
                    doc = LangChainDocument(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "type": "docx",
                            "chunk_id": i,
                            "chunk_size": len(chunk),
                            "entities": [ent['text'] for ent in chunk_entities],
                            "processing_method": docling_result['metadata']['processing_method']
                        }
                    )
                    documents.append(doc)
            
            return documents, [], entities_list  # No images for DOCX yet
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path}: {e}")
            raise
    
    async def process_txt_enhanced(self, file_path: str) -> Tuple[List[LangChainDocument], List[Dict], List[Dict]]:
        """Enhanced text processing"""
        documents = []
        entities_list = []
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                content = await file.read()
            
            entities = self.entity_extractor.extract_entities(content)
            entities_list.extend(entities)
            
            if self.semantic_chunker:
                chunks = self.semantic_chunker.semantic_chunk(content)
            else:
                chunks = self.text_splitter.split_text(content)
            
            for i, chunk in enumerate(chunks):
                chunk_entities = self.entity_extractor.extract_entities(chunk)
                
                doc = LangChainDocument(
                    page_content=chunk,
                    metadata={
                        "source": file_path,
                        "type": "txt",
                        "chunk_id": i,
                        "chunk_size": len(chunk),
                        "entities": [ent['text'] for ent in chunk_entities],
                        "processing_method": "enhanced_text"
                    }
                )
                documents.append(doc)
            
            return documents, [], entities_list
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path}: {e}")
            raise

# HuggingFace Embedding Manager (Enhanced)
class HuggingFaceEmbeddingManager:
    """Enhanced HuggingFace embeddings with fallback options"""
    
    def __init__(self, model_name: str = None, cache_dir: str = None, device: str = 'cpu'):
        self.model_name = model_name or config.hf_embedding_model
        self.cache_dir = cache_dir or config.hf_cache_dir
        self.device = device or config.hf_device
        self.embedding_model = None
        self.llamaindex_embedding = None
        
        Path(self.cache_dir).mkdir(parents=True, exist_ok=True)
        self.setup_embeddings()
    
    def setup_embeddings(self):
        """Setup HuggingFace embeddings"""
        if not HUGGINGFACE_AVAILABLE:
            logger.error("HuggingFace transformers not available")
            return False
        
        try:
            self.embedding_model = SentenceTransformer(
                self.model_name,
                cache_folder=self.cache_dir,
                device=self.device
            )
            
            if LLAMAINDEX_AVAILABLE:
                try:
                    self.llamaindex_embedding = HuggingFaceEmbedding(
                        model_name=self.model_name,
                        cache_folder=self.cache_dir,
                        device=self.device
                    )
                    
                    if hasattr(Settings, 'embed_model'):
                        Settings.embed_model = self.llamaindex_embedding
                    
                    logger.info(f"✅ LlamaIndex HuggingFace embeddings initialized")
                except Exception as e:
                    logger.warning(f"LlamaIndex embedding setup failed: {e}")
                    self.llamaindex_embedding = None
            
            logger.info(f"✅ HuggingFace embeddings initialized: {self.model_name}")
            return True
            
        except Exception as e:
            logger.error(f"Failed to initialize HuggingFace embeddings: {e}")
            return False
    
    def embed_query(self, text: str) -> List[float]:
        """Embed a single query text"""
        if not self.embedding_model:
            raise ValueError("Embedding model not initialized")
        
        try:
            embedding = self.embedding_model.encode(text, normalize_embeddings=True)
            return embedding.tolist()
        except Exception as e:
            logger.error(f"Query embedding failed: {e}")
            raise
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed multiple documents"""
        if not self.embedding_model:
            raise ValueError("Embedding model not initialized")
        
        try:
            embeddings = self.embedding_model.encode(texts, normalize_embeddings=True)
            return embeddings.tolist()
        except Exception as e:
            logger.error(f"Document embedding failed: {e}")
            raise

# Weaviate Manager (Enhanced)
class WeaviateManager:
    """Enhanced Weaviate manager with multimodal support"""
    
    def __init__(self):
        self.client = None
        self.vector_store = None
        self.connection_type = None
        self.image_documents = []
        
        if WEAVIATE_AVAILABLE:
            self.setup_weaviate()
        else:
            logger.info("Weaviate not available - will use fallback storage")
    
    def setup_weaviate(self):
        """Setup Weaviate with proper connection"""
        if not WEAVIATE_AVAILABLE:
            return False
        
        try:
            self.client = weaviate.connect_to_local()
            
            if self.client.is_ready():
                logger.info("✅ Weaviate connected successfully (local)")
                self.connection_type = "local"
                return True
            else:
                self.client = None
                
        except Exception as e:
            logger.warning(f"Local Weaviate connection failed: {e}")
            self.client = None
        
        if config.weaviate_api_key and not self.client:
            try:
                self.client = weaviate.connect_to_weaviate_cloud(
                    cluster_url=config.weaviate_url,
                    auth_credentials=weaviate.auth.AuthApiKey(config.weaviate_api_key)
                )
                
                if self.client.is_ready():
                    logger.info("✅ Weaviate connected successfully (cloud)")
                    self.connection_type = "cloud"
                    return True
                else:
                    self.client = None
                    
            except Exception as e:
                logger.warning(f"Weaviate Cloud connection failed: {e}")
                self.client = None
        
        return False
    
    def create_vector_store(self, documents: List[LangChainDocument], embeddings):
        """Create vector store from documents"""
        if not self.client or not documents:
            return None
        
        try:
            self.vector_store = WeaviateVectorStore.from_documents(
                documents,
                embeddings,
                client=self.client,
                index_name="RagDocuments"
            )
            logger.info(f"Vector store created with {len(documents)} documents")
            return self.vector_store
            
        except Exception as e:
            logger.error(f"Failed to create vector store: {e}")
            return None
    
    def add_documents_to_store(self, documents: List[LangChainDocument]):
        """Add documents to existing vector store"""
        if not self.vector_store or not documents:
            return False
        
        try:
            self.vector_store.add_documents(documents)
            logger.info(f"Added {len(documents)} documents to vector store")
            return True
        except Exception as e:
            logger.error(f"Failed to add documents to vector store: {e}")
            return False
    
    def add_image_documents(self, image_docs: List[Dict]):
        """Store image documents separately"""
        self.image_documents.extend(image_docs)
        logger.info(f"Added {len(image_docs)} image documents")
    
    def search_similar(self, query: str, k: int = 5):
        """Search for similar documents"""
        if not self.vector_store:
            return []
        
        try:
            results = self.vector_store.similarity_search(query, k=k)
            return results
        except Exception as e:
            logger.error(f"Vector search failed: {e}")
            return []
    
    def search_images(self, query: str, clip_manager: CLIPMultimodalManager, k: int = 3):
        """Search for similar images using CLIP"""
        if not self.image_documents or not clip_manager.clip_model:
            return []
        
        try:
            query_embedding = clip_manager.embed_text(query)
            
            # Calculate similarities
            similarities = []
            for img_doc in self.image_documents:
                img_embedding = img_doc.get('embedding', np.zeros(512))
                if len(img_embedding) > 0 and len(query_embedding) > 0:
                    similarity = cosine_similarity(
                        query_embedding.reshape(1, -1),
                        img_embedding.reshape(1, -1)
                    )[0][0]
                    similarities.append((img_doc, similarity))
            
            # Sort by similarity and return top k
            similarities.sort(key=lambda x: x[1], reverse=True)
            return [img_doc for img_doc, sim in similarities[:k]]
            
        except Exception as e:
            logger.error(f"Image search failed: {e}")
            return []
    
    def close(self):
        """Close Weaviate connection"""
        if self.client:
            try:
                self.client.close()
                logger.info("Weaviate connection closed")
            except Exception as e:
                logger.warning(f"Error closing Weaviate connection: {e}")

# Enhanced RAG Bot
class EnhancedRAGBot:
    """Enhanced RAG Bot with multimodal, multi-hop reasoning, and advanced features"""
    
    def __init__(self):
        # Initialize core components
        self.embedding_manager = HuggingFaceEmbeddingManager()
        self.clip_manager = CLIPMultimodalManager()
        self.entity_extractor = EntityExtractor()
        self.weaviate_manager = WeaviateManager()
        
        # Initialize enhanced document processor
        self.document_processor = EnhancedDocumentProcessor(
            self.clip_manager, 
            self.entity_extractor
        )
        
        # Initialize Azure OpenAI LLM
        try:
            from langchain_openai import AzureChatOpenAI
            self.llm = AzureChatOpenAI(
                azure_endpoint=config.azure_openai_endpoint,
                api_key=config.azure_openai_api_key,
                api_version=config.azure_openai_api_version,
                deployment_name=config.azure_openai_deployment_name,
                model_name=config.azure_openai_model_name,
                temperature=config.azure_openai_temperature,
                max_tokens=config.azure_openai_max_tokens
            )
            logger.info("✅ Azure OpenAI LLM initialized")
        except Exception as e:
            logger.error(f"Failed to initialize LLM: {e}")
            self.llm = None
        
        # Initialize reasoning and recommendation engines
        self.multi_hop_reasoner = MultiHopReasoner(self.weaviate_manager, self.llm)
        self.recommendation_engine = RecommendationEngine(self.llm, self.entity_extractor)
        
        # Initialize document storage
        self.documents = []
        self.all_entities = []
        
        # Setup enhanced prompts
        self.setup_enhanced_prompts()
    
    def setup_enhanced_prompts(self):
        """Setup enhanced prompt templates"""
        self.multimodal_qa_prompt = PromptTemplate(
            template="""You are an advanced AI assistant with multimodal capabilities. Analyze both text and visual information to provide comprehensive answers.

Text Context:
{text_context}

Image Context:
{image_context}

Entities Identified:
{entities}

Multi-hop Reasoning Steps:
{reasoning_steps}

Question: {question}

Instructions:
1. Synthesize information from text, images, entities, and reasoning steps
2. Provide a detailed, well-structured answer
3. Reference specific sources when making claims
4. If images are relevant, describe their content and relationship to the question
5. Use the multi-hop reasoning to build a logical argument
6. Be precise and factual, avoiding speculation

Comprehensive Answer:""",
            input_variables=["text_context", "image_context", "entities", "reasoning_steps", "question"]
        )
        
        self.conclusion_prompt = PromptTemplate(
            template="""Based on the comprehensive analysis provided, create a concise conclusion that summarizes the key findings and insights.

Analysis: {analysis}
Key Entities: {entities}
Sources: {sources}

Create a conclusion that:
1. Summarizes the main findings
2. Highlights key insights
3. Provides actionable takeaways
4. Is concise but comprehensive (2-3 sentences)

Conclusion:""",
            input_variables=["analysis", "entities", "sources"]
        )
    
    async def process_and_store_document(self, file_path: str, file_type: str) -> Tuple[str, int, int, int, float]:
        """Enhanced document processing and storage"""
        start_time = datetime.now()
        
        try:
            doc_id = str(uuid.uuid4())
            
            # Process document with enhanced processor
            text_docs, image_docs, entities = await self.document_processor.process_document(file_path, file_type)
            
            if not text_docs and not image_docs:
                raise ValueError("No content extracted from document")
            
            # Store text documents in Weaviate
            if text_docs:
                if self.weaviate_manager.vector_store is None:
                    embedding_wrapper = HuggingFaceEmbeddingsWrapper(self.embedding_manager)
                    self.weaviate_manager.create_vector_store(text_docs, embedding_wrapper)
                else:
                    self.weaviate_manager.add_documents_to_store(text_docs)
                
                self.documents.extend(text_docs)
            
            # Store image documents
            if image_docs:
                self.weaviate_manager.add_image_documents(image_docs)
            
            # Store entities
            self.all_entities.extend(entities)
            
            processing_time = (datetime.now() - start_time).total_seconds()
            logger.info(f"Document {doc_id} processed successfully in {processing_time:.2f}s")
            
            return doc_id, len(text_docs), len(image_docs), len(entities), processing_time
            
        except Exception as e:
            logger.error(f"Error processing document {file_path}: {e}")
            raise
    
    async def enhanced_query(self, query: str, top_k: int = 5, use_multimodal: bool = True, 
                           enable_multi_hop: bool = True, include_recommendations: bool = True,
                           summary_length: str = "medium") -> Dict[str, Any]:
        """Enhanced query processing with all advanced features"""
        start_time = datetime.now()
        
        try:
            if not self.documents and not self.weaviate_manager.image_documents:
                raise ValueError("No documents have been processed yet")
            
            # Step 1: Multi-hop reasoning (if enabled)
            reasoning_steps = []
            if enable_multi_hop and self.llm:
                reasoning_results = self.multi_hop_reasoner.multi_hop_reason(query)
                reasoning_steps = [f"Step {i+1}: {step['answer']}" for i, step in enumerate(reasoning_results)]
            
            # Step 2: Retrieve relevant text documents
            text_docs = []
            if self.weaviate_manager.vector_store:
                text_docs = self.weaviate_manager.search_similar(query, k=top_k)
            
            # Step 3: Retrieve relevant images (if multimodal enabled)
            image_docs = []
            if use_multimodal and self.weaviate_manager.image_documents:
                image_docs = self.weaviate_manager.search_images(query, self.clip_manager, k=3)
            
            # Step 4: Extract query-specific entities
            query_entities = self.entity_extractor.extract_entities(query)
            
            # Step 5: Prepare contexts
            text_context = ""
            if text_docs:
                text_context = "\n\n".join([
                    f"Source {i+1}: {doc.page_content}"
                    for i, doc in enumerate(text_docs)
                ])
            
            image_context = ""
            if image_docs:
                image_descriptions = []
                for i, img_doc in enumerate(image_docs):
                    img_desc = f"Image {i+1} (Page {img_doc['page']}): Visual content from document"
                    image_descriptions.append(img_desc)
                image_context = "\n".join(image_descriptions)
            
            entities_context = ", ".join([ent['text'] for ent in query_entities[:10]])
            reasoning_context = "\n".join(reasoning_steps)
            
            # Step 6: Generate comprehensive answer
            answer = ""
            if self.llm:
                try:
                    # Create multimodal message
                    content = []
                    
                    # Add text prompt
                    prompt_text = self.multimodal_qa_prompt.format(
                        text_context=text_context,
                        image_context=image_context,
                        entities=entities_context,
                        reasoning_steps=reasoning_context,
                        question=query
                    )
                    
                    content.append({
                        "type": "text",
                        "text": prompt_text
                    })
                    
                    # Add relevant images
                    for img_doc in image_docs[:2]:  # Limit to 2 images
                        if img_doc.get('base64'):
                            content.append({
                                "type": "text",
                                "text": f"\n[Image from page {img_doc['page']}]:\n"
                            })
                            content.append({
                                "type": "image_url",
                                "image_url": {
                                    "url": f"data:image/png;base64,{img_doc['base64']}"
                                }
                            })
                    
                    # Generate response
                    message = HumanMessage(content=content)
                    response = self.llm.invoke([message])
                    answer = response.content
                    
                except Exception as e:
                    logger.error(f"LLM invocation failed: {e}")
                    answer = f"Found relevant information but failed to generate response: {str(e)}"
            else:
                answer = "LLM not available for answer generation"
            
            # Step 7: Generate conclusion
            conclusion = ""
            if self.llm and answer:
                try:
                    conclusion_messages = [
                        SystemMessage(content="You are an expert at creating insightful conclusions."),
                        HumanMessage(content=self.conclusion_prompt.format(
                            analysis=answer[:1000],
                            entities=entities_context,
                            sources=str(len(text_docs + image_docs))
                        ))
                    ]
                    
                    conclusion_response = self.llm.invoke(conclusion_messages)
                    conclusion = conclusion_response.content
                    
                except Exception as e:
                    logger.warning(f"Conclusion generation failed: {e}")
                    conclusion = "Key insights derived from comprehensive analysis of available sources."
            
            # Step 8: Generate recommendations
            recommendations = []
            if include_recommendations:
                recommendations = self.recommendation_engine.generate_recommendations(
                    query, answer, query_entities, text_docs + image_docs
                )
            
            # Step 9: Generate summary
            summary = answer[:200] + "..." if len(answer) > 200 else answer
            
            # Step 10: Prepare response
            sources = []
            for i, doc in enumerate(text_docs):
                sources.append({
                    "id": i,
                    "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    "metadata": doc.metadata,
                    "type": "text"
                })
            
            image_sources = []
            for i, img_doc in enumerate(image_docs):
                image_sources.append({
                    "id": i,
                    "image_id": img_doc['image_id'],
                    "page": img_doc['page'],
                    "metadata": img_doc['metadata'],
                    "type": "image"
                })
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return {
                "answer": answer,
                "summary": summary,
                "conclusion": conclusion,
                "sources": sources,
                "entities": [{"text": ent['text'], "label": ent['label'], "confidence": ent['confidence']} for ent in query_entities],
                "recommendations": recommendations,
                "confidence_score": 0.8,  # Calculate based on various factors
                "processing_time": processing_time,
                "multi_hop_reasoning": reasoning_steps,
                "image_sources": image_sources
            }
            
        except Exception as e:
            logger.error(f"Error processing enhanced query: {e}")
            raise
    
    def close(self):
        """Clean up resources"""
        try:
            if self.weaviate_manager:
                self.weaviate_manager.close()
        except Exception as e:
            logger.warning(f"Error closing Weaviate manager: {e}")
        
        logger.info("Enhanced RAG Bot resources closed")

# LangChain compatible embeddings wrapper
class HuggingFaceEmbeddingsWrapper:
    """Wrapper to make HuggingFace embeddings compatible with LangChain"""
    
    def __init__(self, embedding_manager: HuggingFaceEmbeddingManager):
        self.embedding_manager = embedding_manager
    
    def embed_query(self, text: str) -> List[float]:
        """Embed query text"""
        return self.embedding_manager.embed_query(text)
    
    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        """Embed multiple documents"""
        return self.embedding_manager.embed_documents(texts)

# Global variable
enhanced_rag_bot = None

# Lifespan event handler
@asynccontextmanager
async def lifespan(app: FastAPI):
    # Startup
    global enhanced_rag_bot
    logger.info("🚀 Enhanced Multimodal RAG Bot API starting up...")
    
    try:
        enhanced_rag_bot = EnhancedRAGBot()
        logger.info("✅ Enhanced Multimodal RAG Bot initialized successfully")
    except Exception as e:
        logger.error(f"Failed to initialize Enhanced RAG Bot: {e}")
        enhanced_rag_bot = None
    
    yield
    
    # Shutdown
    logger.info("🛑 Enhanced Multimodal RAG Bot API shutting down...")
    if enhanced_rag_bot:
        try:
            enhanced_rag_bot.close()
            logger.info("✅ Enhanced RAG Bot closed successfully")
        except Exception as e:
            logger.warning(f"Error during shutdown: {e}")

# FastAPI application
app = FastAPI(
    title="Enhanced Multimodal RAG Bot with Advanced Features", 
    version="3.0.0",
    description="Advanced multimodal RAG system with Docling, CLIP, multi-hop reasoning, entity extraction, and recommendations",
    lifespan=lifespan
)

# CORS middleware
app.add_middleware(
    CORSMiddleware,
    allow_origins=config.cors_origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Security
security = HTTPBearer(auto_error=False)

def verify_token(credentials: HTTPAuthorizationCredentials = Security(security)):
    return credentials

# API Endpoints
@app.post("/upload", response_model=DocumentUploadResponse)
async def upload_document(
    file: UploadFile = File(...),
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Upload and process a document with enhanced multimodal processing"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        if not file.filename:
            raise HTTPException(status_code=400, detail="No file provided")
        
        file_extension = file.filename.split('.')[-1].lower()
        if file_extension not in config.allowed_file_types:
            raise HTTPException(
                status_code=400, 
                detail=f"File type {file_extension} not supported. Allowed types: {config.allowed_file_types}"
            )
        
        # Save uploaded file
        upload_dir = Path("uploads")
        upload_dir.mkdir(exist_ok=True)
        file_path = upload_dir / f"{uuid.uuid4()}_{file.filename}"
        
        async with aiofiles.open(file_path, 'wb') as f:
            content = await file.read()
            if len(content) > config.max_file_size:
                raise HTTPException(status_code=400, detail="File too large")
            await f.write(content)
        
        # Process document
        doc_id, chunks_created, images_processed, entities_extracted, processing_time = await enhanced_rag_bot.process_and_store_document(
            str(file_path), file_extension
        )
        
        # Clean up uploaded file
        file_path.unlink()
        
        return DocumentUploadResponse(
            message="Document processed successfully with enhanced multimodal processing",
            document_id=doc_id,
            chunks_created=chunks_created,
            images_processed=images_processed,
            entities_extracted=entities_extracted,
            processing_time=processing_time
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Upload error: {e}")
        if 'file_path' in locals() and file_path.exists():
            file_path.unlink()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query", response_model=QueryResponse)
async def query_documents(
    request: QueryRequest,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Enhanced multimodal query with all advanced features"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        result = await enhanced_rag_bot.enhanced_query(
            query=request.query,
            top_k=request.top_k,
            use_multimodal=request.use_multimodal,
            enable_multi_hop=request.enable_multi_hop,
            include_recommendations=request.include_recommendations,
            summary_length=request.summary_length
        )
        
        return QueryResponse(**result)
        
    except Exception as e:
        logger.error(f"Query error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/health")
async def health_check():
    """Enhanced health check endpoint"""
    weaviate_status = False
    clip_status = False
    docling_status = False
    spacy_status = False
    connection_type = None
    
    if enhanced_rag_bot:
        weaviate_status = enhanced_rag_bot.weaviate_manager.client is not None
        connection_type = enhanced_rag_bot.weaviate_manager.connection_type
        clip_status = enhanced_rag_bot.clip_manager.clip_model is not None
        docling_status = enhanced_rag_bot.document_processor.docling_processor.converter is not None
        spacy_status = enhanced_rag_bot.entity_extractor.nlp is not None
    
    return {
        "status": "healthy" if enhanced_rag_bot else "degraded", 
        "timestamp": datetime.now().isoformat(),
        "weaviate_connected": weaviate_status,
        "weaviate_connection_type": connection_type,
        "clip_initialized": clip_status,
        "docling_available": docling_status,
        "spacy_available": spacy_status,
        "enhanced_rag_bot_initialized": enhanced_rag_bot is not None,
        "huggingface_embeddings": HUGGINGFACE_AVAILABLE,
        "sentence_transformers_available": SENTENCE_TRANSFORMERS_AVAILABLE,
        "llamaindex_available": LLAMAINDEX_AVAILABLE,
        "docling_available_global": DOCLING_AVAILABLE
    }

@app.get("/stats")
async def get_enhanced_stats(credentials: HTTPAuthorizationCredentials = Depends(verify_token)):
    """Get enhanced system statistics"""
    if enhanced_rag_bot is None:
        return {
            "error": "Enhanced RAG Bot not initialized",
            "total_documents": 0,
            "total_images": 0,
            "total_entities": 0
        }
    
    return {
        "total_text_documents": len(enhanced_rag_bot.documents),
        "total_image_documents": len(enhanced_rag_bot.weaviate_manager.image_documents),
        "total_entities": len(enhanced_rag_bot.all_entities),
        "vector_store_initialized": enhanced_rag_bot.weaviate_manager.vector_store is not None,
        "weaviate_connection": enhanced_rag_bot.weaviate_manager.connection_type,
        "clip_model": enhanced_rag_bot.clip_manager.model_name,
        "embedding_model": enhanced_rag_bot.embedding_manager.model_name,
        "huggingface_device": enhanced_rag_bot.embedding_manager.device,
        "multimodal_capabilities": {
            "text_processing": True,
            "image_processing": enhanced_rag_bot.clip_manager.clip_model is not None,
            "entity_extraction": enhanced_rag_bot.entity_extractor.nlp is not None,
            "semantic_chunking": enhanced_rag_bot.document_processor.semantic_chunker is not None,
            "multi_hop_reasoning": enhanced_rag_bot.llm is not None,
            "docling_processing": enhanced_rag_bot.document_processor.docling_processor.converter is not None
        }
    }

@app.get("/documents")
async def list_enhanced_documents(credentials: HTTPAuthorizationCredentials = Depends(verify_token)):
    """List all processed documents with enhanced information"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        documents_info = []
        seen_sources = set()
        
        # Process text documents
        for doc in enhanced_rag_bot.documents:
            source = doc.metadata.get('source', 'Unknown')
            if source not in seen_sources:
                seen_sources.add(source)
                
                # Count chunks and entities for this source
                source_chunks = [d for d in enhanced_rag_bot.documents if d.metadata.get('source') == source]
                source_entities = []
                for chunk in source_chunks:
                    source_entities.extend(chunk.metadata.get('entities', []))
                
                documents_info.append({
                    "source": source,
                    "type": doc.metadata.get('type', 'Unknown'),
                    "chunk_count": len(source_chunks),
                    "entity_count": len(set(source_entities)),
                    "processing_method": doc.metadata.get('processing_method', 'Unknown'),
                    "sample_content": doc.page_content[:100] + "..." if len(doc.page_content) > 100 else doc.page_content,
                    "sample_entities": list(set(source_entities))[:5]
                })
        
        # Add image document info
        image_sources = {}
        for img_doc in enhanced_rag_bot.weaviate_manager.image_documents:
            source = img_doc['metadata'].get('source', 'Unknown')
            if source not in image_sources:
                image_sources[source] = 0
            image_sources[source] += 1
        
        return {
            "total_unique_documents": len(documents_info),
            "total_text_chunks": len(enhanced_rag_bot.documents),
            "total_image_documents": len(enhanced_rag_bot.weaviate_manager.image_documents),
            "total_entities": len(enhanced_rag_bot.all_entities),
            "documents": documents_info,
            "image_sources": image_sources,
            "top_entities": [{"text": ent['text'], "label": ent['label']} for ent in enhanced_rag_bot.all_entities[:10]]
        }
        
    except Exception as e:
        logger.error(f"Error listing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.delete("/documents")
async def clear_enhanced_documents(credentials: HTTPAuthorizationCredentials = Depends(verify_token)):
    """Clear all processed documents and reset system"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        document_count = len(enhanced_rag_bot.documents)
        image_count = len(enhanced_rag_bot.weaviate_manager.image_documents)
        entity_count = len(enhanced_rag_bot.all_entities)
        
        # Clear all data
        enhanced_rag_bot.documents.clear()
        enhanced_rag_bot.weaviate_manager.image_documents.clear()
        enhanced_rag_bot.all_entities.clear()
        enhanced_rag_bot.clip_manager.image_data_store.clear()
        
        # Reset vector store
        enhanced_rag_bot.weaviate_manager.vector_store = None
        
        logger.info(f"Cleared {document_count} text documents, {image_count} images, {entity_count} entities")
        
        return {
            "message": f"Successfully cleared all data and reset system",
            "cleared_text_documents": document_count,
            "cleared_image_documents": image_count,
            "cleared_entities": entity_count
        }
        
    except Exception as e:
        logger.error(f"Error clearing documents: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/entities")
async def get_entities(
    limit: int = 50,
    entity_type: Optional[str] = None,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Get extracted entities with filtering"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        entities = enhanced_rag_bot.all_entities
        
        # Filter by entity type if specified
        if entity_type:
            entities = [ent for ent in entities if ent.get('label') == entity_type.upper()]
        
        # Count occurrences
        entity_counts = Counter(ent['text'].lower() for ent in entities)
        
        # Get top entities
        top_entities = []
        for entity_text, count in entity_counts.most_common(limit):
            # Find original entity data
            original_entity = next(
                (ent for ent in entities if ent['text'].lower() == entity_text), 
                {"text": entity_text, "label": "UNKNOWN", "confidence": 0.0}
            )
            
            top_entities.append({
                "text": original_entity['text'],
                "label": original_entity['label'],
                "confidence": original_entity['confidence'],
                "count": count,
                "extraction_method": original_entity.get('method', 'unknown')
            })
        
        # Get entity type distribution
        type_distribution = Counter(ent['label'] for ent in entities)
        
        return {
            "total_entities": len(entities),
            "unique_entities": len(entity_counts),
            "top_entities": top_entities,
            "entity_types": dict(type_distribution),
            "available_types": list(type_distribution.keys())
        }
        
    except Exception as e:
        logger.error(f"Error getting entities: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/images")
async def get_images(
    limit: int = 20,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Get processed images information"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        images = enhanced_rag_bot.weaviate_manager.image_documents[:limit]
        
        image_info = []
        for img in images:
            image_info.append({
                "image_id": img['image_id'],
                "page": img['page'],
                "source": img['metadata'].get('source', 'Unknown'),
                "has_embedding": len(img.get('embedding', [])) > 0,
                "embedding_dimension": len(img.get('embedding', [])),
                "has_base64": bool(img.get('base64'))
            })
        
        # Get source distribution
        source_distribution = Counter(img['metadata'].get('source', 'Unknown') for img in enhanced_rag_bot.weaviate_manager.image_documents)
        
        return {
            "total_images": len(enhanced_rag_bot.weaviate_manager.image_documents),
            "images": image_info,
            "source_distribution": dict(source_distribution),
            "clip_model": enhanced_rag_bot.clip_manager.model_name
        }
        
    except Exception as e:
        logger.error(f"Error getting images: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/search/multimodal")
async def multimodal_search(
    query: str,
    search_images: bool = True,
    search_text: bool = True,
    top_k: int = 5,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Perform multimodal search across text and images"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        results = {
            "query": query,
            "text_results": [],
            "image_results": [],
            "search_time": 0
        }
        
        start_time = datetime.now()
        
        # Search text documents
        if search_text and enhanced_rag_bot.weaviate_manager.vector_store:
            text_docs = enhanced_rag_bot.weaviate_manager.search_similar(query, k=top_k)
            
            for i, doc in enumerate(text_docs):
                results["text_results"].append({
                    "rank": i + 1,
                    "content": doc.page_content[:300] + "..." if len(doc.page_content) > 300 else doc.page_content,
                    "metadata": doc.metadata,
                    "entities": doc.metadata.get('entities', [])
                })
        
        # Search image documents
        if search_images and enhanced_rag_bot.weaviate_manager.image_documents:
            image_docs = enhanced_rag_bot.weaviate_manager.search_images(
                query, enhanced_rag_bot.clip_manager, k=min(top_k, 5)
            )
            
            for i, img_doc in enumerate(image_docs):
                results["image_results"].append({
                    "rank": i + 1,
                    "image_id": img_doc['image_id'],
                    "page": img_doc['page'],
                    "source": img_doc['metadata'].get('source', 'Unknown'),
                    "metadata": img_doc['metadata']
                })
        
        results["search_time"] = (datetime.now() - start_time).total_seconds()
        
        return results
        
    except Exception as e:
        logger.error(f"Multimodal search error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/test/clip")
async def test_clip_embeddings(
    text: str = "This is a test sentence for CLIP embedding",
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Test CLIP embeddings for text"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        start_time = datetime.now()
        
        # Test text embedding
        text_embedding = enhanced_rag_bot.clip_manager.embed_text(text)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "model_name": enhanced_rag_bot.clip_manager.model_name,
            "text": text,
            "embedding_dimension": len(text_embedding),
            "processing_time": processing_time,
            "sample_embedding": text_embedding[:10].tolist() if len(text_embedding) > 0 else [],
            "clip_model_available": enhanced_rag_bot.clip_manager.clip_model is not None
        }
        
    except Exception as e:
        logger.error(f"CLIP test error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/test/entities")
async def test_entity_extraction(
    text: str = "Apple Inc. was founded in 1976 by Steve Jobs. The company's revenue in 2023 was $394.3 billion.",
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Test entity extraction"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        start_time = datetime.now()
        
        entities = enhanced_rag_bot.entity_extractor.extract_entities(text)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "text": text,
            "entities": entities,
            "entity_count": len(entities),
            "processing_time": processing_time,
            "spacy_available": enhanced_rag_bot.entity_extractor.nlp is not None,
            "extraction_methods": list(set(ent.get('method', 'unknown') for ent in entities))
        }
        
    except Exception as e:
        logger.error(f"Entity extraction test error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/query/reasoning")
async def test_multi_hop_reasoning(
    query: str,
    credentials: HTTPAuthorizationCredentials = Depends(verify_token)
):
    """Test multi-hop reasoning capabilities"""
    if enhanced_rag_bot is None:
        raise HTTPException(status_code=503, detail="Enhanced RAG Bot not initialized")
    
    try:
        start_time = datetime.now()
        
        reasoning_steps = enhanced_rag_bot.multi_hop_reasoner.multi_hop_reason(query)
        
        processing_time = (datetime.now() - start_time).total_seconds()
        
        return {
            "original_query": query,
            "reasoning_steps": reasoning_steps,
            "step_count": len(reasoning_steps),
            "processing_time": processing_time,
            "llm_available": enhanced_rag_bot.llm is not None
        }
        
    except Exception as e:
        logger.error(f"Multi-hop reasoning test error: {e}")
        raise HTTPException(status_code=500, detail=str(e))

@app.get("/config/enhanced")
async def get_enhanced_config(credentials: HTTPAuthorizationCredentials = Depends(verify_token)):
    """Get enhanced system configuration"""
    return {
        "multimodal_capabilities": {
            "clip_model": config.clip_model_name,
            "text_embeddings": config.hf_embedding_model,
            "docling_available": DOCLING_AVAILABLE,
            "spacy_available": SPACY_AVAILABLE
        },
        "vector_databases": {
            "weaviate_available": WEAVIATE_AVAILABLE,
            "weaviate_url": config.weaviate_url
        },
        "language_models": {
            "azure_openai_configured": bool(config.azure_openai_endpoint and config.azure_openai_api_key),
            "llamaindex_available": LLAMAINDEX_AVAILABLE
        },
        "processing_features": {
            "semantic_chunking": True,
            "entity_extraction": SPACY_AVAILABLE,
            "multi_hop_reasoning": True,
            "recommendation_engine": True,
            "advanced_document_processing": DOCLING_AVAILABLE
        },
        "file_support": {
            "max_file_size": config.max_file_size,
            "allowed_file_types": config.allowed_file_types
        },
        "system_settings": {
            "debug": config.debug,
            "hf_device": config.hf_device,
            "persist_dir": config.persist_dir
        }
    }

# Error handlers
@app.exception_handler(HTTPException)
async def http_exception_handler(request, exc):
    from fastapi.responses import JSONResponse
    return JSONResponse(
        status_code=exc.status_code,
        content={
            "error": exc.detail,
            "status_code": exc.status_code,
            "timestamp": datetime.now().isoformat()
        }
    )

@app.exception_handler(Exception)
async def general_exception_handler(request, exc):
    from fastapi.responses import JSONResponse
    logger.error(f"Unhandled exception: {exc}")
    return JSONResponse(
        status_code=500,
        content={
            "error": "Internal server error",
            "status_code": 500,
            "timestamp": datetime.now().isoformat(),
            "detail": str(exc) if config.debug else "Internal server error"
        }
    )

if __name__ == "__main__":
    import sys
    
    # Suppress specific warnings
    warnings.filterwarnings("ignore", category=DeprecationWarning, module="websockets")
    warnings.filterwarnings("ignore", category=DeprecationWarning, module="uvicorn")
    warnings.filterwarnings("ignore", message=".*swigvarlink.*")
    
    # Validate required environment variables
    required_vars = [
        'AZURE_OPENAI_ENDPOINT',
        'AZURE_OPENAI_API_KEY',
        'AZURE_OPENAI_DEPLOYMENT_NAME'
    ]
    
    missing_vars = [var for var in required_vars if not os.getenv(var)]
    if missing_vars:
        logger.error(f"Missing required environment variables: {missing_vars}")
        logger.error("Please set these variables in your .env file or environment")
        sys.exit(1)
    
    # Create required directories
    for directory in ["uploads", config.persist_dir, config.hf_cache_dir]:
        Path(directory).mkdir(exist_ok=True)
    
    logger.info("🚀 Starting Enhanced Multimodal RAG Bot API server...")
    logger.info(f"Weaviate available: {WEAVIATE_AVAILABLE}")
    logger.info(f"Docling available: {DOCLING_AVAILABLE}")
    logger.info(f"CLIP model: {config.clip_model_name}")
    logger.info(f"LlamaIndex available: {LLAMAINDEX_AVAILABLE}")
    logger.info(f"HuggingFace available: {HUGGINGFACE_AVAILABLE}")
    logger.info(f"Spacy available: {SPACY_AVAILABLE}")
    logger.info(f"SentenceTransformers available: {SENTENCE_TRANSFORMERS_AVAILABLE}")
    logger.info(f"Embedding model: {config.hf_embedding_model}")
    logger.info(f"Device: {config.hf_device}")
    logger.info(f"Storage directory: {config.persist_dir}")
    
    try:
        if config.debug:
            # For development with reload, use string import
            uvicorn.run(
                "enhanced_multimodal_rag:app",
                host="0.0.0.0",
                port=8000,
                reload=True,
                log_level="info"
            )
        else:
            # For production, use app object directly
            uvicorn.run(
                app,
                host="0.0.0.0",
                port=8000,
                log_level="info"
            )
    except KeyboardInterrupt:
        logger.info("Server stopped by user")
    except Exception as e:
        logger.error(f"Server error: {e}")
        sys.exit(1)