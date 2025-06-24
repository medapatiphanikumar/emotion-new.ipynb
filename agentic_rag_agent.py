# agents/agentic_rag_agent.py

import os
import logging
from pathlib import Path
from typing import List, Dict, Any, Optional, TypedDict
from tqdm import tqdm
import json
import time
import re
import unicodedata
import base64

# LangChain imports
from langchain.docstore.document import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from langchain_community.vectorstores.utils import DistanceStrategy
from langchain_core.prompts import PromptTemplate
from langchain_text_splitters import RecursiveCharacterTextSplitter
from transformers import AutoTokenizer

# Updated HuggingFace embeddings import
try:
    from langchain_huggingface import HuggingFaceEmbeddings
except ImportError:
    from langchain_community.embeddings import HuggingFaceEmbeddings

# Azure OpenAI
from langchain_openai import AzureChatOpenAI

# Document loaders
from langchain_community.document_loaders import TextLoader
try:
    from langchain_community.document_loaders import PyPDFLoader
except ImportError:
    PyPDFLoader = None
try:
    from langchain_community.document_loaders import Docx2txtLoader
except ImportError:
    Docx2txtLoader = None

# PDF processing libraries
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False

try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False

# Smolagents imports for retrieval tool
try:
    from smolagents import Tool, InferenceClientModel, ToolCallingAgent
    from langchain_core.vectorstores import VectorStore
    SMOLAGENTS_AVAILABLE = True
except ImportError:
    SMOLAGENTS_AVAILABLE = False

# LangGraph imports
try:
    from langgraph.graph import START, StateGraph
    LANGGRAPH_AVAILABLE = True
except ImportError:
    LANGGRAPH_AVAILABLE = False

# Import settings
from core.config import settings

logger = logging.getLogger(__name__)

def sanitize_filename(filename: str) -> str:
    """Sanitize filename for safe filesystem operations"""
    # Normalize unicode characters
    filename = unicodedata.normalize('NFKD', filename)
    # Replace problematic characters
    filename = re.sub(r'[^\w\s\-_\.]', '_', filename)
    # Replace multiple underscores/spaces with single underscore
    filename = re.sub(r'[_\s]+', '_', filename)
    # Remove leading/trailing underscores
    filename = filename.strip('_')
    # Limit length
    if len(filename) > 200:
        filename = filename[:200]
    return filename

class DocumentExtractionTool(Tool):
    """ReACT tool for document text extraction with reasoning"""
    name = "document_extractor"
    description = "Extracts text content from PDF, DOCX, or TXT documents using intelligent extraction methods with reasoning"
    inputs = {
        "file_path": {
            "type": "string", 
            "description": "Path to the document file to extract text from"
        },
        "extraction_method": {
            "type": "string",
            "description": "Specific extraction method to try: 'auto', 'pymupdf', 'pdfplumber', 'pypdf2', 'docx', 'text'"
        },
        "reasoning": {
            "type": "string",
            "description": "Reasoning about why this extraction method should be used"
        }
    }
    output_type = "string"

    def __init__(self, **kwargs):
        super().__init__(**kwargs)

    def forward(self, file_path: str, extraction_method: str = "auto", reasoning: str = "") -> str:
        """Extract text using ReACT reasoning approach"""
        
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        # ReACT Thought process
        thoughts = []
        thoughts.append(f"THOUGHT: I need to extract text from {file_path.name} (type: {file_ext})")
        thoughts.append(f"REASONING: {reasoning}")
        
        # Determine best extraction strategy
        if extraction_method == "auto":
            if file_ext == '.pdf':
                thoughts.append("THOUGHT: This is a PDF file. I should try multiple extraction methods in order of reliability.")
                extraction_methods = ['pymupdf', 'pdfplumber', 'pypdf2']
            elif file_ext in ['.docx', '.doc']:
                thoughts.append("THOUGHT: This is a DOCX file. I should use DOCX-specific extraction.")
                extraction_methods = ['docx']
            elif file_ext in ['.txt', '.md']:
                thoughts.append("THOUGHT: This is a text file. I can read it directly.")
                extraction_methods = ['text']
            else:
                thoughts.append("THOUGHT: Unknown file type. I'll try text extraction as fallback.")
                extraction_methods = ['text']
        else:
            extraction_methods = [extraction_method]
        
        # ACTION: Try extraction methods
        for method in extraction_methods:
            thoughts.append(f"ACTION: Attempting extraction using {method}")
            
            try:
                if method == 'pymupdf' and PYMUPDF_AVAILABLE:
                    result = self._extract_with_pymupdf(file_path)
                elif method == 'pdfplumber' and PDFPLUMBER_AVAILABLE:
                    result = self._extract_with_pdfplumber(file_path)
                elif method == 'pypdf2' and PYPDF2_AVAILABLE:
                    result = self._extract_with_pypdf2(file_path)
                elif method == 'docx':
                    result = self._extract_with_docx(file_path)
                elif method == 'text':
                    result = self._extract_with_text(file_path)
                else:
                    thoughts.append(f"OBSERVATION: Method {method} is not available or supported")
                    continue
                
                # OBSERVATION: Evaluate extraction quality
                if result and result.strip():
                    char_count = len(result)
                    word_count = len(result.split())
                    thoughts.append(f"OBSERVATION: Successfully extracted {char_count} characters, {word_count} words using {method}")
                    
                    # Quality assessment
                    if char_count > 100 and word_count > 10:
                        thoughts.append("THOUGHT: This appears to be good quality extraction with substantial content")
                        thoughts.append(f"ACTION: Returning extracted text from {method}")
                        
                        return "\n".join(thoughts) + f"\n\nEXTRACTED_TEXT:\n{result}"
                    else:
                        thoughts.append("THOUGHT: Extraction successful but content seems limited, trying next method")
                else:
                    thoughts.append(f"OBSERVATION: Method {method} returned empty or no content")
                    
            except Exception as e:
                thoughts.append(f"OBSERVATION: Method {method} failed with error: {str(e)}")
        
        # If all methods fail
        thoughts.append("THOUGHT: All extraction methods failed or returned insufficient content")
        thoughts.append("ACTION: Returning error message with reasoning")
        
        return "\n".join(thoughts) + "\n\nERROR: Could not extract meaningful text from document"

    def _extract_with_pymupdf(self, file_path: Path) -> str:
        """Extract using PyMuPDF"""
        pdf_document = fitz.open(str(file_path))
        content = ""
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            content += page.get_text() + "\n"
        pdf_document.close()
        return content.strip()

    def _extract_with_pdfplumber(self, file_path: Path) -> str:
        """Extract using pdfplumber"""
        content = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
        return content.strip()

    def _extract_with_pypdf2(self, file_path: Path) -> str:
        """Extract using PyPDF2"""
        content = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                content += page.extract_text() + "\n"
        return content.strip()

    def _extract_with_docx(self, file_path: Path) -> str:
        """Extract using python-docx"""
        try:
            from docx import Document
            doc = Document(file_path)
            content = ""
            for paragraph in doc.paragraphs:
                content += paragraph.text + "\n"
            
            # Also extract from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        content += cell.text + " "
                    content += "\n"
            return content.strip()
        except ImportError:
            raise Exception("python-docx not available")

    def _extract_with_text(self, file_path: Path) -> str:
        """Extract plain text"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read().strip()

class ReACTDocumentProcessor:
    """ReACT-based document processor that reasons through extraction challenges"""
    
    def __init__(self, llm=None):
        self.llm = llm
        self.extraction_tool = DocumentExtractionTool()
        
    def process_document_with_react(self, file_path: str, file_id: str) -> List[Document]:
        """Process document using ReACT reasoning"""
        
        file_path = Path(file_path)
        
        # Initial ReACT reasoning
        initial_reasoning = self._analyze_document_characteristics(file_path)
        
        # Extract text using ReACT tool
        extraction_result = self.extraction_tool.forward(
            file_path=str(file_path),
            extraction_method="auto",
            reasoning=initial_reasoning
        )
        
        # Parse the result
        if "EXTRACTED_TEXT:" in extraction_result:
            text_content = extraction_result.split("EXTRACTED_TEXT:")[1].strip()
            reasoning_log = extraction_result.split("EXTRACTED_TEXT:")[0].strip()
            
            if text_content and not text_content.startswith("ERROR:"):
                # Create document with reasoning metadata
                doc = Document(
                    page_content=text_content,
                    metadata={
                        "source": file_id,
                        "file_path": str(file_path),
                        "file_type": file_path.suffix.lower(),
                        "extraction_method": "react",
                        "reasoning_log": reasoning_log,
                        "extraction_quality": self._assess_extraction_quality(text_content)
                    }
                )
                return [doc]
        
        # If extraction failed, try advanced ReACT reasoning with LLM
        if self.llm:
            return self._llm_guided_extraction(file_path, file_id, extraction_result)
        
        return []
    
    def _analyze_document_characteristics(self, file_path: Path) -> str:
        """Analyze document to determine best extraction approach"""
        file_size = file_path.stat().st_size
        file_ext = file_path.suffix.lower()
        
        reasoning = f"File analysis: {file_path.name} is {file_size} bytes, type {file_ext}. "
        
        if file_ext == '.pdf':
            if file_size > 10_000_000:  # 10MB
                reasoning += "Large PDF detected - may need robust extraction method like PyMuPDF."
            elif file_size < 50_000:  # 50KB
                reasoning += "Small PDF - might be mostly text or could be scanned."
            else:
                reasoning += "Medium-sized PDF - standard extraction should work."
        elif file_ext in ['.docx', '.doc']:
            reasoning += "Microsoft Word document - should extract cleanly with python-docx."
        elif file_ext in ['.txt', '.md']:
            reasoning += "Text file - direct reading should be sufficient."
        
        return reasoning
    
    def _assess_extraction_quality(self, text: str) -> Dict[str, Any]:
        """Assess the quality of extracted text"""
        if not text:
            return {"quality": "empty", "score": 0}
        
        char_count = len(text)
        word_count = len(text.split())
        line_count = len(text.split('\n'))
        
        # Simple quality metrics
        avg_word_length = char_count / max(word_count, 1)
        
        quality_score = 0
        quality_indicators = []
        
        if char_count > 1000:
            quality_score += 3
            quality_indicators.append("substantial_content")
        elif char_count > 100:
            quality_score += 2
            quality_indicators.append("moderate_content")
        else:
            quality_score += 1
            quality_indicators.append("minimal_content")
        
        if 3 <= avg_word_length <= 10:
            quality_score += 2
            quality_indicators.append("normal_word_length")
        
        if word_count > 50:
            quality_score += 2
            quality_indicators.append("sufficient_words")
        
        # Check for common OCR artifacts
        ocr_artifacts = ['�', '|||', '...', '___']
        if any(artifact in text for artifact in ocr_artifacts):
            quality_score -= 1
            quality_indicators.append("possible_ocr_artifacts")
        
        if quality_score >= 6:
            quality = "high"
        elif quality_score >= 4:
            quality = "medium" 
        else:
            quality = "low"
        
        return {
            "quality": quality,
            "score": quality_score,
            "char_count": char_count,
            "word_count": word_count,
            "line_count": line_count,
            "avg_word_length": avg_word_length,
            "indicators": quality_indicators
        }
    
    def _llm_guided_extraction(self, file_path: Path, file_id: str, previous_attempt: str) -> List[Document]:
        """Use LLM to guide extraction when standard methods fail"""
        
        react_prompt = f"""
You are an expert document processor using ReACT (Reasoning and Acting) methodology.

Previous extraction attempt:
{previous_attempt}

File: {file_path.name}
Type: {file_path.suffix.lower()}
Size: {file_path.stat().st_size} bytes

Think step by step about why the extraction might have failed and what alternative approaches to try:

THOUGHT: [Analyze why extraction failed]
ACTION: [Recommend specific extraction strategy]
OBSERVATION: [What would indicate success/failure]

Available methods:
- PyMuPDF (good for complex PDFs)
- pdfplumber (good for tables and layouts)
- PyPDF2 (basic PDF extraction)
- OCR (for scanned documents)
- Manual analysis

Provide your reasoning and recommended next steps.
"""
        
        try:
            response = self.llm.invoke(react_prompt)
            reasoning = response.content
            
            # Extract recommended action from LLM response
            if "OCR" in reasoning.upper():
                # Try OCR approach (placeholder for future OCR integration)
                return self._attempt_ocr_extraction(file_path, file_id, reasoning)
            elif "PYMUPDF" in reasoning.upper():
                return self._retry_extraction_with_method(file_path, file_id, "pymupdf", reasoning)
            elif "PDFPLUMBER" in reasoning.upper():
                return self._retry_extraction_with_method(file_path, file_id, "pdfplumber", reasoning)
            else:
                # Default retry with different approach
                return self._retry_extraction_with_method(file_path, file_id, "text", reasoning)
                
        except Exception as e:
            logger.error(f"LLM-guided extraction failed: {e}")
            return []
    
    def _retry_extraction_with_method(self, file_path: Path, file_id: str, method: str, reasoning: str) -> List[Document]:
        """Retry extraction with specific method and reasoning"""
        try:
            result = self.extraction_tool.forward(
                file_path=str(file_path),
                extraction_method=method,
                reasoning=f"LLM-guided retry: {reasoning}"
            )
            
            if "EXTRACTED_TEXT:" in result:
                text_content = result.split("EXTRACTED_TEXT:")[1].strip()
                if text_content and not text_content.startswith("ERROR:"):
                    doc = Document(
                        page_content=text_content,
                        metadata={
                            "source": file_id,
                            "file_path": str(file_path),
                            "file_type": file_path.suffix.lower(),
                            "extraction_method": f"react_llm_guided_{method}",
                            "llm_reasoning": reasoning,
                            "extraction_quality": self._assess_extraction_quality(text_content)
                        }
                    )
                    return [doc]
        except Exception as e:
            logger.error(f"Retry extraction failed: {e}")
        
        return []
    
    def _attempt_ocr_extraction(self, file_path: Path, file_id: str, reasoning: str) -> List[Document]:
        """Placeholder for OCR extraction (can be extended with pytesseract/easyocr)"""
        logger.info(f"OCR extraction recommended for {file_path.name}")
        
        # For now, return empty - can be extended with actual OCR
        doc = Document(
            page_content="OCR extraction not yet implemented. This document may be scanned/image-based.",
            metadata={
                "source": file_id,
                "file_path": str(file_path),
                "file_type": file_path.suffix.lower(),
                "extraction_method": "react_ocr_placeholder",
                "llm_reasoning": reasoning,
                "needs_ocr": True
            }
        )
        return [doc]

class RetrieverTool(Tool):
    """Semantic similarity retrieval tool for document search"""
    name = "retriever"
    description = "Using semantic similarity, retrieves some documents from the knowledge base that have the closest embeddings to the input query."
    inputs = {
        "query": {
            "type": "string",
            "description": "The query to perform. This should be semantically close to your target documents. Use the affirmative form rather than a question.",
        }
    }
    output_type = "string"

    def __init__(self, vectordb: VectorStore, **kwargs):
        super().__init__(**kwargs)
        self.vectordb = vectordb

    def forward(self, query: str) -> str:
        assert isinstance(query, str), "Your search query must be a string"

        docs = self.vectordb.similarity_search(
            query,
            k=7,
        )

        return "\nRetrieved documents:\n" + "".join(
            [
                f"===== Document {str(i)} =====\n" + doc.page_content
                for i, doc in enumerate(docs)
            ]
        )

class DocumentRAGState(TypedDict):
    """State for the document RAG agent"""
    query: str
    documents: List[Document]
    context: str
    response: str
    chat_history: List[Dict[str, str]]
    file_ids: List[str]

class DocumentRAGAgent:
    """Agentic RAG system for document chat (PDF, TXT, DOCX)"""
    
    def __init__(self, 
                 embedding_model_name: str = "thenlper/gte-small",
                 chunk_size: int = 1000,
                 chunk_overlap: int = 200,
                 vectorstore_dir: str = None):
        
        self.embedding_model_name = embedding_model_name
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        
        # Use settings for vectorstore directory
        if vectorstore_dir is None:
            self.vectorstore_dir = settings.DATA_DIR / "vectorstores"
        else:
            self.vectorstore_dir = Path(vectorstore_dir)
        
        self.vectorstore_dir.mkdir(parents=True, exist_ok=True)
        
        # Initialize components
        self._init_embeddings()
        self._init_text_splitter()
        self._init_llm()
        self._init_prompts()
        
        # Document storage
        self.vectorstores: Dict[str, FAISS] = {}
        self.document_metadata: Dict[str, Dict] = {}
        
        # Initialize agent graph if LangGraph is available
        if LANGGRAPH_AVAILABLE:
            self._init_agent_graph()
    
    def _init_embeddings(self):
        """Initialize embedding model"""
        try:
            self.embeddings = HuggingFaceEmbeddings(
                model_name=self.embedding_model_name,
                model_kwargs={'device': 'cpu'},
                encode_kwargs={'normalize_embeddings': True}
            )
            logger.info(f"Initialized embeddings: {self.embedding_model_name}")
        except Exception as e:
            logger.error(f"Failed to initialize embeddings: {e}")
            raise
    
    def _init_text_splitter(self):
        """Initialize text splitter with tokenizer"""
        try:
            # Use tokenizer-aware splitter for better chunking (updated approach)
            tokenizer = AutoTokenizer.from_pretrained(self.embedding_model_name)
            self.text_splitter = RecursiveCharacterTextSplitter.from_huggingface_tokenizer(
                tokenizer,
                chunk_size=200,  # Smaller chunks for better retrieval
                chunk_overlap=20,
                add_start_index=True,
                strip_whitespace=True,
                separators=["\n\n", "\n", ".", " ", ""],
            )
            logger.info("Initialized tokenizer-aware text splitter with optimized settings")
        except Exception as e:
            logger.warning(f"Failed to initialize tokenizer-aware splitter, using default: {e}")
            # Fallback to regular text splitter
            self.text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=self.chunk_size,
                chunk_overlap=self.chunk_overlap,
                add_start_index=True,
                strip_whitespace=True,
                separators=["\n\n", "\n", ".", " ", ""]
            )
    
    def _init_llm(self):
        """Initialize Azure OpenAI LLM"""
        try:
            # Initialize Azure OpenAI LLM
            self.llm = AzureChatOpenAI(
                azure_endpoint=settings.AZURE_OPENAI_ENDPOINT,
                azure_deployment=settings.AZURE_OPENAI_DEPLOYMENT,
                api_key=settings.AZURE_OPENAI_KEY,
                api_version=settings.AZURE_OPENAI_API_VERSION,
                temperature=settings.LLM_TEMPERATURE,
                max_tokens=settings.MAX_TOKENS
            )
            logger.info("Initialized Azure OpenAI LLM")
        except Exception as e:
            logger.error(f"Failed to initialize Azure OpenAI LLM: {e}")
            self.llm = None
    
    def _init_prompts(self):
        """Initialize prompt templates"""
        self.qa_prompt = PromptTemplate(
            template="""You are a helpful assistant that answers questions based on the provided document context.

Context from documents:
{context}

Chat History:
{chat_history}

Question: {question}

Instructions:
- Answer the question based only on the provided context
- If the answer is not in the context, say "I don't have enough information to answer this question based on the provided documents"
- Be specific and cite relevant information from the documents
- Maintain a conversational tone while being informative

Answer:""",
            input_variables=["context", "chat_history", "question"]
        )
        
        self.context_prompt = PromptTemplate(
            template="""Based on the following documents, provide relevant context for the question: {question}

Documents:
{documents}

Relevant Context:""",
            input_variables=["question", "documents"]
        )
    
    def _init_agent_graph(self):
        """Initialize LangGraph agent"""
        try:
            # Define the agent workflow
            workflow = StateGraph(DocumentRAGState)
            
            # Add nodes
            workflow.add_node("retrieve_context", self._retrieve_context_node)
            workflow.add_node("generate_response", self._generate_response_node)
            
            # Define edges
            workflow.add_edge(START, "retrieve_context")
            workflow.add_edge("retrieve_context", "generate_response")
            
            # Compile the graph
            self.agent_graph = workflow.compile()
            logger.info("Initialized LangGraph agent")
            
        except Exception as e:
            logger.warning(f"Failed to initialize LangGraph agent: {e}")
            self.agent_graph = None
    
    def load_document(self, file_path: str, file_id: str) -> List[Document]:
        """Load document using ReACT reasoning approach"""
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        logger.info(f"Loading document with ReACT: {file_path} (type: {file_ext})")
        logger.info(f"File size: {file_path.stat().st_size} bytes")
        
        # Try ReACT processor first if available
        if hasattr(self, 'react_processor') and self.react_processor:
            try:
                logger.info("Using ReACT processor for intelligent extraction")
                docs = self.react_processor.process_document_with_react(str(file_path), file_id)
                if docs:
                    logger.info(f"ReACT processor successfully extracted {len(docs)} documents")
                    return docs
                else:
                    logger.warning("ReACT processor failed, falling back to standard methods")
            except Exception as e:
                logger.warning(f"ReACT processor error: {e}, falling back to standard methods")
        else:
            logger.info("ReACT processor not available, using standard extraction")
        
        # Fallback to standard extraction methods
        try:
            docs = []
            
            if file_ext == '.txt':
                logger.info("Loading as text file")
                loader = TextLoader(str(file_path), encoding='utf-8')
                docs = loader.load()
            elif file_ext == '.pdf':
                logger.info("Loading as PDF file")
                # Try our enhanced PDF loading first
                docs = self._load_pdf_enhanced(file_path, file_id)
                
                # If that fails, try the original loader
                if not docs and PyPDFLoader:
                    logger.info("Trying PyPDFLoader as fallback")
                    try:
                        loader = PyPDFLoader(str(file_path))
                        docs = loader.load()
                        logger.info(f"PyPDFLoader loaded {len(docs)} documents")
                    except Exception as e:
                        logger.warning(f"PyPDFLoader failed: {e}")
                        
            elif file_ext in ['.docx', '.doc']:
                logger.info("Loading as DOCX file")
                if Docx2txtLoader:
                    logger.info("Using Docx2txtLoader")
                    loader = Docx2txtLoader(str(file_path))
                    docs = loader.load()
                else:
                    logger.info("Using DOCX fallback")
                    docs = self._load_docx_fallback(file_path)
            elif file_ext == '.md':
                logger.info("Loading as Markdown file")
                loader = TextLoader(str(file_path), encoding='utf-8')
                docs = loader.load()
            else:
                raise ValueError(f"Unsupported file type: {file_ext}")
            
            logger.info(f"Standard extraction result: {len(docs)} documents")
            
            # Debug: Log content details for each document
            for i, doc in enumerate(docs):
                content_length = len(doc.page_content) if doc.page_content else 0
                logger.info(f"Document {i}: content_length={content_length}")
                if content_length > 0:
                    preview = doc.page_content[:100].replace('\n', ' ').strip()
                    logger.info(f"Document {i} preview: {preview}...")
                else:
                    logger.warning(f"Document {i} has empty content")
            
            # Filter out empty documents and add metadata
            valid_docs = []
            for i, doc in enumerate(docs):
                if doc.page_content and doc.page_content.strip():
                    doc.metadata.update({
                        "source": file_id,
                        "file_path": str(file_path),
                        "file_type": file_ext,
                        "chunk_id": i,
                        "extraction_method": "standard"
                    })
                    valid_docs.append(doc)
                    logger.info(f"Added valid document {i} with {len(doc.page_content)} characters")
                else:
                    logger.warning(f"Skipping empty document at index {i}")
            
            logger.info(f"Loaded {len(valid_docs)} valid documents from {file_id}")
            return valid_docs
            
        except Exception as e:
            logger.error(f"Failed to load document {file_id}: {e}")
            logger.error(f"File path: {file_path}")
            logger.error(f"File exists: {file_path.exists()}")
            logger.error(f"File size: {file_path.stat().st_size if file_path.exists() else 'N/A'}")
            
            # Final fallback: try text loading
            try:
                logger.info(f"Attempting final fallback text loading for {file_id}")
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                if content.strip():
                    logger.info(f"Fallback successful, content length: {len(content)}")
                    return [Document(
                        page_content=content,
                        metadata={
                            "source": file_id,
                            "file_path": str(file_path),
                            "file_type": file_ext,
                            "loader": "fallback_text",
                            "extraction_method": "fallback"
                        }
                    )]
                else:
                    logger.warning("Fallback text loading returned empty content")
            except Exception as fallback_error:
                logger.error(f"Fallback loading also failed for {file_id}: {fallback_error}")
            
            return []
    
    def _load_pdf_enhanced(self, file_path: Path, file_id: str) -> List[Document]:
        """Enhanced PDF loader that tries all available methods"""
        docs = []
        
        logger.info(f"Enhanced PDF loading for {file_id}")
        logger.info(f"Available processors: PyMuPDF={PYMUPDF_AVAILABLE}, pdfplumber={PDFPLUMBER_AVAILABLE}, PyPDF2={PYPDF2_AVAILABLE}")
        
        # Try PyMuPDF first (most reliable)
        if PYMUPDF_AVAILABLE:
            try:
                logger.info("Attempting PDF extraction with PyMuPDF")
                pdf_document = fitz.open(str(file_path))
                logger.info(f"PDF has {len(pdf_document)} pages")
                
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    text = page.get_text()
                    logger.info(f"Page {page_num + 1}: extracted {len(text)} characters")
                    
                    if text.strip():
                        doc = Document(
                            page_content=text,
                            metadata={"page": page_num + 1, "loader": "PyMuPDF"}
                        )
                        docs.append(doc)
                        logger.info(f"Added page {page_num + 1} with content preview: {text[:100]}...")
                    else:
                        logger.warning(f"Page {page_num + 1} is empty")
                        
                pdf_document.close()
                logger.info(f"PyMuPDF extracted {len(docs)} non-empty pages")
                
                if docs:
                    return docs
                else:
                    logger.warning("PyMuPDF found no content")
                    
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        # Try pdfplumber second
        if PDFPLUMBER_AVAILABLE:
            try:
                logger.info("Attempting PDF extraction with pdfplumber")
                with pdfplumber.open(file_path) as pdf:
                    logger.info(f"PDF has {len(pdf.pages)} pages")
                    
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        logger.info(f"Page {page_num + 1}: extracted {len(text) if text else 0} characters")
                        
                        if text and text.strip():
                            doc = Document(
                                page_content=text,
                                metadata={"page": page_num + 1, "loader": "pdfplumber"}
                            )
                            docs.append(doc)
                            logger.info(f"Added page {page_num + 1} with content preview: {text[:100]}...")
                        else:
                            logger.warning(f"Page {page_num + 1} is empty or None")
                            
                logger.info(f"pdfplumber extracted {len(docs)} non-empty pages")
                
                if docs:
                    return docs
                else:
                    logger.warning("pdfplumber found no content")
                    
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        
        # Try PyPDF2 as final fallback
        if PYPDF2_AVAILABLE:
            try:
                logger.info("Attempting PDF extraction with PyPDF2")
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    logger.info(f"PDF has {len(pdf_reader.pages)} pages")
                    
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        logger.info(f"Page {page_num + 1}: extracted {len(text)} characters")
                        
                        if text.strip():
                            doc = Document(
                                page_content=text,
                                metadata={"page": page_num + 1, "loader": "PyPDF2"}
                            )
                            docs.append(doc)
                            logger.info(f"Added page {page_num + 1} with content preview: {text[:100]}...")
                        else:
                            logger.warning(f"Page {page_num + 1} is empty")
                            
                logger.info(f"PyPDF2 extracted {len(docs)} non-empty pages")
                
                if docs:
                    return docs
                else:
                    logger.warning("PyPDF2 found no content")
                    
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {e}")
        
        logger.error("All PDF extraction methods failed or found no content")
        return []

    def _load_pdf_fallback(self, file_path: Path) -> List[Document]:
        """Enhanced PDF loader with multiple fallback options"""
        docs = []
        
        # Try PyMuPDF first (most reliable)
        if PYMUPDF_AVAILABLE:
            try:
                logger.info("Attempting PDF extraction with PyMuPDF")
                pdf_document = fitz.open(str(file_path))
                for page_num in range(len(pdf_document)):
                    page = pdf_document.load_page(page_num)
                    text = page.get_text()
                    if text.strip():
                        doc = Document(
                            page_content=text,
                            metadata={"page": page_num + 1, "loader": "PyMuPDF"}
                        )
                        docs.append(doc)
                pdf_document.close()
                logger.info(f"PyMuPDF extracted {len(docs)} pages")
                return docs
            except Exception as e:
                logger.warning(f"PyMuPDF failed: {e}")
        
        # Try pdfplumber second
        if PDFPLUMBER_AVAILABLE:
            try:
                logger.info("Attempting PDF extraction with pdfplumber")
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages):
                        text = page.extract_text()
                        if text and text.strip():
                            doc = Document(
                                page_content=text,
                                metadata={"page": page_num + 1, "loader": "pdfplumber"}
                            )
                            docs.append(doc)
                logger.info(f"pdfplumber extracted {len(docs)} pages")
                return docs
            except Exception as e:
                logger.warning(f"pdfplumber failed: {e}")
        
        # Try PyPDF2 as final fallback
        if PYPDF2_AVAILABLE:
            try:
                logger.info("Attempting PDF extraction with PyPDF2")
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page_num, page in enumerate(pdf_reader.pages):
                        text = page.extract_text()
                        if text.strip():
                            doc = Document(
                                page_content=text,
                                metadata={"page": page_num + 1, "loader": "PyPDF2"}
                            )
                            docs.append(doc)
                logger.info(f"PyPDF2 extracted {len(docs)} pages")
                return docs
            except Exception as e:
                logger.error(f"PyPDF2 also failed: {e}")
        
        logger.error("All PDF extraction methods failed")
        return []
    
    def _load_docx_fallback(self, file_path: Path) -> List[Document]:
        """Fallback DOCX loader using python-docx"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            if text.strip():
                return [Document(page_content=text, metadata={})]
            return []
        except ImportError:
            logger.error("python-docx not available for DOCX fallback")
            return []
        except Exception as e:
            logger.error(f"DOCX fallback failed: {e}")
            return []
    
    def load_document_from_base64(self, base64_content: str, file_id: str, file_type: str) -> List[Document]:
        """Load document from base64 encoded content"""
        try:
            logger.info(f"Loading document from base64: {file_id} (type: {file_type})")
            
            # Decode base64 content
            binary_content = base64.b64decode(base64_content)
            logger.info(f"Decoded {len(binary_content)} bytes from base64")
            
            # Create temporary file
            temp_dir = self.vectorstore_dir / "temp"
            temp_dir.mkdir(exist_ok=True)
            temp_file = temp_dir / f"{file_id}_temp{file_type}"
            
            try:
                # Write binary content to temporary file
                with open(temp_file, 'wb') as f:
                    f.write(binary_content)
                
                logger.info(f"Created temporary file: {temp_file}")
                
                # Load document using existing methods
                docs = self.load_document(str(temp_file), file_id)
                
                # Clean up temporary file
                temp_file.unlink()
                
                return docs
                
            except Exception as e:
                # Clean up temporary file on error
                if temp_file.exists():
                    temp_file.unlink()
                raise e
                
        except Exception as e:
            logger.error(f"Failed to load document from base64: {e}")
            return []

    def encode_document_to_base64(self, file_path: str) -> str:
        """Encode document file to base64"""
        try:
            with open(file_path, 'rb') as f:
                binary_content = f.read()
            encoded_content = base64.b64encode(binary_content).decode('utf-8')
            logger.info(f"Encoded {len(binary_content)} bytes to base64 ({len(encoded_content)} chars)")
            return encoded_content
        except Exception as e:
            logger.error(f"Failed to encode document to base64: {e}")
            return ""

    def process_document_from_base64(self, file_id: str, base64_content: str, file_type: str) -> bool:
        """Process document from base64 encoded content for RAG"""
        try:
            logger.info(f"Processing document from base64: {file_id}")
            
            # Load documents from base64
            docs = self.load_document_from_base64(base64_content, file_id, file_type)
            
            if not docs:
                logger.error(f"No documents loaded from base64 for {file_id}")
                return False
            
            # Continue with standard processing
            return self._process_documents_for_rag(file_id, docs)
            
        except Exception as e:
            logger.error(f"Failed to process document from base64 {file_id}: {e}")
            return False

    def _process_documents_for_rag(self, file_id: str, docs: List[Document]) -> bool:
        """Common document processing logic for RAG"""
        try:
            # Sanitize file_id for safe filesystem operations
            safe_file_id = sanitize_filename(file_id)
            logger.info(f"Processing {len(docs)} documents for RAG: {file_id} -> {safe_file_id}")
            
            # Split documents into chunks with improved processing
            logger.info(f"Splitting documents for {file_id}...")
            docs_processed = []
            unique_texts = {}
            
            print("Splitting documents...")
            for doc in tqdm(docs, desc="Processing documents"):
                try:
                    new_docs = self.text_splitter.split_documents([doc])
                    for new_doc in new_docs:
                        if new_doc.page_content and new_doc.page_content.strip():
                            text_content = new_doc.page_content.strip()
                            if text_content not in unique_texts:
                                unique_texts[text_content] = True
                                docs_processed.append(new_doc)
                except Exception as e:
                    logger.warning(f"Error splitting document chunk: {e}")
                    continue
            
            if not docs_processed:
                logger.error(f"No valid chunks created for {file_id}")
                return False
            
            logger.info(f"Created {len(docs_processed)} unique chunks for {file_id}")
            
            # Create vector store with improved embedding
            logger.info(f"Embedding documents... This may take a few minutes")
            try:
                vectorstore = FAISS.from_documents(
                    documents=docs_processed,
                    embedding=self.embeddings,
                    distance_strategy=DistanceStrategy.COSINE,
                )
            except Exception as e:
                logger.error(f"Failed to create vector store: {e}")
                return False
            
            # Save vector store using sanitized filename
            vectorstore_path = self.vectorstore_dir / f"{safe_file_id}_vectorstore"
            logger.info(f"Saving vector store to: {vectorstore_path}")
            
            # Ensure the parent directory exists
            vectorstore_path.mkdir(parents=True, exist_ok=True)
            
            try:
                vectorstore.save_local(str(vectorstore_path))
                logger.info(f"Successfully saved vector store to {vectorstore_path}")
            except Exception as e:
                logger.error(f"Failed to save vector store: {e}")
                return False
            
            # Store in memory using original file_id as key
            self.vectorstores[file_id] = vectorstore
            self.document_metadata[file_id] = {
                "num_chunks": len(docs_processed),
                "vectorstore_path": str(vectorstore_path),
                "safe_file_id": safe_file_id,
                "processed": True
            }
            
            # Create marker file using sanitized filename
            marker_file = self.vectorstore_dir / f"{safe_file_id}_processed.marker"
            marker_data = {
                "file_id": file_id,
                "safe_file_id": safe_file_id,
                "processed": True,
                "num_chunks": len(docs_processed),
                "timestamp": time.time()
            }
            try:
                marker_file.write_text(json.dumps(marker_data))
            except Exception as e:
                logger.warning(f"Failed to create marker file: {e}")
            
            logger.info(f"Successfully processed {file_id} for RAG")
            return True
            
        except Exception as e:
            logger.error(f"Failed to process documents for RAG {file_id}: {e}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return False

    def process_document_for_rag(self, file_id: str, content: str = None, schema: Dict = None) -> bool:
        """Process document for RAG"""
        try:
            logger.info(f"Processing document {file_id} for RAG")
            
            # If content is provided directly, use it
            if content:
                logger.info(f"Using provided content for {file_id}")
                docs = [Document(page_content=content, metadata={"source": file_id})]
            else:
                # Find and load the document file
                logger.info(f"Looking for file with file_id: {file_id}")
                logger.info(f"Upload directory: {settings.UPLOAD_DIR}")
                
                file_path = None
                allowed_extensions = ['.pdf', '.txt', '.docx', '.doc', '.md']
                
                # List all files in upload directory for debugging
                if settings.UPLOAD_DIR.exists():
                    all_files = list(settings.UPLOAD_DIR.iterdir())
                    logger.info(f"All files in upload directory: {[f.name for f in all_files]}")
                else:
                    logger.error(f"Upload directory does not exist: {settings.UPLOAD_DIR}")
                    return False
                
                for ext in allowed_extensions:
                    potential_path = settings.UPLOAD_DIR / f"{file_id}{ext}"
                    logger.info(f"Checking for file: {potential_path}")
                    if potential_path.exists():
                        file_path = potential_path
                        logger.info(f"Found file: {file_path}")
                        break
                
                if not file_path:
                    logger.error(f"No supported file found for {file_id}")
                    logger.error(f"Searched for: {[f'{file_id}{ext}' for ext in allowed_extensions]}")
                    return False
                
                logger.info(f"Loading document from: {file_path}")
                docs = self.load_document(str(file_path), file_id)
                logger.info(f"Loaded {len(docs)} documents")
            
            if not docs:
                logger.error(f"No documents loaded for {file_id}")
                return False
            
            # Use common processing logic
            return self._process_documents_for_rag(file_id, docs)
            
        except Exception as e:
            logger.error(f"Failed to process document {file_id} for RAG: {e}")
            import traceback
            logger.error(f"Full traceback: {traceback.format_exc()}")
            return False
    
    def _load_vectorstore(self, file_id: str) -> Optional[FAISS]:
        """Load vector store for a file"""
        try:
            if file_id in self.vectorstores:
                return self.vectorstores[file_id]
            
            # Try with sanitized filename
            safe_file_id = sanitize_filename(file_id)
            
            # Try original filename first
            vectorstore_path = self.vectorstore_dir / f"{file_id}_vectorstore"
            if not vectorstore_path.exists():
                # Try sanitized filename
                vectorstore_path = self.vectorstore_dir / f"{safe_file_id}_vectorstore"
            
            if vectorstore_path.exists():
                logger.info(f"Loading vectorstore from: {vectorstore_path}")
                vectorstore = FAISS.load_local(
                    str(vectorstore_path),
                    self.embeddings,
                    allow_dangerous_deserialization=True
                )
                self.vectorstores[file_id] = vectorstore
                return vectorstore
            
            logger.warning(f"No vectorstore found for {file_id} (tried both original and sanitized paths)")
            return None
        except Exception as e:
            logger.error(f"Failed to load vectorstore for {file_id}: {e}")
            return None
    
    def _retrieve_context_node(self, state: DocumentRAGState) -> DocumentRAGState:
        """Retrieve relevant context from documents"""
        query = state["query"]
        file_ids = state.get("file_ids", [])
        
        all_docs = []
        
        for file_id in file_ids:
            vectorstore = self._load_vectorstore(file_id)
            if vectorstore:
                # Retrieve relevant documents
                relevant_docs = vectorstore.similarity_search(query, k=5)
                all_docs.extend(relevant_docs)
        
        # Combine context
        context = "\n\n".join([doc.page_content for doc in all_docs])
        
        state["documents"] = all_docs
        state["context"] = context
        return state
    
    def _generate_response_node(self, state: DocumentRAGState) -> DocumentRAGState:
        """Generate response using LLM"""
        query = state["query"]
        context = state["context"]
        chat_history = state.get("chat_history", [])
        
        # Format chat history
        history_text = ""
        if chat_history:
            for exchange in chat_history[-3:]:  # Last 3 exchanges
                if "user" in exchange and "assistant" in exchange:
                    history_text += f"User: {exchange['user']}\nAssistant: {exchange['assistant']}\n\n"
        
        if self.llm and context:
            try:
                # Use LLM to generate response
                prompt = self.qa_prompt.format(
                    context=context,
                    chat_history=history_text,
                    question=query
                )
                response = self.llm.invoke(prompt)
                state["response"] = response.content
            except Exception as e:
                logger.error(f"LLM generation failed: {e}")
                state["response"] = f"I found relevant information in the documents, but encountered an error generating a response: {str(e)}\n\nRelevant context:\n{context[:500]}..."
        else:
            # Fallback: return context directly
            if context:
                state["response"] = f"Based on the documents, here's what I found:\n\n{context[:1000]}..."
            else:
                state["response"] = "I couldn't find relevant information in the provided documents to answer your question."
        
        return state
    
    def chat_with_documents(self, 
                          query: str, 
                          file_ids: List[str], 
                          context: Dict[str, Any] = None,
                          chat_history: List[Dict[str, str]] = None) -> str:
        """Enhanced chat with documents using agentic RAG approach"""
        try:
            if self.agent_graph:
                # Use LangGraph agent
                state = DocumentRAGState(
                    query=query,
                    documents=[],
                    context="",
                    response="",
                    chat_history=chat_history or [],
                    file_ids=file_ids
                )
                
                result = self.agent_graph.invoke(state)
                return result["response"]
            else:
                # Enhanced fallback with retrieval tool approach
                all_docs = []
                combined_vectorstore = None
                
                for file_id in file_ids:
                    vectorstore = self._load_vectorstore(file_id)
                    if vectorstore:
                        if combined_vectorstore is None:
                            combined_vectorstore = vectorstore
                        else:
                            # Merge vectorstores for multi-document retrieval
                            relevant_docs = vectorstore.similarity_search(query, k=5)
                            all_docs.extend(relevant_docs)
                
                if not combined_vectorstore and not all_docs:
                    return "I couldn't find any relevant information in the specified documents."
                
                # Use retrieval tool approach for better context gathering
                if SMOLAGENTS_AVAILABLE and combined_vectorstore:
                    try:
                        retriever_tool = RetrieverTool(combined_vectorstore)
                        context_text = retriever_tool.forward(query)
                    except Exception as e:
                        logger.warning(f"Retrieval tool failed, using standard approach: {e}")
                        relevant_docs = combined_vectorstore.similarity_search(query, k=7)
                        context_text = "\n\n".join([doc.page_content for doc in relevant_docs])
                else:
                    # Standard retrieval
                    if combined_vectorstore:
                        relevant_docs = combined_vectorstore.similarity_search(query, k=7)
                    else:
                        relevant_docs = all_docs[:7]
                    context_text = "\n\n".join([doc.page_content for doc in relevant_docs])
                
                # Format chat history
                history_text = ""
                if chat_history:
                    for exchange in chat_history[-3:]:
                        if "user" in exchange and "assistant" in exchange:
                            history_text += f"User: {exchange['user']}\nAssistant: {exchange['assistant']}\n\n"
                
                if self.llm:
                    try:
                        # Enhanced prompt for better responses
                        prompt = f"""Given the question and supporting documents below, give a comprehensive answer to the question.
Respond only to the question asked, response should be concise and relevant to the question.
Provide the number of the source document when relevant.
If you cannot find information, do not give up and try to provide the best answer based on available context.

Previous conversation:
{history_text}

Question:
{query}

{context_text}

Answer:"""
                        
                        response = self.llm.invoke(prompt)
                        return response.content
                    except Exception as e:
                        logger.error(f"LLM generation failed: {e}")
                        return f"I found relevant information but encountered an error: {str(e)}\n\nContext: {context_text[:500]}..."
                else:
                    return f"Based on the documents:\n\n{context_text[:1000]}..."
                    
        except Exception as e:
            logger.error(f"Chat with documents failed: {e}")
            return f"I encountered an error while processing your request: {str(e)}"

def create_document_rag_system() -> DocumentRAGAgent:
    """Factory function to create document RAG system"""
    try:
        return DocumentRAGAgent()
    except Exception as e:
        logger.error(f"Failed to create document RAG system: {e}")
        raise